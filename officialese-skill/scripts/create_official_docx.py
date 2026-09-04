#!/usr/bin/env python
"""Create a starter DOCX using the uploaded SOE official-document format."""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from embed_fonts import embed_bundled_fonts  # noqa: E402  随技能分发的字体嵌入器


TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
KAITI_FONT = "楷体_GB2312"
HEITI_FONT = "黑体"
SONGTI_FONT = "宋体"
# 西文字母与阿拉伯数字统一 Times New Roman；中文走 eastAsia，Word 在同一 run 内
# 按字符类型自动分派，无需拆分 run。页码是例外，见 add_page_field。
WESTERN_FONT = "Times New Roman"

TITLE_SIZE = 22  # 二号
BODY_SIZE = 16  # 三号
PAGE_NUMBER_SIZE = 14  # 四号
BODY_LINE_PT = 28
TITLE_LINE_PT = 30
# 版式以「字」为单位：一个全角汉字宽 = 正文字号 = 三号 16 磅。
CHAR_PT = BODY_SIZE
TWO_CHAR_INDENT_PT = BODY_SIZE * 2
SIGNATURE_RIGHT_INDENT_PT = BODY_SIZE * 4

PAGE_WIDTH_CM = 21.0
LEFT_MARGIN_CM = 2.8
RIGHT_MARGIN_CM = 2.6
TEXT_WIDTH_PT = Cm(PAGE_WIDTH_CM - LEFT_MARGIN_CM - RIGHT_MARGIN_CM).pt  # 版心宽 ≈ 442.2 磅

ATTACHMENT_LABEL = "附件："
ATTACHMENT_LABEL_CHARS = 3.0   # 「附件：」占 3 个字
ATTACHMENT_START_CHARS = 2.0   # 附件说明起行左空 2 字，与正文一致

STYLE_TITLE = "Official Title"
STYLE_SUBTITLE = "Official Subtitle"
STYLE_BODY = "Official Body"
STYLE_H1 = "Official Heading 1"
STYLE_H2 = "Official Heading 2"
STYLE_H3 = "Official Heading 3"
STYLE_H4 = "Official Heading 4"


def set_east_asia_font(obj, font_name: str, western_font: str | None = None) -> None:
    """CJK 用 ``font_name``；西文字母与阿拉伯数字用 ``western_font``（默认
    Times New Roman）。Word 依 rFonts 在同一 run 内按字符类型分派字体，因此
    中英混排无需拆分 run。

    传 ``western_font=font_name`` 可让整个 run 都用中文字体——页码即如此：
    GB/T 9704 规定页码为四号宋体阿拉伯数字，数字不走 Times。"""
    western = western_font or WESTERN_FONT
    rpr = obj._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), western)
    rfonts.set(qn("w:hAnsi"), western)
    rfonts.set(qn("w:cs"), western)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False,
                 western_font: str | None = None) -> None:
    run.font.name = western_font or WESTERN_FONT
    set_east_asia_font(run, font_name, western_font)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_style_font(style, font_name: str, size_pt: int, bold: bool = False,
                   western_font: str | None = None) -> None:
    style.font.name = western_font or WESTERN_FONT
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    set_east_asia_font(style, font_name, western_font)


def set_paragraph_format(
    paragraph,
    *,
    line_pt: int = BODY_LINE_PT,
    first_indent: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    right_indent_pt: float = 0,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(TWO_CHAR_INDENT_PT) if first_indent else Pt(0)
    fmt.right_indent = Pt(right_indent_pt)
    paragraph.alignment = alignment


def configure_style(
    doc: Document,
    name: str,
    font_name: str,
    size_pt: int,
    *,
    bold: bool = False,
    line_pt: int = BODY_LINE_PT,
    first_indent: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    styles = doc.styles
    style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(style, font_name, size_pt, bold)
    fmt = style.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(TWO_CHAR_INDENT_PT) if first_indent else Pt(0)
    style.paragraph_format.alignment = alignment


def configure_styles(doc: Document) -> None:
    set_style_font(doc.styles["Normal"], BODY_FONT, BODY_SIZE, False)
    configure_style(
        doc,
        STYLE_TITLE,
        TITLE_FONT,
        TITLE_SIZE,
        line_pt=TITLE_LINE_PT,
        first_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    configure_style(
        doc,
        STYLE_SUBTITLE,
        KAITI_FONT,
        BODY_SIZE,
        line_pt=BODY_LINE_PT,
        first_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    configure_style(doc, STYLE_BODY, BODY_FONT, BODY_SIZE)
    configure_style(doc, STYLE_H1, HEITI_FONT, BODY_SIZE, bold=False)
    configure_style(doc, STYLE_H2, KAITI_FONT, BODY_SIZE, bold=True)
    configure_style(doc, STYLE_H3, BODY_FONT, BODY_SIZE, bold=True)
    configure_style(doc, STYLE_H4, BODY_FONT, BODY_SIZE, bold=False)


def add_text_paragraph(
    doc: Document,
    text: str,
    *,
    font: str = BODY_FONT,
    bold: bool = False,
    style: str = STYLE_BODY,
    first_indent: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    right_indent_pt: float = 0,
) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_format(
        p,
        first_indent=first_indent,
        alignment=alignment,
        right_indent_pt=right_indent_pt,
    )
    run = p.add_run(text)
    set_run_font(run, font, BODY_SIZE, bold)


def add_center_line(doc: Document, text: str, font: str, size: int, style: str, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_format(
        p,
        line_pt=TITLE_LINE_PT if size == TITLE_SIZE else BODY_LINE_PT,
        first_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    run = p.add_run(text)
    set_run_font(run, font, size, bold)


def add_blank_line(doc: Document, line_pt: int = BODY_LINE_PT) -> None:
    p = doc.add_paragraph(style=STYLE_BODY)
    set_paragraph_format(p, line_pt=line_pt, first_indent=False)


def display_width(text: str) -> float:
    """文本宽度，单位为「字」：全角字符计 1，半角字符（阿拉伯数字、西文、
    半角点号）计 0.5。附件悬挂缩进与成文日期居中都按这个尺子算。"""
    width = 0.0
    for ch in text:
        width += 1.0 if unicodedata.east_asian_width(ch) in ("W", "F", "A") else 0.5
    return width


def add_hanging_paragraph(
    doc: Document,
    text: str,
    *,
    left_chars: float,
    first_line_chars: float,
    font: str = BODY_FONT,
    bold: bool = False,
    style: str = STYLE_BODY,
    line_pt: int = BODY_LINE_PT,
):
    """悬挂缩进段落：首行从 ``first_line_chars`` 字处起排，回行后统一落在
    ``left_chars`` 字处。附件说明靠它实现「转行与首行名称对齐」。"""
    p = doc.add_paragraph(style=style)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.left_indent = Pt(left_chars * CHAR_PT)
    fmt.first_line_indent = Pt((first_line_chars - left_chars) * CHAR_PT)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, font, BODY_SIZE, bold)
    return p


def normalize_attachment_name(name: str) -> str:
    """附件名称不加书名号，后不加标点符号；顺带去掉调用方可能重复写入的
    「附件：」前缀和序号。"""
    cleaned = str(name).strip()
    while cleaned.startswith(ATTACHMENT_LABEL) or cleaned.startswith("附件:"):
        cleaned = cleaned.split("：", 1)[-1].split(":", 1)[-1].strip()
    cleaned = re.sub(r"^\d+\s*[.、．]\s*", "", cleaned)
    cleaned = cleaned.rstrip("。；;，,、.")
    if cleaned.startswith("《") and cleaned.endswith("》"):
        cleaned = cleaned[1:-1].strip().rstrip("。；;，,、.")
    return cleaned


def add_page_field(paragraph) -> None:
    # 页码整体走宋体（含阿拉伯数字），依 GB/T 9704「页码用四号半角宋体阿拉伯
    # 数字」；这是全文西文走 Times New Roman 的唯一例外。
    run = paragraph.add_run("-")
    set_run_font(run, SONGTI_FONT, PAGE_NUMBER_SIZE, western_font=SONGTI_FONT)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    result = OxmlElement("w:t")
    result.text = "1"

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    set_run_font(field_run, SONGTI_FONT, PAGE_NUMBER_SIZE, western_font=SONGTI_FONT)
    field_run._r.append(field_begin)
    field_run._r.append(instr)
    field_run._r.append(field_separate)
    field_run._r.append(result)
    field_run._r.append(field_end)

    run = paragraph.add_run("-")
    set_run_font(run, SONGTI_FONT, PAGE_NUMBER_SIZE, western_font=SONGTI_FONT)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(2.35)

    doc.settings.odd_and_even_pages_header_footer = True

    odd_footer = section.footer.paragraphs[0]
    odd_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(odd_footer)

    even_footer = section.even_page_footer.paragraphs[0]
    even_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_page_field(even_footer)


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 1:
        add_text_paragraph(doc, text, font=HEITI_FONT, bold=False, style=STYLE_H1)
    elif level == 2:
        add_text_paragraph(doc, text, font=KAITI_FONT, bold=True, style=STYLE_H2)
    elif level == 3:
        add_text_paragraph(doc, text, font=BODY_FONT, bold=True, style=STYLE_H3)
    else:
        add_text_paragraph(doc, text, font=BODY_FONT, bold=False, style=STYLE_H4)


def add_sections(doc: Document, sections: list[dict]) -> None:
    for section in sections:
        title = section.get("heading") or section.get("title")
        if title:
            add_heading(doc, title, int(section.get("level", 1)))
        for para in section.get("paragraphs", []):
            add_text_paragraph(doc, para)
        children = section.get("children", [])
        if children:
            add_sections(doc, children)


def add_attachments(doc: Document, attachments: list[str]) -> None:
    """附件说明。版式取自 assets/templates/文件字体格式.doc：

    - 「附件：」左空 2 字起排，与正文首行缩进齐；
    - 多个附件用阿拉伯数字序号，第 2 条起左空 5 字，序号与首条的「1.」对齐；
    - 每条名称回行时不顶格，悬挂对齐到本条名称的起始位置（首条为左空 6 字）；
    - 名称不加书名号，后不加标点符号。
    """
    names = [normalize_attachment_name(a) for a in attachments if str(a).strip()]
    if not names:
        return
    add_blank_line(doc)

    if len(names) == 1:
        # 「附件：」占 3 字，名称从第 5 字起排，回行悬挂在同一列。
        name_col = ATTACHMENT_START_CHARS + ATTACHMENT_LABEL_CHARS
        add_hanging_paragraph(
            doc,
            f"{ATTACHMENT_LABEL}{names[0]}",
            left_chars=name_col,
            first_line_chars=ATTACHMENT_START_CHARS,
        )
        return

    for idx, name in enumerate(names, start=1):
        serial = f"{idx}."
        serial_chars = display_width(serial)
        if idx == 1:
            # 首条与「附件：」同行：名称列 = 2 + 3 + 序号宽（常见为 6 字）。
            name_col = ATTACHMENT_START_CHARS + ATTACHMENT_LABEL_CHARS + serial_chars
            first_line_chars = ATTACHMENT_START_CHARS
            text = f"{ATTACHMENT_LABEL}{serial}{name}"
        else:
            # 后续各条序号左空 5 字，与首条序号对齐；名称列随本条序号宽度走。
            first_line_chars = ATTACHMENT_START_CHARS + ATTACHMENT_LABEL_CHARS
            name_col = first_line_chars + serial_chars
            text = f"{serial}{name}"
        add_hanging_paragraph(
            doc,
            text,
            left_chars=name_col,
            first_line_chars=first_line_chars,
        )


def add_date_line(doc: Document, date: str, issuer: str | None) -> None:
    """成文日期排在单位名称正下方一行的正中间。

    做法不是按字数估算日期宽度再右空若干字——日期里「2026年6月8日」这类中西
    文混排会被 Word 的中西文自动间距撑宽，估算必偏。改为把日期段落的左右缩进
    卡在单位名称占位的两端（右空 4 字，左缩进 = 版心宽 - 4 字 - 署名宽），再让
    Word 在这个盒子里居中：无论日期实际排多宽，都与署名同心。
    """
    right_pt = float(SIGNATURE_RIGHT_INDENT_PT)
    issuer_pt = display_width(issuer) * CHAR_PT if issuer else 0.0
    # 署名太短（撑不下日期）或没有署名时，退回 GB/T 9704 的右空 4 字。
    if not issuer or display_width(date) >= display_width(issuer):
        add_text_paragraph(
            doc,
            date,
            first_indent=False,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent_pt=right_pt,
        )
        return

    p = doc.add_paragraph(style=STYLE_BODY)
    set_paragraph_format(
        p,
        first_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        right_indent_pt=right_pt,
    )
    p.paragraph_format.left_indent = Pt(max(0.0, TEXT_WIDTH_PT - right_pt - issuer_pt))
    run = p.add_run(date)
    set_run_font(run, BODY_FONT, BODY_SIZE)


def add_signature_block(doc: Document, issuer: str | None, date: str | None) -> None:
    if not issuer and not date:
        return

    # 正文（或附件说明）之后空两行，再排发文机关署名与成文日期。
    add_blank_line(doc)
    add_blank_line(doc)
    if issuer:
        add_text_paragraph(
            doc,
            issuer,
            first_indent=False,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent_pt=SIGNATURE_RIGHT_INDENT_PT,
        )
    if date:
        add_date_line(doc, date, issuer)


def build_docx(data: dict, output: Path) -> None:
    doc = Document()
    configure_styles(doc)
    setup_document(doc)

    issuer_title = data.get("issuer_title")
    if issuer_title:
        add_center_line(doc, issuer_title, TITLE_FONT, TITLE_SIZE, STYLE_TITLE)

    add_center_line(doc, data.get("title", "关于XXXX的通知"), TITLE_FONT, TITLE_SIZE, STYLE_TITLE)

    subtitle = data.get("subtitle")
    if subtitle:
        add_center_line(doc, subtitle, KAITI_FONT, BODY_SIZE, STYLE_SUBTITLE)

    add_blank_line(doc)

    recipient = data.get("recipient")
    if recipient:
        add_text_paragraph(doc, recipient, first_indent=False)

    for para in data.get("body", []):
        add_text_paragraph(doc, para)

    add_sections(doc, data.get("sections", []))
    add_attachments(doc, data.get("attachments", []))
    add_signature_block(doc, data.get("issuer"), data.get("date"))

    doc.save(output)
    # 嵌入随附的可嵌入字体（仿宋_GB2312、楷体_GB2312），使交付件在未装这两款
    # 字体的机器上不掉字；方正小标宋许可禁止嵌入，自动跳过。校验不过则保留
    # 未嵌入版本，不影响生成。
    embed_bundled_fonts(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create an SOE-style Chinese official DOCX starter.")
    parser.add_argument("output", type=Path, help="Output .docx path")
    parser.add_argument("--input", type=Path, help="JSON document specification")
    parser.add_argument("--title", help="Main title")
    parser.add_argument("--issuer-title", help="Issuing unit line above title")
    parser.add_argument("--subtitle", help="Centered subtitle or department line")
    parser.add_argument("--recipient", help="Recipient line, ending with Chinese colon")
    parser.add_argument("--body", action="append", default=[], help="Body paragraph; repeat as needed")
    parser.add_argument("--attachment", action="append", default=[], help="Attachment name; repeat as needed")
    parser.add_argument("--issuer", help="Signature unit")
    parser.add_argument("--date", help="Chinese date, e.g. 2026年6月8日")
    parser.add_argument("--doc-type", help="Optional document type note, e.g. 通知/请示/报告/函")
    parser.add_argument("--direction", choices=["上行文", "平行文", "下行文", "默认"], help="Optional writing direction note")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.input:
        data = json.loads(args.input.read_text(encoding="utf-8-sig"))
    else:
        data = {}

    for key in ["title", "issuer_title", "subtitle", "recipient", "issuer", "date", "doc_type", "direction"]:
        value = getattr(args, key)
        if value:
            data[key] = value
    if args.body:
        data["body"] = args.body
    if args.attachment:
        data["attachments"] = args.attachment
    build_docx(data, args.output)


if __name__ == "__main__":
    main()
