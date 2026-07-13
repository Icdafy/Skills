#!/usr/bin/env python
"""Create a 投委会议题 DOCX from a JSON spec, using the company officialese format.

Every run sets ascii/hAnsi to Times New Roman and eastAsia to the role's Chinese
font, so digits and Latin text render in Times New Roman automatically, matching
the gold examples. Inline bold is marked with **...** in any text field.

JSON spec example:
{
  "title": "审议关于XX股份有限公司2026年第一次临时股东会参会及表决事项的议题",
  "submit_line": "提交部门/子公司：投管公司",
  "recipient": "投委会：",
  "intro": "XX股份有限公司（以下简称“XX”）为产投公司参股企业，……具体如下：",
  "blocks": [
    {"type": "h1", "text": "一、会议基本信息"},
    {"type": "para", "text": "会议名称：XX股份有限公司2026年第一次临时股东会"},
    {"type": "h1", "text": "二、会议审议事项"},
    {"type": "h2", "text": "（一）《关于XXXX的议案》"},
    {"type": "para", "text": "**经对比遴选，……。**其余正文……"},
    {"type": "table", "caption": "XX近三年主要财务指标如下：",
     "header": ["指标名称", "2023年", "2024年", "2025年"],
     "rows": [["营业收入（万元）", "6,963.89", "-16,636.00", "8,189.79"]]},
    {"type": "h1", "text": "三、请示事项"},
    {"type": "para", "text": "经核查，……现提请投委会……。"}
  ],
  "attachments": ["XX关于召开2026年第一次临时股东会的通知", "……议案", "……表决票"]
}
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
KAITI_FONT = "楷体_GB2312"
HEITI_FONT = "黑体"
SONGTI_FONT = "宋体"
WESTERN_FONT = "Times New Roman"

TITLE_SIZE = 22  # 二号
BODY_SIZE = 16  # 三号
PAGE_NUMBER_SIZE = 14  # 四号
TABLE_SIZE = 10.5  # 五号
BODY_LINE_PT = 28
TITLE_LINE_PT = 30
TWO_CHAR_INDENT_PT = BODY_SIZE * 2

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*", re.S)


def set_run_font(run, east_asia_font: str, size_pt: float, bold: bool = False) -> None:
    """Western chars -> Times New Roman; CJK chars -> east_asia_font."""
    run.font.name = WESTERN_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), WESTERN_FONT)
    rfonts.set(qn("w:hAnsi"), WESTERN_FONT)
    rfonts.set(qn("w:cs"), WESTERN_FONT)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def add_runs_with_inline_bold(paragraph, text: str, east_asia_font: str, size_pt: float, bold: bool = False) -> None:
    pos = 0
    for match in BOLD_PATTERN.finditer(text):
        if match.start() > pos:
            set_run_font(paragraph.add_run(text[pos:match.start()]), east_asia_font, size_pt, bold)
        set_run_font(paragraph.add_run(match.group(1)), east_asia_font, size_pt, True)
        pos = match.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), east_asia_font, size_pt, bold)


def set_paragraph_format(
    paragraph,
    *,
    line_pt: int = BODY_LINE_PT,
    first_indent: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    left_indent_pt: float = 0,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(TWO_CHAR_INDENT_PT) if first_indent else Pt(0)
    if left_indent_pt:
        fmt.left_indent = Pt(left_indent_pt)
    paragraph.alignment = alignment


def add_paragraph(
    doc: Document,
    text: str,
    *,
    font: str = BODY_FONT,
    size: float = BODY_SIZE,
    bold: bool = False,
    line_pt: int = BODY_LINE_PT,
    first_indent: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    left_indent_pt: float = 0,
):
    p = doc.add_paragraph()
    set_paragraph_format(
        p,
        line_pt=line_pt,
        first_indent=first_indent,
        alignment=alignment,
        left_indent_pt=left_indent_pt,
    )
    add_runs_with_inline_bold(p, text, font, size, bold)
    return p


def add_blank_line(doc: Document, line_pt: int = BODY_LINE_PT) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, line_pt=line_pt, first_indent=False)


def add_page_field(paragraph) -> None:
    set_run_font(paragraph.add_run("-"), SONGTI_FONT, PAGE_NUMBER_SIZE)
    field_run = paragraph.add_run()
    set_run_font(field_run, SONGTI_FONT, PAGE_NUMBER_SIZE)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, result, end):
        field_run._r.append(element)
    set_run_font(paragraph.add_run("-"), SONGTI_FONT, PAGE_NUMBER_SIZE)


def setup_document(doc: Document, page_numbers: bool) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(2.35)

    if page_numbers:
        doc.settings.odd_and_even_pages_header_footer = True
        odd_footer = section.footer.paragraphs[0]
        odd_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_field(odd_footer)
        even_footer = section.even_page_footer.paragraphs[0]
        even_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_page_field(even_footer)


def add_table(doc: Document, block: dict) -> None:
    caption = block.get("caption")
    if caption:
        add_paragraph(doc, caption)

    header = block.get("header") or []
    rows = block.get("rows") or []
    all_rows = ([header] if header else []) + rows
    if not all_rows:
        return
    n_cols = max(len(r) for r in all_rows)

    table = doc.add_table(rows=len(all_rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Autofit to window width so the table never overflows the text area.
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tbl_pr.append(layout)

    for r_idx, row_data in enumerate(all_rows):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = str(row_data[c_idx]) if c_idx < len(row_data) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt = p.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.first_line_indent = Pt(0)
            is_header = bool(header) and r_idx == 0
            add_runs_with_inline_bold(p, text, BODY_FONT, TABLE_SIZE, bold=is_header)


def add_attachments(doc: Document, attachments: list[str]) -> None:
    if not attachments:
        return
    add_blank_line(doc)
    if len(attachments) == 1:
        add_paragraph(doc, f"附件：{attachments[0]}")
        return
    add_paragraph(doc, f"附件：1.{attachments[0]}")
    # "附件：" occupies 3 characters; align following numbers under "1."
    hang_pt = TWO_CHAR_INDENT_PT + BODY_SIZE * 3
    for idx, name in enumerate(attachments[1:], start=2):
        add_paragraph(doc, f"{idx}.{name}", first_indent=False, left_indent_pt=hang_pt)


def add_block(doc: Document, block: dict) -> None:
    kind = block.get("type", "para")
    text = block.get("text", "")
    if kind == "h1":
        add_paragraph(doc, text, font=HEITI_FONT)
    elif kind == "h2":
        add_paragraph(doc, text, font=KAITI_FONT, bold=True)
    elif kind == "h3":
        add_paragraph(doc, text, bold=True)
    elif kind == "h4":
        add_paragraph(doc, text)
    elif kind == "table":
        add_table(doc, block)
    elif kind == "blank":
        add_blank_line(doc)
    else:
        add_paragraph(doc, text)


def build_docx(data: dict, output: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = WESTERN_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    setup_document(doc, page_numbers=bool(data.get("page_numbers", False)))

    for line in str(data.get("title", "审议关于XXXX的议题")).splitlines():
        add_paragraph(
            doc,
            line,
            font=TITLE_FONT,
            size=TITLE_SIZE,
            line_pt=TITLE_LINE_PT,
            first_indent=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    submit_line = data.get("submit_line", "提交部门/子公司：投管公司")
    if submit_line:
        add_paragraph(
            doc,
            submit_line,
            font=KAITI_FONT,
            line_pt=TITLE_LINE_PT,
            first_indent=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    add_blank_line(doc)

    recipient = data.get("recipient", "投委会：")
    if recipient:
        add_paragraph(doc, recipient, first_indent=False)

    intro = data.get("intro")
    if intro:
        add_paragraph(doc, intro)

    for block in data.get("blocks", []):
        add_block(doc, block)

    add_attachments(doc, data.get("attachments", []))

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a 投委会议题 DOCX from a JSON spec.")
    parser.add_argument("input", type=Path, help="JSON spec path")
    parser.add_argument("output", type=Path, help="Output .docx path")
    args = parser.parse_args()
    data = json.loads(args.input.read_text(encoding="utf-8-sig"))
    build_docx(data, args.output)
    print(f"Saved: {args.output}")


if __name__ == "__main__":
    main()
