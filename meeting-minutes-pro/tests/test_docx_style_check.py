#!/usr/bin/env python3
"""Tests for the delivery-time DOCX style readback (fonts/sizes/weights vs the
layout spec). Skipped when python-docx is unavailable."""

from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:  # pragma: no cover
    HAS_DOCX = False


def _load(name: str):
    script = Path(__file__).parents[1] / "scripts" / f"{name}.py"
    spec = importlib.util.spec_from_file_location(name, script)
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


if HAS_DOCX:
    CM = _load("create_minutes_docx")
    DSC = _load("docx_style_check")


def _build(path: Path) -> "Document":
    document = Document()
    CM.add_title(document, "某某公司访谈纪要")
    CM.add_content_paragraph(document, "　　一、完整总结概述")
    CM.add_content_paragraph(document, "　　（一）项目背景")
    CM.add_content_paragraph(document, "　　COO 负责 5G 模组，营收 3.5亿元。")
    document.save(path)
    return document


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class StyleReadbackTests(unittest.TestCase):
    def test_clean_document_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "m.docx"
            _build(path)
            self.assertEqual(DSC.check_docx_style(path), [])

    def test_wrong_font_is_flagged(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "m.docx"
            document = _build(path)
            # Corrupt a body run's font, then re-read.
            body = document.paragraphs[-1]
            body.runs[0].font.name = "SimSun"
            document.save(path)
            problems = DSC.check_docx_style(path)
            self.assertTrue(problems)
            self.assertTrue(any("SimSun" in p for p in problems))

    def test_wrong_bold_is_flagged(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "m.docx"
            document = _build(path)
            # First-level heading must not be bold; make it bold.
            heading = document.paragraphs[1]
            heading.runs[0].bold = True
            document.save(path)
            problems = DSC.check_docx_style(path)
            self.assertTrue(any("加粗" in p for p in problems))


if __name__ == "__main__":
    unittest.main()
