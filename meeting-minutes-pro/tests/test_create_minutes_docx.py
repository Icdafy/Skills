#!/usr/bin/env python3
"""Rendering-layer tests: role→font mapping, western-run splitting, page
geometry and the page-number field of the generated DOCX. Skipped when
python-docx is unavailable (install the skill runtime to run them)."""

from __future__ import annotations

import importlib.util
import sys
import unittest
from pathlib import Path

try:
    import docx  # noqa: F401
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    HAS_DOCX = True
except ImportError:  # pragma: no cover - runtime without python-docx
    HAS_DOCX = False

if HAS_DOCX:
    SCRIPT = Path(__file__).parents[1] / "scripts" / "create_minutes_docx.py"
    SPEC = importlib.util.spec_from_file_location("create_minutes_docx", SCRIPT)
    CM = importlib.util.module_from_spec(SPEC)
    assert SPEC.loader is not None
    sys.modules[SPEC.name] = CM
    SPEC.loader.exec_module(CM)


def _east_asia(run):
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn("w:eastAsia"))


def _runs(paragraph):
    return [(r.text, r.font.name, _east_asia(r), r.bold)
            for r in paragraph.runs if r.text]


def _footer_pattern(paragraph):
    parts = []
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            parts.append("".join(node.text or "" for node in child.iter(qn("w:t"))))
        elif child.tag == qn("w:fldSimple"):
            parts.append("{PAGE}" if child.get(qn("w:instr")) == "PAGE" else "{FIELD}")
    return "".join(parts)


def _part_text(part):
    return "".join(node.text or "" for node in part._element.iter(qn("w:t")))


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class RenderRoleTests(unittest.TestCase):
    def setUp(self) -> None:
        self.doc = Document()

    def last(self):
        return self.doc.paragraphs[-1]

    def test_title_uses_xiaobiaosong(self) -> None:
        CM.add_title(self.doc, "某某公司访谈纪要")
        self.assertEqual(_runs(self.last())[0][1], "方正小标宋简体")

    def test_first_level_hei_not_bold(self) -> None:
        CM.add_content_paragraph(self.doc, "　　一、完整总结概述")
        text, name, ea, bold = _runs(self.last())[0]
        self.assertEqual((name, ea), ("SimHei", "SimHei"))
        self.assertNotEqual(bold, True)

    def test_second_level_kai_bold(self) -> None:
        CM.add_content_paragraph(self.doc, "　　（一）项目背景")
        _, name, _, bold = _runs(self.last())[0]
        self.assertEqual(name, "楷体_GB2312")
        self.assertTrue(bold)

    def test_third_level_fangsong_bold_with_times_numeral(self) -> None:
        CM.add_content_paragraph(self.doc, "　　1.核心技术指标")
        runs = _runs(self.last())
        self.assertEqual(runs[0][:2], ("1", "Times New Roman"))
        self.assertTrue(runs[0][3])
        self.assertEqual(runs[1][1], "仿宋_GB2312")
        self.assertTrue(self.last().paragraph_format.keep_together)

    def test_fourth_level_fangsong_not_bold(self) -> None:
        CM.add_content_paragraph(self.doc, "　　（1）测试环境")
        names = {name for _, name, _, _ in _runs(self.last())}
        self.assertIn("仿宋_GB2312", names)
        self.assertNotEqual(_runs(self.last())[-1][3], True)


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class WesternRunTests(unittest.TestCase):
    def setUp(self) -> None:
        self.doc = Document()

    def test_letters_and_alnum_go_to_times(self) -> None:
        CM.add_content_paragraph(self.doc, "　　COO 负责 5G 模组与 A4 检测。")
        runs = _runs(self.doc.paragraphs[-1])
        times = {t for t, name, _, _ in runs if name == "Times New Roman"}
        fangsong = {t for t, name, _, _ in runs if name == "仿宋_GB2312"}
        self.assertEqual(times, {"COO", "5G", "A4"})
        self.assertTrue(all("模组" in t or "负责" in t or "检测" in t or t.strip() == ""
                            for t in fangsong))

    def test_decimal_led_body_is_not_a_heading(self) -> None:
        # #3 end-to-end: a body line opening with a decimal renders as body
        # (FangSong, not kept-together), and the decimal itself is Times.
        CM.add_content_paragraph(self.doc, "　　3.5亿元用于研发投入。")
        para = self.doc.paragraphs[-1]
        self.assertNotEqual(para.paragraph_format.keep_together, True)
        runs = _runs(para)
        self.assertEqual(runs[0][:2], ("3.5", "Times New Roman"))
        self.assertEqual(runs[1][1], "仿宋_GB2312")
        self.assertNotEqual(runs[1][3], True)


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class PageLayoutTests(unittest.TestCase):
    def test_geometry_and_page_number_field(self) -> None:
        doc = Document()
        CM.configure_document(doc)
        section = doc.sections[0]
        # Compare in centimetres: page size is stored as integer twips, so raw
        # EMU round-trips are off by a few hundred EMU.
        self.assertEqual(round(section.page_width.cm, 1), 21.0)
        self.assertEqual(round(section.page_height.cm, 1), 29.7)
        self.assertEqual(round(section.top_margin.cm, 1), 3.7)
        self.assertEqual(round(section.bottom_margin.cm, 1), 3.5)
        self.assertEqual(round(section.left_margin.cm, 1), 2.8)
        self.assertEqual(round(section.right_margin.cm, 1), 2.6)
        self.assertTrue(doc.settings.odd_and_even_pages_header_footer)
        self.assertFalse(section.different_first_page_header_footer)
        self.assertEqual(section.footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(
            section.even_page_footer.paragraphs[0].alignment,
            WD_ALIGN_PARAGRAPH.LEFT,
        )
        for footer in (section.footer, section.even_page_footer):
            paragraph = footer.paragraphs[0]
            self.assertEqual(_footer_pattern(paragraph), "- {PAGE} -")
            for run in paragraph._p.iter(qn("w:r")):
                properties = run.find(qn("w:rPr"))
                fonts = properties.find(qn("w:rFonts"))
                self.assertEqual(fonts.get(qn("w:ascii")), "宋体")
                self.assertEqual(fonts.get(qn("w:hAnsi")), "宋体")
                self.assertEqual(fonts.get(qn("w:eastAsia")), "宋体")
                self.assertEqual(properties.find(qn("w:sz")).get(qn("w:val")), "28")
        for header in (section.header, section.even_page_header, section.first_page_header):
            self.assertEqual(_part_text(header).strip(), "")

    def test_qa_separator_is_blank_exact_spacing(self) -> None:
        doc = Document()
        CM.add_qa_separator(doc)
        para = doc.paragraphs[-1]
        self.assertEqual(para.text, "")
        self.assertEqual(para.paragraph_format.line_spacing_rule, WD_LINE_SPACING.EXACTLY)
        self.assertAlmostEqual(para.paragraph_format.line_spacing, Pt(28))

    def test_first_line_indent_is_character_based(self) -> None:
        # #5: two-character indent via w:firstLineChars (scales with font size),
        # with the absolute w:firstLine kept as a fallback.
        doc = Document()
        CM.add_content_paragraph(doc, "　　正文一段。")
        ind = doc.paragraphs[-1]._p.pPr.find(qn("w:ind"))
        self.assertEqual(ind.get(qn("w:firstLineChars")), "200")
        self.assertIsNotNone(ind.get(qn("w:firstLine")))


if __name__ == "__main__":
    unittest.main()
