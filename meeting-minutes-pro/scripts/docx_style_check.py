#!/usr/bin/env python3
"""Read a generated minutes DOCX back and confirm every run carries the font,
size and weight the layout spec prescribes — without rendering to PDF.

The content-fidelity check in ``check_all.py`` proves the DOCX carries the right
*text*; this proves it carries the right *formatting*. Each body paragraph's
expected face is re-derived from ``format_spec.paragraph_role`` (the same
function the renderer uses), western runs must be Times New Roman, and the
centred title/subtitle are checked against their fixed faces. Headers must be
empty, while odd/even footers must carry the exact page-number field, font,
size and alignment from ``format_spec``. Any drift is therefore caught here
instead of only by a human eyeballing the PDF.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from format_spec import (  # noqa: E402  (needs the path shim above)
    BODY_SIZE,
    KAI_FONT,
    NUMBER_FONT,
    PAGE_NUMBER_FONT,
    PAGE_NUMBER_PREFIX,
    PAGE_NUMBER_SIZE,
    PAGE_NUMBER_SUFFIX,
    SUBTITLE_SIZE,
    TITLE_FONT,
    TITLE_SIZE,
    WESTERN_SEGMENT,
    paragraph_role,
)


def _east_asia(run) -> str | None:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn("w:eastAsia"))


def _is_western(text: str) -> bool:
    return WESTERN_SEGMENT.fullmatch(text) is not None


def _expected(paragraph, center_seen: int) -> tuple[str, int, bool] | None:
    """Return (east-asia font, size pt, bold) expected for this paragraph, or
    None to skip it (blank line / spacer)."""
    text = paragraph.text.strip()
    if not text:
        return None
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        # First centred line is the title; any later centred line is the subtitle.
        if center_seen == 0:
            return TITLE_FONT, TITLE_SIZE, False
        return KAI_FONT, SUBTITLE_SIZE, False
    _role, font, bold = paragraph_role(text)
    return font, BODY_SIZE, bold


def _part_text(part) -> str:
    """Visible text in a header/footer part, including field result text."""
    return "".join(node.text or "" for node in part._element.iter(qn("w:t")))


def _footer_pattern(paragraph) -> str:
    """Represent the generated simple PAGE field as ``{PAGE}``."""
    parts: list[str] = []
    for child in paragraph._p:
        if child.tag == qn("w:r"):
            parts.append("".join(node.text or "" for node in child.iter(qn("w:t"))))
        elif child.tag == qn("w:fldSimple"):
            instruction = (child.get(qn("w:instr")) or "").strip()
            parts.append("{PAGE}" if instruction == "PAGE" else f"{{{instruction}}}")
    return "".join(parts)


def _footer_run_problems(paragraph, label: str) -> list[str]:
    problems: list[str] = []
    expected_size = str(PAGE_NUMBER_SIZE * 2)
    for run in paragraph._p.iter(qn("w:r")):
        text = "".join(node.text or "" for node in run.iter(qn("w:t")))
        if not text:
            continue
        properties = run.find(qn("w:rPr"))
        fonts = None if properties is None else properties.find(qn("w:rFonts"))
        for attribute in ("ascii", "hAnsi", "eastAsia"):
            actual = None if fonts is None else fonts.get(qn(f"w:{attribute}"))
            if actual != PAGE_NUMBER_FONT:
                problems.append(
                    f"{label}中「{text}」{attribute} 字体为 {actual}，应为 {PAGE_NUMBER_FONT}"
                )
        size = None if properties is None else properties.find(qn("w:sz"))
        actual_size = None if size is None else size.get(qn("w:val"))
        if actual_size != expected_size:
            try:
                got = None if actual_size is None else int(actual_size) / 2
            except ValueError:
                got = actual_size
            problems.append(
                f"{label}中「{text}」字号为 {got} 磅，应为 {PAGE_NUMBER_SIZE} 磅"
            )
    return problems


def _page_frame_problems(document) -> list[str]:
    problems: list[str] = []
    if not document.settings.odd_and_even_pages_header_footer:
        problems.append("未启用奇偶页不同的页眉页脚")

    expected_pattern = f"{PAGE_NUMBER_PREFIX}{{PAGE}}{PAGE_NUMBER_SUFFIX}"
    for index, section in enumerate(document.sections, start=1):
        prefix = f"第 {index} 节"
        if section.different_first_page_header_footer:
            problems.append(f"{prefix}启用了首页不同页眉页脚，应关闭")
        for label, header in (
            ("奇数页页眉", section.header),
            ("偶数页页眉", section.even_page_header),
            ("首页页眉", section.first_page_header),
        ):
            if _part_text(header).strip():
                problems.append(f"{prefix}{label}应为空，实际为「{_part_text(header).strip()}」")

        for label, footer, alignment in (
            ("奇数页页脚", section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
            ("偶数页页脚", section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
        ):
            paragraphs = footer.paragraphs
            numbered = [p for p in paragraphs if "{PAGE}" in _footer_pattern(p)]
            if len(numbered) != 1:
                problems.append(f"{prefix}{label}应包含且仅包含一个 PAGE 页码域")
                continue
            paragraph = numbered[0]
            extras = [
                p for p in paragraphs
                if p._p is not paragraph._p and _footer_pattern(p).strip()
            ]
            if extras:
                problems.append(f"{prefix}{label}含页码以外的内容")
            actual_pattern = _footer_pattern(paragraph)
            if actual_pattern != expected_pattern:
                problems.append(
                    f"{prefix}{label}页码格式为「{actual_pattern}」，应为「{expected_pattern}」"
                )
            if paragraph.alignment != alignment:
                want = "右对齐" if alignment == WD_ALIGN_PARAGRAPH.RIGHT else "左对齐"
                problems.append(f"{prefix}{label}未{want}")
            problems.extend(_footer_run_problems(paragraph, f"{prefix}{label}"))
    return problems


def check_docx_style(path: Path) -> list[str]:
    document = Document(str(path))
    problems: list[str] = []
    center_seen = 0
    for paragraph in document.paragraphs:
        expected = _expected(paragraph, center_seen)
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and paragraph.text.strip():
            center_seen += 1
        if expected is None:
            continue
        ea_font, size, bold = expected
        preview = paragraph.text.strip()[:20]
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            want_font = NUMBER_FONT if _is_western(run.text) else ea_font
            if run.font.name != want_font:
                problems.append(
                    f"「{preview}」中「{run.text[:12]}」字体为 {run.font.name}，应为 {want_font}"
                )
            if not _is_western(run.text) and _east_asia(run) != ea_font:
                problems.append(
                    f"「{preview}」中「{run.text[:12]}」东亚字体为 {_east_asia(run)}，应为 {ea_font}"
                )
            if run.font.size != Pt(size):
                got = None if run.font.size is None else round(run.font.size.pt, 1)
                problems.append(
                    f"「{preview}」中「{run.text[:12]}」字号为 {got} 磅，应为 {size} 磅"
                )
            if bool(run.bold) != bold:
                problems.append(
                    f"「{preview}」中「{run.text[:12]}」加粗为 {bool(run.bold)}，应为 {bold}"
                )
    problems.extend(_page_frame_problems(document))
    return problems


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path, help="generated minutes DOCX")
    args = parser.parse_args()
    if not args.docx.is_file():
        print(f"找不到 DOCX：{args.docx}")
        return 2
    problems = check_docx_style(args.docx)
    if problems:
        print("DOCX 样式回读未通过：")
        for problem in problems:
            print(f"- {problem}")
        return 1
    print("DOCX 样式回读通过：正文及页眉页脚的字体、字号、加粗与版式均符合规约。")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
