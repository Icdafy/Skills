#!/usr/bin/env python3
"""Create a fixed-format Chinese meeting-minutes DOCX from validated plain text.

Layout knowledge (heading grammar, role→font mapping, western-run rule, sizes,
geometry, fonts) is imported from ``format_spec`` — the same module the text
validator uses — so the rendered DOCX can never disagree with what
``quality_check`` accepted. After saving, the bundled GB2312 faces are embedded
into the package (see ``embed_fonts``) so the file renders faithfully on
machines that lack 仿宋_GB2312 / 楷体_GB2312; pass ``--no-embed`` to skip.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from format_spec import (  # noqa: E402  (needs the path shim above)
    BODY_LINE_SPACING,
    BODY_SIZE,
    BOTTOM_MARGIN_CM,
    FANGSONG_FONT,
    FIRST_LINE_INDENT_CHARS,
    FIRST_LINE_INDENT_PT,
    INDENT,
    KAI_FONT,
    LEFT_MARGIN_CM,
    NUMBER_FONT,
    PAGE_HEIGHT_CM,
    PAGE_NUMBER_SIZE,
    PAGE_WIDTH_CM,
    RIGHT_MARGIN_CM,
    SUBTITLE_LINE_SPACING,
    SUBTITLE_SIZE,
    TITLE_FONT,
    TITLE_LINE_SPACING,
    TITLE_SIZE,
    TOP_MARGIN_CM,
    WESTERN_SEGMENT,
    paragraph_role,
)
from embed_fonts import (  # noqa: E402
    default_font_paths,
    embed_fonts_into_docx,
    verify_embedded_fonts,
)
from font_preflight import required_font_status  # noqa: E402
from quality_check import validate  # noqa: E402

# Page numbers follow the explicit footer rule: `-1-`, 三号 (16 pt),
# 仿宋_GB2312 throughout, odd pages right / even pages left.
PAGE_NUMBER_FONT = FANGSONG_FONT


def set_east_asia_font(run, font_name: str, size: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_western_font(run, east_asia_font: str, size: float, bold: bool = False) -> None:
    """Latin letters and Arabic numerals are set in Times New Roman; the
    surrounding East Asian face is kept as the CJK fallback."""
    run.font.name = NUMBER_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def add_text_runs(paragraph, text: str, font_name: str, size: float, bold: bool = False) -> None:
    """Add runs for ``text``, switching western (ASCII) runs to Times New Roman."""
    cursor = 0
    for match in WESTERN_SEGMENT.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_east_asia_font(run, font_name, size, bold)
        run = paragraph.add_run(match.group())
        set_western_font(run, font_name, size, bold)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_east_asia_font(run, font_name, size, bold)


def set_exact_line_spacing(paragraph, points: float) -> None:
    paragraph.paragraph_format.line_spacing = Pt(points)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def set_first_line_indent(paragraph) -> None:
    """Indent the first line by two characters. w:firstLineChars scales with the
    font size (the public-document convention); the absolute w:firstLine is kept
    as a fallback for renderers that ignore the character measure."""
    paragraph.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    indent = paragraph._p.get_or_add_pPr().get_or_add_ind()
    indent.set(qn("w:firstLineChars"), str(FIRST_LINE_INDENT_CHARS * 100))


def set_body_layout(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_first_line_indent(paragraph)
    paragraph.paragraph_format.widow_control = True
    set_exact_line_spacing(paragraph, BODY_LINE_SPACING)


def append_page_field(paragraph) -> None:
    leading = paragraph.add_run("-")
    set_east_asia_font(leading, PAGE_NUMBER_FONT, PAGE_NUMBER_SIZE)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), PAGE_NUMBER_FONT)
    fonts.set(qn("w:hAnsi"), PAGE_NUMBER_FONT)
    fonts.set(qn("w:eastAsia"), PAGE_NUMBER_FONT)
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(PAGE_NUMBER_SIZE * 2))
    properties.append(size)
    field_run.append(properties)
    text = OxmlElement("w:t")
    text.text = "1"
    field_run.append(text)
    field.append(field_run)
    paragraph._p.append(field)
    trailing = paragraph.add_run("-")
    set_east_asia_font(trailing, PAGE_NUMBER_FONT, PAGE_NUMBER_SIZE)


def format_footer(footer, alignment: WD_ALIGN_PARAGRAPH) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    append_page_field(paragraph)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(TOP_MARGIN_CM)
    section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
    section.left_margin = Cm(LEFT_MARGIN_CM)
    section.right_margin = Cm(RIGHT_MARGIN_CM)
    section.start_type = WD_SECTION_START.NEW_PAGE
    document.settings.odd_and_even_pages_header_footer = True
    format_footer(section.footer, WD_ALIGN_PARAGRAPH.RIGHT)
    format_footer(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)


def add_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    set_exact_line_spacing(paragraph, TITLE_LINE_SPACING)
    add_text_runs(paragraph, title, TITLE_FONT, TITLE_SIZE)


def add_subtitle(document: Document, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    set_exact_line_spacing(paragraph, SUBTITLE_LINE_SPACING)
    add_text_runs(paragraph, subtitle, KAI_FONT, SUBTITLE_SIZE)
    blank = document.add_paragraph()
    blank.paragraph_format.keep_with_next = True
    set_exact_line_spacing(blank, SUBTITLE_LINE_SPACING)


def add_content_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_body_layout(paragraph)
    content = text.removeprefix(INDENT).strip()
    role, font_name, bold = paragraph_role(content)
    if role != "body":
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = True
    add_text_runs(paragraph, content, font_name, BODY_SIZE, bold)


def add_qa_separator(document: Document) -> None:
    """Blank spacer line rendered between consecutive Q/A groups."""
    paragraph = document.add_paragraph()
    set_exact_line_spacing(paragraph, BODY_LINE_SPACING)


def missing_font_families() -> list[str]:
    return [item["family"] for item in required_font_status() if not item["installed"]]


def read_lines(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8-sig").splitlines()


def embed_bundled_fonts(output: Path) -> dict:
    """Embed the bundled GB2312 faces into ``output`` and verify the result."""
    report = embed_fonts_into_docx(output, default_font_paths())
    report["verify"] = verify_embedded_fonts(output)
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Create fixed-format meeting minutes DOCX.")
    parser.add_argument("--input", type=Path, required=True, help="validated UTF-8 plain-text minutes")
    parser.add_argument("--output", type=Path, required=True, help="destination .docx path")
    parser.add_argument("--subtitle", default=None, help="optional centered department or subtitle line")
    parser.add_argument(
        "--mode", choices=("auto", "minutes", "qa", "qa-summary"), default="auto"
    )
    parser.add_argument(
        "--no-embed",
        action="store_true",
        help="不把随附字体嵌入 DOCX（默认嵌入，使文件在未装字体的机器上仍忠实呈现）",
    )
    parser.add_argument(
        "--allow-line",
        type=int,
        action="append",
        default=[],
        metavar="行号",
        help="与 quality_check.py 一致：放行该行经确认属转录稿真实内容的表述",
    )
    args = parser.parse_args()

    if not args.input.is_file():
        parser.error(f"找不到输入文件：{args.input}")
    validation_errors = validate(args.input, args.mode, frozenset(args.allow_line))
    if validation_errors:
        parser.error("输入文本未通过校验：\n- " + "\n- ".join(validation_errors))
    missing_fonts = missing_font_families()
    if missing_fonts:
        parser.error(
            "缺少固定版式所需字体："
            + "、".join(missing_fonts)
            + "。先运行 font_preflight.py --check；取得用户许可后可运行 "
            "font_preflight.py --install-user 安装随技能提供的字体。"
        )
    lines = read_lines(args.input)
    title = next((line.strip() for line in lines if line.strip()), None)
    if not title:
        parser.error("输入文件不包含标题。")

    document = Document()
    configure_document(document)
    add_title(document, title)
    if args.subtitle:
        add_subtitle(document, args.subtitle)

    title_consumed = False
    pending_blank = False
    for line in lines:
        if not line.strip():
            pending_blank = title_consumed
            continue
        if not title_consumed and line.strip() == title:
            title_consumed = True
            continue
        # Source blank lines before a new Q/A group become a spacer paragraph
        # so consecutive Q/A groups stay visually separated in the DOCX.
        if pending_blank and line.removeprefix(INDENT).strip().startswith("问："):
            add_qa_separator(document)
        pending_blank = False
        add_content_paragraph(document, line)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(f"已生成：{args.output}")

    if not args.no_embed:
        report = embed_bundled_fonts(args.output)
        embedded = "、".join(item["font"] for item in report["embedded"]) or "无"
        print(f"已嵌入字体：{embedded}")
        for item in report["skipped"]:
            print(f"未嵌入 {item['font']}：{item['reason']}")
        if not report["verify"]["ok"]:
            parser.error(
                "字体嵌入校验未通过，请勿交付该 DOCX：\n"
                + "\n".join(
                    f"- {font['font']}：{font.get('error', '未知')}"
                    for font in report["verify"]["fonts"]
                    if not font.get("valid")
                )
                or "- 未设置 w:embedTrueTypeFonts"
            )


if __name__ == "__main__":
    main()
