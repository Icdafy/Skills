#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py —— 立项报告 Word 生成器

把一份结构化的报告内容（JSON 或 python dict）渲染成与样本版式一致的 .docx：
- A4 页面（上 3.7cm、下 3.5cm、左 2.8cm、右 2.6cm）
- 标题（方正小标宋简体二号 22pt、居中、固定行距 30 磅）
- 正文（仿宋_GB2312 三号 16pt、两端对齐、首行缩进 2 字符、固定行距 28 磅）
- 四级编号标题（一、/（一）/ 1. /（1））：
  一级黑体三号不加粗；二级楷体_GB2312 三号加粗；三级仿宋_GB2312 三号加粗；四级仿宋_GB2312 三号不加粗
- 目录（Word 原生 TOC 域，打开后可右键更新页码）
- 表格（统一细边框 #BFBFBF、表头浅蓝底 #D9E2F3、单元格内边距、DXA 宽度、
  所有单元格内容水平+垂直居中）
- 页脚页码（奇偶页外侧，四号宋体，格式 -1-）

设计原则：脚本只管"排版"，不管"写作"。你（调用方）负责把写好的各章内容
按下面的 BLOCK 结构组织好传进来。脚本对内容零编造。

依赖：python-docx  →  pip install python-docx --break-system-packages

用法：
    python build_docx.py content.json out.docx
或在 python 中：
    from build_docx import build
    build(content_dict, "out.docx")

content 结构（dict）：
{
  "cover": {
     "title_lines": ["关于投资XX公司的", "立项报告"],   # 居中大标题，1~2行
     "org": "西安未央城建产投投资管理有限公司",          # 平台名
     "date": "2026年5月"                                  # 年月
  },
  "toc": true,                          # 是否插入目录域（默认 true）
  "summary": "要点概述正文（可选，多段用 \\n 分隔）",   # 可选，放目录后
  "summary_title": "要点概述",          # 可选，默认"要点概述"
  "blocks": [ ...见下... ]              # 正文，按顺序渲染
}

blocks 里每个元素是一个 dict，type 决定渲染方式：
  {"type":"h1","text":"一、项目概况"}            # 一级标题（你自己带"一、"前缀）
  {"type":"h2","text":"（一）核心人员"}           # 二级
  {"type":"h3","text":"1.首席科学家XXX"}          # 三级
  {"type":"h4","text":"（1）XXX"}                 # 四级
  {"type":"p","text":"正文段落……"}               # 普通段落
  {"type":"p","text":"……","bold":true}          # 加粗段落（样本里常见整段加粗的"金句"）
  {"type":"bullet","items":["要点1","要点2"]}     # 项目符号列表
  {"type":"table","header":["列1","列2"],          # 表格；header 可省略（无表头）
      "rows":[["a","b"],["c","d"]],
      "widths":[3000,6360]}                       # widths 可选，DXA，合计≈9360
  {"type":"note","text":"【待补充：尽调阶段确认】"} # 灰色提示，用于缺口标注
  {"type":"pagebreak"}                            # 分页

标题层级与编号：脚本不自动编号，"一、""（一）"等前缀由你写在 text 里——
这与样本完全一致（样本就是手工编号），好处是你能灵活处理"要点概述""附件"等无编号块。
"""

import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 版式常量（对齐公文格式样本） ----------
FONT_TITLE = "方正小标宋简体"
FONT_BODY = "仿宋_GB2312"
FONT_H1 = "黑体"            # 使用系统原本黑体，不随技能打包替换
FONT_H2 = "楷体_GB2312"
FONT_FOOTER = "宋体"
FONT_EN = "Times New Roman"
BODY_SZ = 16               # 三号
H1_SZ = 16                 # 一级标题黑体三号，不加粗
H2_SZ = 16                 # 二级标题楷体_GB2312 三号，加粗
H3_SZ = 16                 # 三级标题仿宋_GB2312 三号，加粗
H4_SZ = 16                 # 四级标题仿宋_GB2312 三号，不加粗
COVER_TITLE_SZ = 22        # 二号方正小标宋简体
COVER_ORG_SZ = 16          # 落款/机构三号仿宋_GB2312
FOOTER_SZ = 14             # 四号宋体
BODY_LINE_PT = 28
TITLE_LINE_PT = 30
TABLE_BORDER = "BFBFBF"    # 表格边框灰
HEADER_FILL = "D9E2F3"     # 表头浅蓝底
CONTENT_WIDTH = 8844       # A4：21cm - 2.8cm - 2.6cm ≈ 15.6cm
TABLE_SZ = 10.5            # 数据表默认小五，避免宽表溢出；必要时调用方可拆表


def _set_run_font(run, size=BODY_SZ, bold=False, color=None, font_name=FONT_BODY):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_EN
    # 中文字体需要单独设 eastAsia
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), FONT_EN)
    rfonts.set(qn('w:hAnsi'), FONT_EN)
    rfonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_para(doc, text, size=BODY_SZ, bold=False, align=None, color=None,
              space_before=0, space_after=6, outline=None, line=None,
              font_name=FONT_BODY):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line:
        if isinstance(line, (int, float)) and line > 5:
            pf.line_spacing = Pt(line)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        else:
            pf.line_spacing = line
    # 中文首行不缩进标题；正文段落首行缩进 2 字符
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color, font_name=font_name)
    if outline is not None:
        _set_outline_level(p, outline)
    return p


def _set_outline_level(paragraph, level):
    """给段落打 outlineLevel，使其能进 TOC（0=H1,1=H2,...）。"""
    pPr = paragraph._p.get_or_add_pPr()
    ol = pPr.find(qn('w:outlineLvl'))
    if ol is None:
        ol = OxmlElement('w:outlineLvl')
        pPr.append(ol)
    ol.set(qn('w:val'), str(level))


def _indent_first_line(paragraph, chars=2):
    """正文段落首行缩进 2 字符。"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))


def _style_cell(cell, fill=None, color=TABLE_BORDER, sz=4,
                margins=(60, 60, 100, 100), valign='center'):
    """一次性构建 tcPr 的子元素，严格按 OOXML schema 顺序：
       tcBorders → shd → tcMar → vAlign 。顺序错了 Word 校验会报错。
       valign 默认 center：单元格内容垂直居中（配合段落水平居中，实现表格内容整体居中）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    # 清掉可能已存在的同名元素，避免重复
    for tag in ('w:tcBorders', 'w:shd', 'w:tcMar', 'w:vAlign'):
        for e in tcPr.findall(qn(tag)):
            tcPr.remove(e)
    # 1) 边框
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)
    # 2) 底色
    if fill:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    # 3) 内边距
    top, bottom, left, right = margins
    m = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)
    # 4) 垂直居中（水平居中在 _fill_cell 的段落上设置）
    if valign:
        v = OxmlElement('w:vAlign')
        v.set(qn('w:val'), valign)
        tcPr.append(v)


def _fill_cell(cell, text, bold=False, header=False, size=BODY_SZ):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中；垂直居中见 _style_cell
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run(str(text) if text is not None else "")
    _set_run_font(run, size=size, bold=bold or header, font_name=FONT_BODY)
    _style_cell(cell, fill=HEADER_FILL if header else None)


def _add_table(doc, header, rows, widths=None):
    ncols = len(header) if header else (len(rows[0]) if rows else 1)
    if not widths or len(widths) != ncols:
        widths = [CONTENT_WIDTH // ncols] * ncols
    nrows = (1 if header else 0) + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 设置表格总宽（用已存在的 tblW，避免重复/顺序错误）
    _set_tbl_width(table, sum(widths))
    r = 0
    if header:
        for j, h in enumerate(header):
            _fill_cell(table.rows[0].cells[j], h, header=True, size=TABLE_SZ)
        r = 1
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            _fill_cell(table.rows[r + i].cells[j], val, size=TABLE_SZ)
    # 列宽
    for row in table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = Twips(w)
    return table


def _set_tbl_width(table, total):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total))
    tblW.set(qn('w:type'), 'dxa')


def _add_toc(doc):
    """插入 Word 原生 TOC 域；打开文档后需手动/自动更新（右键→更新域）。"""
    _add_para(doc, "目录", size=H1_SZ, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=8, line=BODY_LINE_PT, font_name=FONT_H1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（打开文档后右键此处『更新域』生成目录与页码）"
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    for el in (fldChar1, instr, fldChar2, t, fldChar3):
        run._r.append(el)


def _add_page_number_footer(section):
    _add_footer_page_field(section.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    try:
        _add_footer_page_field(section.even_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)
    except Exception:
        pass


def _add_footer_page_field(p, align):
    p.alignment = align
    run = p.add_run()
    run.add_text("-")
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    for el in (f1, instr, f2):
        run._r.append(el)
    run.add_text("-")
    _set_run_font(run, size=FOOTER_SZ, font_name=FONT_FOOTER)


def _enable_odd_even_footers(doc):
    settings = doc.settings.element
    even = settings.find(qn('w:evenAndOddHeaders'))
    if even is None:
        settings.append(OxmlElement('w:evenAndOddHeaders'))


def _set_a4_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def _fix_zoom(doc):
    """python-docx 默认 settings.xml 里的 w:zoom 可能缺 percent 属性，补上。"""
    try:
        settings = doc.settings.element
        zoom = settings.find(qn('w:zoom'))
        if zoom is not None and zoom.get(qn('w:percent')) is None:
            zoom.set(qn('w:percent'), '100')
    except Exception:
        pass


def build(content, out_path):
    doc = Document()
    _fix_zoom(doc)
    _enable_odd_even_footers(doc)
    sec = doc.sections[0]
    _set_a4_page(sec)
    _add_page_number_footer(sec)

    # 设置 Normal 默认字体
    normal = doc.styles['Normal']
    normal.font.name = FONT_EN
    normal.font.size = Pt(BODY_SZ)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # ---------- 封面 ----------
    cover = content.get("cover", {})
    for _ in range(6):
        doc.add_paragraph()
    for line in cover.get("title_lines", ["立项报告"]):
        _add_para(doc, line, size=COVER_TITLE_SZ, bold=False,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
                  line=TITLE_LINE_PT, font_name=FONT_TITLE)
    for _ in range(6):
        doc.add_paragraph()
    if cover.get("org"):
        _add_para(doc, cover["org"], size=COVER_ORG_SZ, bold=False,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
                  line=BODY_LINE_PT, font_name=FONT_BODY)
    if cover.get("date"):
        _add_para(doc, cover["date"], size=COVER_ORG_SZ, bold=False,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line=BODY_LINE_PT,
                  font_name=FONT_BODY)
    doc.add_page_break()

    # ---------- 目录 ----------
    if content.get("toc", True):
        _add_toc(doc)
        doc.add_page_break()

    # ---------- 要点概述（可选） ----------
    if content.get("summary"):
        _add_para(doc, content.get("summary_title", "要点概述"),
                  size=H1_SZ, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=8, outline=0, line=BODY_LINE_PT,
                  font_name=FONT_H1)
        for para in str(content["summary"]).split("\n"):
            if para.strip():
                p = _add_para(doc, para.strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              line=BODY_LINE_PT, font_name=FONT_BODY)
                _indent_first_line(p)
        doc.add_page_break()

    # ---------- 正文 ----------
    for blk in content.get("blocks", []):
        t = blk.get("type")
        if t == "h1":
            p = _add_para(doc, blk["text"], size=H1_SZ, bold=False,
                          space_before=8, space_after=6, outline=0,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_H1)
            _indent_first_line(p)
        elif t == "h2":
            p = _add_para(doc, blk["text"], size=H2_SZ, bold=True,
                          space_before=6, space_after=4, outline=1,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_H2)
            _indent_first_line(p)
        elif t == "h3":
            p = _add_para(doc, blk["text"], size=H3_SZ, bold=True,
                          space_before=4, space_after=4, outline=2,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_BODY)
            _indent_first_line(p)
        elif t == "h4":
            p = _add_para(doc, blk["text"], size=H4_SZ, bold=False,
                          space_before=2, space_after=2, outline=3,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_BODY)
            _indent_first_line(p)
        elif t == "p":
            p = _add_para(doc, blk["text"], bold=blk.get("bold", False),
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_BODY)
            _indent_first_line(p)
        elif t == "bullet":
            for item in blk.get("items", []):
                p = doc.add_paragraph(style=None)
                p.paragraph_format.left_indent = Twips(480)
                p.paragraph_format.line_spacing = Pt(BODY_LINE_PT)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                run = p.add_run("• " + str(item))
                _set_run_font(run, size=BODY_SZ, font_name=FONT_BODY)
        elif t == "table":
            _add_table(doc, blk.get("header"), blk.get("rows", []), blk.get("widths"))
            doc.add_paragraph()  # 表后空行
        elif t == "note":
            _add_para(doc, blk["text"], color="808080", line=BODY_LINE_PT,
                      font_name=FONT_BODY)
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            # 未知类型，按普通段落处理，避免内容丢失
            if blk.get("text"):
                _add_para(doc, blk["text"], align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line=BODY_LINE_PT, font_name=FONT_BODY)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python build_docx.py content.json out.docx")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    path = build(data, sys.argv[2])
    print("已生成:", path)
