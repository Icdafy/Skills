#!/usr/bin/env python3
"""Build and audit WeiYang-style meeting-minute DOCX files."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


TITLE_FONT = "方正小标宋简体"
H1_FONT = "黑体"
H2_FONT = "楷体_GB2312"
BODY_FONT = "仿宋_GB2312"
DIGIT_FONT = "Times New Roman"
ALLOWED_CJK_FONTS = {TITLE_FONT, H1_FONT, H2_FONT, BODY_FONT, "楷体GB2312", "仿宋GB2312"}
ALLOWED_ALL_FONTS = ALLOWED_CJK_FONTS | {DIGIT_FONT}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

CN_NUM = "一二三四五六七八九十"
H1_RE = re.compile(rf"^[{CN_NUM}]+、")
H2_RE = re.compile(rf"^（[{CN_NUM}]+）")
H3_RE = re.compile(r"^\d+[.．]")
H4_RE = re.compile(r"^（\d+）")
META_RE = re.compile(r"^(访谈时间|访谈地点|访谈对象|访谈人员|会议时间|会议地点|参会人员|会议主题)[:：]")
DIGIT_RE = re.compile(r"(\d+)")


def set_run_font(run, font_name: str, size_pt: float = 16, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def add_text_runs(paragraph, text: str, font_name: str, size_pt: float, bold: bool = False) -> None:
    for part in DIGIT_RE.split(text):
        if not part:
            continue
        if part.isdigit():
            run = paragraph.add_run(part)
            set_run_font(run, DIGIT_FONT, size_pt, bold)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, font_name, size_pt, bold)


def set_paragraph_format(paragraph, *, title: bool = False, indent: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(30 if title else 28)
    fmt.first_line_indent = Pt(0 if title or not indent else 32)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if title else WD_ALIGN_PARAGRAPH.JUSTIFY


def classify_line(line: str, is_title: bool) -> tuple[str, str, float, bool, bool]:
    stripped = line.lstrip("　 ").rstrip()
    if is_title:
        return stripped, TITLE_FONT, 22, False, False
    if H1_RE.match(stripped):
        return stripped, H1_FONT, 16, False, True
    if H2_RE.match(stripped):
        return stripped, H2_FONT, 16, True, True
    if H3_RE.match(stripped):
        return stripped, BODY_FONT, 16, True, True
    if H4_RE.match(stripped):
        return stripped, BODY_FONT, 16, False, True
    if META_RE.match(stripped):
        return stripped, BODY_FONT, 16, False, True
    return stripped, BODY_FONT, 16, False, True


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)


def build_docx(input_text: Path, output_docx: Path) -> None:
    text = input_text.read_text(encoding="utf-8-sig")
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    content_lines = [line for line in lines if line.strip()]
    if not content_lines:
        raise ValueError("Input minutes text is empty.")

    document = Document()
    configure_document(document)

    first = True
    for line in content_lines:
        clean, font_name, size_pt, bold, indent = classify_line(line, first)
        paragraph = document.add_paragraph()
        set_paragraph_format(paragraph, title=first, indent=indent)
        add_text_runs(paragraph, clean, font_name, size_pt, bold)
        first = False

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def run_text(run) -> str:
    return "".join(node.text or "" for node in run.xpath(".//w:t", namespaces=NS))


def run_fonts(run) -> set[str]:
    rfonts = run.find("w:rPr/w:rFonts", namespaces=NS)
    if rfonts is None:
        return set()
    return rfonts_values(rfonts)


def rfonts_values(rfonts) -> set[str]:
    values = set()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        value = rfonts.get(f"{{{W_NS}}}{attr}")
        if value:
            values.add(value)
    return values


def attr_val(element, name: str = "val") -> str | None:
    if element is None:
        return None
    return element.get(f"{{{W_NS}}}{name}")


def style_context(style_root):
    styles: dict[str, tuple[set[str], str | None]] = {}
    doc_default_fonts: set[str] = set()
    if style_root is None:
        return styles, doc_default_fonts

    default_rfonts = style_root.find("w:docDefaults/w:rPrDefault/w:rPr/w:rFonts", namespaces=NS)
    if default_rfonts is not None:
        doc_default_fonts = rfonts_values(default_rfonts)

    for style in style_root.xpath(".//w:style", namespaces=NS):
        style_id = attr_val(style, "styleId")
        if not style_id:
            continue
        rfonts = style.find("w:rPr/w:rFonts", namespaces=NS)
        fonts = rfonts_values(rfonts) if rfonts is not None else set()
        based_on = attr_val(style.find("w:basedOn", namespaces=NS))
        styles[style_id] = (fonts, based_on)
    return styles, doc_default_fonts


def inherited_style_fonts(style_id: str | None, styles: dict[str, tuple[set[str], str | None]]) -> set[str]:
    fonts: set[str] = set()
    seen: set[str] = set()
    current = style_id
    while current and current not in seen:
        seen.add(current)
        style_fonts, based_on = styles.get(current, (set(), None))
        fonts.update(style_fonts)
        current = based_on
    return fonts


def audit_docx(path: Path) -> list[str]:
    issues: list[str] = []
    parts = ["word/document.xml"]
    with ZipFile(path) as zf:
        style_root = None
        if "word/styles.xml" in zf.namelist():
            style_root = etree.fromstring(zf.read("word/styles.xml"))
        styles, doc_default_fonts = style_context(style_root)

        parts.extend(
            name
            for name in zf.namelist()
            if re.match(r"word/(header|footer)\d+\.xml$", name)
        )
        for part_name in parts:
            root = etree.fromstring(zf.read(part_name))
            run_index = 0
            for paragraph in root.xpath(".//w:p", namespaces=NS):
                paragraph_style = attr_val(paragraph.find("w:pPr/w:pStyle", namespaces=NS))
                paragraph_fonts = inherited_style_fonts(paragraph_style, styles)
                for run in paragraph.xpath(".//w:r", namespaces=NS):
                    run_index += 1
                    text = run_text(run)
                    if not text:
                        continue
                    direct_fonts = run_fonts(run)
                    run_style = attr_val(run.find("w:rPr/w:rStyle", namespaces=NS))
                    effective_fonts = set(direct_fonts)
                    if not effective_fonts:
                        effective_fonts.update(inherited_style_fonts(run_style, styles))
                        effective_fonts.update(paragraph_fonts)
                        effective_fonts.update(doc_default_fonts)

                    has_digit = any(ch.isdigit() for ch in text)
                    if has_digit:
                        if text.strip() and not text.strip().isdigit():
                            issues.append(f"{part_name} run {run_index}: digit mixed with non-digit text {text!r}")
                        if DIGIT_FONT not in effective_fonts:
                            issues.append(
                                f"{part_name} run {run_index}: digit run lacks {DIGIT_FONT}: "
                                f"{text!r} fonts={sorted(effective_fonts)}"
                            )
                    else:
                        disallowed = effective_fonts - ALLOWED_CJK_FONTS
                        if effective_fonts and disallowed:
                            issues.append(
                                f"{part_name} run {run_index}: disallowed font(s) "
                                f"{sorted(disallowed)} for text {text!r}"
                            )
    return issues


def main() -> int:
    parser = argparse.ArgumentParser(description="Build or audit WeiYang-style meeting-minute DOCX files.")
    parser.add_argument("input", help="Input UTF-8 minutes text, or DOCX when --audit-only is used.")
    parser.add_argument("output", nargs="?", help="Output DOCX path when building.")
    parser.add_argument("--audit", action="store_true", help="Audit the generated DOCX after building.")
    parser.add_argument("--audit-only", action="store_true", help="Audit an existing DOCX and do not build.")
    args = parser.parse_args()

    input_path = Path(args.input)
    if args.audit_only:
        issues = audit_docx(input_path)
        if issues:
            print("\n".join(issues), file=sys.stderr)
            return 1
        print("DOCX font audit passed.")
        return 0

    if not args.output:
        parser.error("output is required unless --audit-only is used")

    output_path = Path(args.output)
    build_docx(input_path, output_path)
    if args.audit:
        issues = audit_docx(output_path)
        if issues:
            print("\n".join(issues), file=sys.stderr)
            return 1
        print("DOCX built and font audit passed.")
    else:
        print(f"DOCX written: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
