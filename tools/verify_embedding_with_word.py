#!/usr/bin/env python3
"""决定性验证：嵌入 DOCX 的字体是否**真的被渲染器使用**。

为什么需要单独一个工具：单元测试只能断言 fontTable 里写了 `w:charset`、反混淆能
还原出有效字体——这些都是**间接证据**。它们证明不了 Word/LibreOffice 打开文件时
会不会真的采用这份嵌入字体。正是这个盲区，曾让一版"看起来成功、实际无效"的实现
通过了全部校验：反混淆自洽、verify ok=true，而 Word 始终在静默回退系统字体。

原理（必须带阴性对照才成立）：
  1. 把随附字体的**内部字体名**改成本机未安装的名字（如 VerifyProbeFS）；
  2. 用该名字排版中文，导出 PDF，看 PDF 里实际用的是哪个字体；
     - 阴性对照（不嵌入）：字体没装，必然回退系统字体。若这一步没回退，
       说明实验不敏感（比如该名字碰巧已安装），结论作废；
     - 嵌入后：若不再回退，则该字形只可能来自嵌入的那一份 → 证明成立。
       Word 会把它改名为 `___WRD_EMBED_SUB_*`，LibreOffice 保留原名。

依赖 fontTools 与 pypdf，以及 Microsoft Word（Windows COM）或 LibreOffice。
属开发期验证工具，不随技能分发；改动任何 embed_fonts 后建议跑一次。

用法：
    python tools/verify_embedding_with_word.py --skill gongsi-qingkuang
    python tools/verify_embedding_with_word.py --skill meeting-minutes-pro \
        --renderer libreoffice
"""

from __future__ import annotations

import argparse
import importlib.util
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
PROBE_NAME = "VerifyProbeFS"
FALLBACK_HINTS = ("SimSun", "NSimSun", "MicrosoftYaHei", "DejaVuSans", "FangSong")


def load_embed_fonts(skill: str):
    path = REPO / skill / "scripts" / "embed_fonts.py"
    if not path.is_file():
        raise SystemExit(f"找不到 {path}")
    spec = importlib.util.spec_from_file_location(f"ef_{skill.replace('-', '_')}", path)
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def make_probe_font(source: Path, dest: Path) -> None:
    from fontTools.ttLib import TTFont

    font = TTFont(source)
    for record in font["name"].names:
        if record.nameID in (1, 3, 4, 6):
            record.string = PROBE_NAME
    font.save(dest)


def make_docx(dest: Path) -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    document = Document()
    run = document.add_paragraph().add_run("嵌入字体验证中文样本")
    run.font.name = PROBE_NAME
    run.font.size = Pt(24)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), PROBE_NAME)
    document.save(dest)


def render_with_word(docx: Path, pdf: Path) -> bool:
    if sys.platform != "win32":
        return False
    script = f"""
$w = New-Object -ComObject Word.Application
$w.Visible = $false
try {{
  $d = $w.Documents.Open('{docx.resolve()}', $false, $true)
  $d.SaveAs2('{pdf.resolve()}', 17)
  $d.Close($false)
}} finally {{ try {{ $w.Quit() }} catch {{}} }}
"""
    handle = Path(tempfile.mktemp(suffix=".ps1"))
    handle.write_text(script, encoding="utf-8-sig")
    try:
        subprocess.run(["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass",
                        "-File", str(handle)], capture_output=True, timeout=180)
    finally:
        handle.unlink(missing_ok=True)
    return pdf.is_file()


def render_with_soffice(docx: Path, pdf: Path) -> bool:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    for candidate in (Path("C:/Program Files/LibreOffice/program/soffice.exe"),
                      Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")):
        if soffice is None and candidate.is_file():
            soffice = str(candidate)
    if not soffice:
        return False
    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run([soffice, "--headless", "--norestore", "--convert-to", "pdf",
                        "--outdir", tmp, str(docx.resolve())],
                       capture_output=True, timeout=180)
        produced = Path(tmp) / (docx.stem + ".pdf")
        if not produced.is_file():
            return False
        shutil.copyfile(produced, pdf)
    return True


def pdf_font_names(pdf: Path) -> set[str]:
    from pypdf import PdfReader

    names: set[str] = set()
    for page in PdfReader(str(pdf)).pages:
        fonts = (page.get("/Resources") or {}).get("/Font") or {}
        for font in fonts.values():
            font = font.get_object()
            if font.get("/BaseFont"):
                names.add(str(font["/BaseFont"]).lstrip("/").split("+")[-1])
            for descendant in font.get("/DescendantFonts") or []:
                descendant = descendant.get_object()
                if descendant.get("/BaseFont"):
                    names.add(str(descendant["/BaseFont"]).lstrip("/").split("+")[-1])
    return names


def uses_embedded(names: set[str]) -> bool:
    return any("WRD_EMBED" in name or PROBE_NAME in name for name in names)


def fell_back(names: set[str]) -> bool:
    return any(name in FALLBACK_HINTS for name in names)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--skill", default="gongsi-qingkuang",
                        help="要验证其 embed_fonts 实现的技能目录名")
    parser.add_argument("--renderer", choices=("word", "libreoffice"), default="word")
    parser.add_argument("--keep", action="store_true", help="保留中间产物便于排查")
    args = parser.parse_args()

    embed = load_embed_fonts(args.skill)
    source = REPO / args.skill / "assets" / "fonts" / "simfang.ttf"
    if not source.is_file():
        raise SystemExit(f"找不到随附字体：{source}")
    render = render_with_word if args.renderer == "word" else render_with_soffice

    work = Path(tempfile.mkdtemp(prefix="embed-verify-"))
    try:
        probe = work / "probe.ttf"
        make_probe_font(source, probe)
        print(f"技能：{args.skill}    渲染器：{args.renderer}")
        print(f"探针字体内部名：{PROBE_NAME}（本机应未安装）")

        control = work / "control.docx"
        make_docx(control)
        if not render(control, work / "control.pdf"):
            raise SystemExit("渲染失败：请确认已安装 Microsoft Word 或 LibreOffice")
        control_fonts = pdf_font_names(work / "control.pdf")
        print(f"\nA) 阴性对照（不嵌入） → {sorted(control_fonts)}")
        if not fell_back(control_fonts):
            print("   ✗ 未回退到系统字体，说明实验不敏感（该字体名可能已安装），结论作废")
            return 1
        print("   ✓ 已回退到系统字体，实验敏感")

        embedded = work / "embedded.docx"
        make_docx(embedded)
        embed.embed_fonts_into_docx(embedded, {PROBE_NAME: probe})
        verified = embed.verify_embedded_fonts(embedded)["ok"]
        if not render(embedded, work / "embedded.pdf"):
            raise SystemExit("渲染失败")
        embedded_fonts = pdf_font_names(work / "embedded.pdf")
        print(f"\nB) 嵌入后（verify ok={verified}） → {sorted(embedded_fonts)}")

        if uses_embedded(embedded_fonts) and not fell_back(embedded_fonts):
            print("   ✓ 渲染器使用了嵌入字体")
            print("\n结论：嵌入生效——未装该字体的机器也能忠实呈现。")
            return 0
        print("   ✗ 渲染器仍在回退系统字体：嵌入未生效")
        print("\n结论：嵌入无效。检查 <w:font> 是否带 w:charset 等字体描述"
              "（缺 charset 时 Word 会静默忽略嵌入字体）。")
        return 1
    finally:
        if args.keep:
            print(f"\n中间产物保留于：{work}")
        else:
            shutil.rmtree(work, ignore_errors=True)


if __name__ == "__main__":
    raise SystemExit(main())
