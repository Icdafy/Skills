#!/usr/bin/env python3
"""Build a source-derived Chinese SOE post-investment report DOCX.

The input is a UTF-8 JSON object shaped like ``assets/report-spec.example.json``.
This renderer is deterministic and contains no project-specific business logic. It
does not bundle or embed fonts; install properly licensed fonts on the rendering
machine and run ``font_preflight.py`` before final production.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
from pathlib import Path
import sys
from typing import Any, Iterable
import uuid

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - clear CLI dependency failure
    raise SystemExit("python-docx is required: python -m pip install python-docx") from exc


FONT_TITLE = "方正小标宋简体"
FONT_BODY = "仿宋_GB2312"
FONT_KAITI = "楷体_GB2312"
FONT_HEITI = "黑体"
FONT_SONG = "宋体"
FONT_WESTERN = "Times New Roman"

TITLE_PT = 22
# The supplied Word template uses a visually large 68 pt red-head run compressed
# horizontally with w:fitText, rather than shrinking the font or allowing a wrap.
REDHEAD_PT = 68
BODY_PT = 16
TABLE_PT = 10.5
FOOTER_PT = 14
IMPRINT_PT = 14
CAPTION_PT = 12
TITLE_LINE_PT = 30
BODY_LINE_PT = 28
TABLE_LINE_PT = 18
TWO_CHARS_OOXML = 200
CONTENT_DXA = 8844


def configure_utf8_stdio() -> None:
    """Prefer readable Unicode CLI output without affecting module imports."""

    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, OSError):
            pass


def spec_fingerprint(spec: dict[str, Any]) -> str:
    canonical = json.dumps(spec, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def _rpr(obj: Any) -> Any:
    return obj._element.get_or_add_rPr()


def set_font(
    run: Any,
    east_asia: str,
    size: float,
    *,
    bold: bool = False,
    color: str | None = None,
    western: str = FONT_WESTERN,
) -> None:
    run.font.name = western
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    fonts = _rpr(run).find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        _rpr(run).insert(0, fonts)
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:cs"), western)
    fonts.set(qn("w:eastAsia"), east_asia)


def fit_text(run: Any, width_twips: int, horizontal_scale: int = 37) -> None:
    """Keep the red-head issuer on one line using the source template's OOXML.

    ``w:fitText`` constrains the run to an exact width while ``w:w`` records the
    horizontal character scale. This mirrors the supplied template's red-head
    construction and remains editable in Word.
    """
    properties = _rpr(run)
    width = properties.find(qn("w:w"))
    if width is None:
        width = OxmlElement("w:w")
        properties.append(width)
    width.set(qn("w:val"), str(horizontal_scale))
    fitted = properties.find(qn("w:fitText"))
    if fitted is None:
        fitted = OxmlElement("w:fitText")
        properties.append(fitted)
    fitted.set(qn("w:val"), str(width_twips))
    fitted.set(qn("w:id"), "1")


def format_paragraph(
    paragraph: Any,
    *,
    align: Any = WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_pt: float = BODY_LINE_PT,
    first_line_chars: int | None = TWO_CHARS_OOXML,
    left_indent_pt: float = 0,
    right_indent_pt: float = 0,
    before_pt: float = 0,
    after_pt: float = 0,
    keep_with_next: bool = False,
) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    fmt.left_indent = Pt(left_indent_pt)
    fmt.right_indent = Pt(right_indent_pt)
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True

    # Use Word's character-based first-line indent so “首行空两字” follows the
    # paragraph font size instead of being frozen at a 32 pt physical distance.
    ppr = paragraph._p.get_or_add_pPr()
    indent = ppr.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        ppr.append(indent)
    for attribute in ("firstLine", "firstLineChars", "hanging", "hangingChars"):
        indent.attrib.pop(qn(f"w:{attribute}"), None)
    if first_line_chars is not None:
        indent.set(qn("w:firstLineChars"), str(first_line_chars))


def set_outline_level(paragraph: Any, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_paragraph(
    doc: Document,
    text: str,
    *,
    font: str = FONT_BODY,
    size: float = BODY_PT,
    bold: bool = False,
    align: Any = WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_pt: float = BODY_LINE_PT,
    first_line_chars: int | None = TWO_CHARS_OOXML,
    left_indent_pt: float = 0,
    right_indent_pt: float = 0,
    before_pt: float = 0,
    after_pt: float = 0,
    keep_with_next: bool = False,
    color: str | None = None,
    outline: int | None = None,
) -> Any:
    paragraph = doc.add_paragraph()
    format_paragraph(
        paragraph,
        align=align,
        line_pt=line_pt,
        first_line_chars=first_line_chars,
        left_indent_pt=left_indent_pt,
        right_indent_pt=right_indent_pt,
        before_pt=before_pt,
        after_pt=after_pt,
        keep_with_next=keep_with_next,
    )
    run = paragraph.add_run(str(text))
    set_font(run, font, size, bold=bold, color=color)
    if outline is not None:
        set_outline_level(paragraph, outline)
    return paragraph


def add_page_field(paragraph: Any) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(FOOTER_PT)

    left = paragraph.add_run("- ")
    set_font(left, FONT_SONG, FOOTER_PT, western=FONT_SONG)

    field = paragraph.add_run()
    set_font(field, FONT_SONG, FOOTER_PT, western=FONT_SONG)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, result, end):
        field._r.append(element)

    right = paragraph.add_run(" -")
    set_font(right, FONT_SONG, FOOTER_PT, western=FONT_SONG)


def setup_document(doc: Document, layout: dict[str, Any] | None = None) -> None:
    layout = layout or {}
    section = doc.sections[0]
    page_width = float(layout.get("page_width_cm", 21.0))
    page_height = float(layout.get("page_height_cm", 29.7))
    section.orientation = WD_ORIENT.PORTRAIT if page_width <= page_height else WD_ORIENT.LANDSCAPE
    section.page_width = Cm(page_width)
    section.page_height = Cm(page_height)
    section.top_margin = Cm(float(layout.get("top_margin_cm", 3.7)))
    section.bottom_margin = Cm(float(layout.get("bottom_margin_cm", 3.5)))
    section.left_margin = Cm(float(layout.get("left_margin_cm", 2.8)))
    section.right_margin = Cm(float(layout.get("right_margin_cm", 2.6)))
    section.header_distance = Cm(float(layout.get("header_distance_cm", 1.5)))
    section.footer_distance = Cm(float(layout.get("footer_distance_cm", 1.75)))

    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))

    odd_footer = section.footer.paragraphs[0]
    odd_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(odd_footer)

    even_footer = section.even_page_footer.paragraphs[0]
    even_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_page_field(even_footer)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_WESTERN
    normal.font.size = Pt(BODY_PT)
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_WESTERN)
    fonts.set(qn("w:hAnsi"), FONT_WESTERN)
    fonts.set(qn("w:eastAsia"), FONT_BODY)

    doc.core_properties.language = "zh-CN"
    doc.core_properties.subject = "年度股权投资项目投后情况报告"


def _set_table_full_width(table: Any) -> None:
    table.autofit = True
    props = table._tbl.tblPr
    width = props.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        props.append(width)
    width.set(qn("w:w"), "5000")
    width.set(qn("w:type"), "pct")
    layout = props.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        props.append(layout)
    layout.set(qn("w:type"), "autofit")


def _set_cell_width_pct(cell: Any, pct: int) -> None:
    props = cell._tc.get_or_add_tcPr()
    for old in props.findall(qn("w:tcW")):
        props.remove(old)
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(pct))
    width.set(qn("w:type"), "pct")
    props.insert(0, width)


def _set_cell_borders(cell: Any, *, color: str = "808080", size: int = 6) -> None:
    props = cell._tc.get_or_add_tcPr()
    borders = props.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def _set_cell_margins(cell: Any, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for edge_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        edge = margins.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            margins.append(edge)
        edge.set(qn("w:w"), str(value))
        edge.set(qn("w:type"), "dxa")


def _set_row_repeat(row: Any) -> None:
    props = row._tr.get_or_add_trPr()
    if props.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        props.append(repeat)


def _set_row_cant_split(row: Any) -> None:
    props = row._tr.get_or_add_trPr()
    if props.find(qn("w:cantSplit")) is None:
        props.append(OxmlElement("w:cantSplit"))


def _fill_cell(cell: Any, text: Any, *, header: bool = False, align: str = "center") -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(TABLE_LINE_PT)
    run = paragraph.add_run("" if text is None else str(text))
    set_font(run, FONT_BODY, TABLE_PT, bold=header)
    _set_cell_borders(cell)
    _set_cell_margins(cell)


def add_table(doc: Document, block: dict[str, Any]) -> Any:
    header = block.get("header") or []
    rows = block.get("rows") or []
    column_count = len(header) if header else max((len(row) for row in rows), default=1)
    widths = block.get("widths") or [1] * column_count
    if len(widths) != column_count or sum(widths) <= 0:
        widths = [1] * column_count

    table = doc.add_table(rows=(1 if header else 0) + len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)

    row_offset = 0
    if header:
        for index, value in enumerate(header):
            _fill_cell(table.rows[0].cells[index], value, header=True)
        _set_row_repeat(table.rows[0])
        _set_row_cant_split(table.rows[0])
        row_offset = 1

    body_align = block.get("align", "center")
    for row_index, values in enumerate(rows, start=row_offset):
        for column_index in range(column_count):
            value = values[column_index] if column_index < len(values) else ""
            align = body_align[column_index] if isinstance(body_align, list) and column_index < len(body_align) else body_align
            _fill_cell(table.rows[row_index].cells[column_index], value, align=str(align))
        _set_row_cant_split(table.rows[row_index])

    total = float(sum(widths))
    percentages = [max(1, round(float(width) / total * 5000)) for width in widths]
    for row in table.rows:
        for column_index, cell in enumerate(row.cells):
            _set_cell_width_pct(cell, percentages[column_index])

    spacer = doc.add_paragraph()
    format_paragraph(spacer, first_line_chars=None, line_pt=8, after_pt=0)
    return table


def _redhead_rule(cell: Any, *, red: str = "FF0000") -> None:
    props = cell._tc.get_or_add_tcPr()
    borders = props.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge_name in ("top", "left", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), red)


def add_redhead(doc: Document, metadata: dict[str, Any]) -> None:
    company = str(metadata.get("company") or "〔公司全称〕")
    paragraph = doc.add_paragraph()
    format_paragraph(
        paragraph,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_pt=48,
        first_line_chars=None,
        before_pt=86,
        after_pt=20,
    )
    run = paragraph.add_run(f"{company}文件")
    set_font(run, FONT_TITLE, REDHEAD_PT, bold=True, color="FF0000")
    fit_text(run, 8195)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)
    left, right = table.rows[0].cells
    for cell in (left, right):
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _redhead_rule(cell)
        _set_cell_width_pct(cell, 2500)
        _set_cell_margins(cell, top=0, start=70, bottom=90, end=70)

    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_p.paragraph_format.line_spacing = Pt(BODY_LINE_PT)
    left_run = left_p.add_run(str(metadata.get("document_number") or "〔文号〕"))
    set_font(left_run, FONT_BODY, BODY_PT)

    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.line_spacing = Pt(BODY_LINE_PT)
    right_run = right_p.add_run(str(metadata.get("signer") or ""))
    set_font(right_run, FONT_KAITI, BODY_PT)

    title = doc.add_paragraph()
    format_paragraph(
        title,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_pt=TITLE_LINE_PT,
        first_line_chars=None,
        before_pt=24,
        after_pt=24,
        keep_with_next=True,
    )
    line1 = title.add_run(company)
    set_font(line1, FONT_TITLE, TITLE_PT)
    line1.add_break(WD_BREAK.LINE)
    line2 = title.add_run(f"关于{metadata.get('report_year') or '〔年度〕'}年度股权投资项目投后情况报告")
    set_font(line2, FONT_TITLE, TITLE_PT)


def add_heading(doc: Document, text: str, level: int) -> Any:
    if level == 1:
        return add_paragraph(
            doc,
            text,
            font=FONT_HEITI,
            size=BODY_PT,
            bold=False,
            before_pt=0,
            after_pt=0,
            keep_with_next=True,
            outline=0,
        )
    if level == 2:
        return add_paragraph(
            doc,
            text,
            font=FONT_KAITI,
            size=BODY_PT,
            bold=True,
            before_pt=0,
            after_pt=0,
            keep_with_next=True,
            outline=1,
        )
    if level == 3:
        return add_paragraph(
            doc,
            text,
            font=FONT_BODY,
            size=BODY_PT,
            bold=True,
            before_pt=0,
            after_pt=0,
            keep_with_next=True,
            outline=2,
        )
    return add_paragraph(
        doc,
        text,
        font=FONT_BODY,
        size=BODY_PT,
        bold=False,
        before_pt=0,
        after_pt=0,
        keep_with_next=True,
        outline=3,
    )


def render_blocks(doc: Document, blocks: Iterable[dict[str, Any]]) -> None:
    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    for block in blocks:
        block_type = str(block.get("type") or "p").lower()
        if block_type.startswith("_"):
            continue
        if block_type in {"h1", "h2", "h3", "h4"}:
            add_heading(doc, str(block.get("text") or ""), int(block_type[1]))
        elif block_type == "p":
            add_paragraph(
                doc,
                str(block.get("text") or ""),
                bold=bool(block.get("bold")),
                align=alignments.get(str(block.get("align") or "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY),
            )
        elif block_type == "caption":
            add_paragraph(
                doc,
                str(block.get("text") or ""),
                font=FONT_HEITI,
                size=CAPTION_PT,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_line_chars=None,
                keep_with_next=True,
            )
        elif block_type == "tnote":
            add_paragraph(
                doc,
                str(block.get("text") or ""),
                size=TABLE_PT,
                align=alignments.get(str(block.get("align") or "left"), WD_ALIGN_PARAGRAPH.LEFT),
                line_pt=TABLE_LINE_PT,
                first_line_chars=None,
            )
        elif block_type == "table":
            add_table(doc, block)
        elif block_type == "pagebreak":
            doc.add_page_break()
        elif block_type == "blank":
            add_paragraph(doc, "", first_line_chars=None)
        else:
            raise ValueError(f"Unsupported block type: {block_type}")


def add_attachment_list(doc: Document, attachments: list[dict[str, Any]]) -> None:
    if not attachments:
        return
    add_paragraph(doc, "", first_line_chars=None)
    first = attachments[0]
    add_paragraph(doc, f"附件：1.{first.get('title', '')}")
    for number, attachment in enumerate(attachments[1:], start=2):
        add_paragraph(
            doc,
            f"{number}.{attachment.get('title', '')}",
            first_line_chars=None,
            left_indent_pt=80,
        )


def add_signature(doc: Document, metadata: dict[str, Any]) -> None:
    add_paragraph(doc, "", first_line_chars=None)
    add_paragraph(doc, "", first_line_chars=None)
    issuer = str(metadata.get("issuer") or metadata.get("company") or "")
    issue_date = str(metadata.get("issue_date") or "")
    if issuer:
        add_paragraph(
            doc,
            issuer,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent_pt=64,
            first_line_chars=None,
        )
    if issue_date:
        add_paragraph(
            doc,
            issue_date,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            right_indent_pt=64,
            first_line_chars=None,
        )
    contact_name = str(metadata.get("contact_name") or "").strip()
    contact_phone = str(metadata.get("contact_phone") or "").strip()
    if contact_name or contact_phone:
        parts = []
        if contact_name:
            parts.append(f"联系人：{contact_name}")
        if contact_phone:
            parts.append(f"联系电话：{contact_phone}")
        add_paragraph(
            doc,
            f"（{'  '.join(parts)}）",
            font=FONT_KAITI,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_chars=None,
        )


def _display_units(text: str) -> float:
    """Approximate rendered width in em: CJK glyphs are full width, others half."""

    return sum(1.0 if ord(character) > 0x2E7F else 0.5 for character in text)


def _imprint_rule(cell: Any) -> None:
    """版记 uses a plain rule above and below the 印发机关／印发日期 row."""

    props = cell._tc.get_or_add_tcPr()
    borders = props.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge_name in ("left", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")
    for edge_name in ("top", "bottom"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "12")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), "auto")


def add_imprint(doc: Document, metadata: dict[str, Any]) -> None:
    """Render 版记（印发机关和印发日期）on the last page, per GB/T 9704-2012.

    Both fields are optional; the validator requires them to be supplied
    together, so an incomplete pair renders nothing here.
    """

    printer = str(metadata.get("printer") or "").strip()
    print_date = str(metadata.get("print_date") or "").strip()
    if not printer or not print_date:
        return
    add_paragraph(doc, "", first_line_chars=None)
    add_paragraph(doc, "", first_line_chars=None)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)
    left, right = table.rows[0].cells
    # Split the row by rendered width so a long 印发机关 name keeps one line.
    date_text = f"{print_date}印发"
    left_units = _display_units(printer)
    right_units = _display_units(date_text)
    total_units = left_units + right_units or 1
    left_pct = min(3500, max(1500, round(5000 * left_units / total_units)))
    for cell, width in ((left, left_pct), (right, 5000 - left_pct)):
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _imprint_rule(cell)
        _set_cell_width_pct(cell, width)
        _set_cell_margins(cell, top=40, start=70, bottom=40, end=70)

    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_p.paragraph_format.line_spacing = Pt(BODY_LINE_PT)
    set_font(left_p.add_run(printer), FONT_BODY, IMPRINT_PT)

    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.paragraph_format.line_spacing = Pt(BODY_LINE_PT)
    set_font(right_p.add_run(date_text), FONT_BODY, IMPRINT_PT)


def add_attachment(doc: Document, attachment: dict[str, Any], fallback_number: int) -> None:
    doc.add_page_break()
    number = attachment.get("number") or fallback_number
    add_paragraph(
        doc,
        f"附件{number}",
        first_line_chars=None,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        keep_with_next=True,
    )
    add_paragraph(
        doc,
        str(attachment.get("title") or "投后管理报告"),
        font=FONT_TITLE,
        size=TITLE_PT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_pt=TITLE_LINE_PT,
        first_line_chars=None,
        after_pt=16,
        keep_with_next=True,
    )
    render_blocks(doc, attachment.get("blocks") or [])


def build_report(spec: dict[str, Any], output: Path, *, force: bool = False) -> Path:
    metadata = spec.get("document") or {}
    doc = Document()
    setup_document(doc, spec.get("layout") or {})
    doc.core_properties.title = f"{metadata.get('company', '')}关于{metadata.get('report_year', '')}年度股权投资项目投后情况报告"
    doc.core_properties.identifier = f"soe-post-investment-report:sha256:{spec_fingerprint(spec)}"

    add_redhead(doc, metadata)
    recipient = str(metadata.get("recipient") or "")
    if recipient:
        add_paragraph(doc, recipient, first_line_chars=None, keep_with_next=True)
    legal_basis = str(metadata.get("legal_basis") or "")
    if legal_basis:
        add_paragraph(doc, legal_basis)

    render_blocks(doc, spec.get("main_blocks") or [])
    attachments = sorted(spec.get("attachments") or [], key=lambda item: int(item.get("number") or 0))
    add_attachment_list(doc, attachments)
    add_signature(doc, metadata)

    for index, attachment in enumerate(attachments, start=1):
        add_attachment(doc, attachment, index)

    add_imprint(doc, metadata)

    output = output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    if output.exists() and not force:
        raise FileExistsError(f"Output already exists (use --force to replace it): {output}")
    staged = output.parent / f".{output.name}.tmp-{uuid.uuid4().hex}"
    try:
        doc.save(staged)
        os.replace(staged, output)
    finally:
        staged.unlink(missing_ok=True)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("spec", type=Path, help="UTF-8 report specification JSON")
    parser.add_argument("output", type=Path, help="Output .docx path")
    parser.add_argument("--template-mode", action="store_true", help="allow a synthetic template specification")
    parser.add_argument("--force", action="store_true", help="replace an existing output atomically")
    return parser.parse_args()


def main() -> int:
    configure_utf8_stdio()
    args = parse_args()
    try:
        spec = json.loads(args.spec.read_text(encoding="utf-8-sig"))
    except Exception as exc:
        print(f"Specification could not be read: {exc}", file=sys.stderr)
        return 2
    if not isinstance(spec, dict):
        print("Specification root must be a JSON object", file=sys.stderr)
        return 2
    from validate_report import Findings, print_findings, validate_spec

    findings = Findings()
    validate_spec(spec, findings, template_mode=args.template_mode)
    if findings.errors:
        print_findings(findings)
        return 2
    try:
        output = build_report(spec, args.output, force=args.force)
    except (OSError, TypeError, ValueError) as exc:
        print(str(exc), file=sys.stderr)
        return 2
    print(f"Generated: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
