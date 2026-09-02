#!/usr/bin/env python3
"""Validate a post-investment report specification, DOCX, and optional PDF."""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import sys
from collections import Counter
from dataclasses import dataclass, field
from datetime import date
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Iterable
from xml.etree import ElementTree as ET
from zipfile import ZipFile

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.table import CT_Tbl
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx is required: python -m pip install python-docx") from exc


def configure_utf8_stdio() -> None:
    """Prefer readable Unicode CLI output without affecting module imports."""

    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, OSError):
            pass


# The source template fixes the first-level headings and the first three
# second-level categories.  The SPV block is a *repeating* slot: the proofed
# source report carries two of them (（四）……SPV项目 and （五）……SPV项目), so the
# contract accepts one or more consecutive SPV slots numbered from 四 onwards.
FIXED_FIRST_H1 = "一、年度股权投资完成总体情况"
FIXED_LEADING_H2 = (
    "（一）存续基金",
    "（二）新设基金",
    "（三）参股公司",
)
FIXED_LAST_H1 = "二、重大投资项目进展情况"
SPV_SLOT_ORDINALS = ("四", "五", "六", "七", "八", "九", "十")
MIN_FIXED_MAIN_HEADINGS = len(FIXED_LEADING_H2) + 3
MAX_FIXED_MAIN_HEADINGS = MIN_FIXED_MAIN_HEADINGS + len(SPV_SLOT_ORDINALS) - 1


def spv_slot_pattern(ordinal: str) -> re.Pattern[str]:
    """Match one SPV slot heading, with or without a project-specific name."""

    return re.compile(rf"^（{ordinal}）(?:[^（）]+)?SPV项目$")


def fixed_main_heading_kinds(count: int) -> tuple[str, ...]:
    """Return the heading levels for a fixed contract of ``count`` headings."""

    return ("h1",) + ("h2",) * (count - 2) + ("h1",)


def default_fixed_main_headings(spv_slots: int = 1) -> tuple[str, ...]:
    headings = [FIXED_FIRST_H1, *FIXED_LEADING_H2]
    headings.extend(f"（{SPV_SLOT_ORDINALS[index]}）SPV项目" for index in range(spv_slots))
    headings.append(FIXED_LAST_H1)
    return tuple(headings)


REQUIRED_MAIN_HEADINGS = default_fixed_main_headings()
REQUIRED_MAIN_BLOCKS = tuple(
    zip(fixed_main_heading_kinds(len(REQUIRED_MAIN_HEADINGS)), REQUIRED_MAIN_HEADINGS)
)
REQUIRED_MAIN_HEADING_KINDS = fixed_main_heading_kinds(len(REQUIRED_MAIN_HEADINGS))
PLACEHOLDER_PATTERNS = (
    # Chinese square brackets are also the correct punctuation for document
    # years (for example, 控股字〔2026〕1号). Flag only semantic placeholder
    # tokens, never a bracketed four-digit year.
    re.compile(
        r"〔(?=[^〕]*(?:示例|待|公司全称|文号|年度|日期|姓名|联系人|联系电话|项目名称|数据|内容|占位|填写|补充|核实|确认))[^〕]+〕"
    ),
    re.compile(r"\{\{[^}]+\}\}"),
    re.compile(r"\b(?:TODO|TBD|FIXME|X{2,})\b", re.IGNORECASE),
    re.compile(r"待(?:补充|填写|核实|确认|完善)"),
)
HEADING_PATTERN = re.compile(r"^(?:[一二三四五六七八九十]+、|（[一二三四五六七八九十]+）|\d+[.．]|（\d+）)\s*\S+")
MATERIAL_NUMBER_PATTERN = re.compile(
    r"(?<![0-9A-Za-z])(?:\d{1,3}(?:,\d{3})+(?:\.\d+)?|\d+(?:\.\d+)?)\s*"
    r"(?:%|％|亿元|万元|元|万股|股|人|家|支|个|项|笔|次|户|平方米|亩|天|月|年)(?![0-9A-Za-z])"
)
BARE_NUMBER_PATTERN = re.compile(r"(?<![0-9A-Za-z])(?:\d{1,3}(?:,\d{3})+(?:\.\d+)?|\d+(?:\.\d+)?)(?![0-9A-Za-z])")
NUMBER_TOKEN_PATTERN = re.compile(r"(?<![0-9A-Za-z])[-+]?(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d+)?(?![0-9A-Za-z])")
TEMPLATE_MARKER_PATTERN = re.compile(
    r"(?:合成示例|本通用模板|示例|synthetic\s+(?:source|fact|example)|template\s+demonstration)",
    re.IGNORECASE,
)
STATUS_CLAIM_PATTERN = re.compile(
    r"(?:尚未|未(?:发生|形成|发现|新增|达到|完成|开展|确认|触发)|"
    r"已经|实际触发|已(?:完成|达到|形成|触发|设立|退出|进入|回收)|"
    r"发生|触发|达到|不存在|存在|稳定|承压|异常|不确定|清晰|增长|下降|增加|减少|新设|新增|退出)"
)
VALID_FACT_STATUSES = {"confirmed", "calculated", "conflicting", "stale", "missing"}
VALID_FACT_DESTINATIONS = {"main body", "attachment", "both", "excluded", "pending user decision"}
VALID_PROJECT_CATEGORIES = {"存续基金", "新设基金", "参股公司", "SPV项目"}
BLOCK_TYPES = {"h1", "h2", "h3", "h4", "p", "caption", "tnote", "table", "pagebreak", "blank"}
FIXED_CATEGORY_SECTIONS = tuple(
    zip(REQUIRED_MAIN_HEADINGS[1:-1], ("存续基金", "新设基金", "参股公司", "SPV项目"))
)
PUBLIC_PATTERNS = {
    "mobile phone number": re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    "landline phone number": re.compile(r"(?<!\d)0\d{2,3}[-—－ ]?\d{7,8}(?!\d)"),
    "email address": re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE),
    "Chinese identity number": re.compile(r"(?<!\d)\d{17}[\dXx](?!\d)"),
    "long account-like number": re.compile(r"(?<!\d)\d{16,19}(?!\d)"),
}
TEMPLATE_TYPOGRAPHY = {
    "redhead": ("方正小标宋简体", 68.0, True, None),
    "title": ("方正小标宋简体", 22.0, False, 30.0),
    "h1": ("黑体", 16.0, False, 28.0),
    "h2": ("楷体_GB2312", 16.0, True, 28.0),
    "h3": ("仿宋_GB2312", 16.0, True, 28.0),
    "h4": ("仿宋_GB2312", 16.0, False, 28.0),
    "p": ("仿宋_GB2312", 16.0, False, 28.0),
    # 表题 in the source template: 小四黑体, centered, on the 28 pt body grid.
    "caption": ("黑体", 12.0, False, 28.0),
    "tnote": ("仿宋_GB2312", 10.5, False, 18.0),
    "footer": ("宋体", 14.0, False, None),
    # 版记（印发机关和印发日期）: GB/T 9704-2012 prescribes 四号仿宋.
    "imprint": ("仿宋_GB2312", 14.0, False, 28.0),
}
DEFAULT_WESTERN_FONT = "Times New Roman"


@dataclass
class Findings:
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    def error(self, message: str) -> None:
        self.errors.append(message)

    def warning(self, message: str) -> None:
        self.warnings.append(message)

    def note(self, message: str) -> None:
        self.notes.append(message)


def normalize_heading(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", text).replace("．", ".")


def normalize_assertion_text(text: str) -> str:
    """Normalize only layout whitespace; punctuation remains evidence-bearing."""

    # The vertical bar is an explicit table-cell separator in fact assertions.
    # It is removed for matching but retained by numeric tokenization, preventing
    # adjacent cell numbers from collapsing into one synthetic number.
    return re.sub(r"[\s\u3000|]+", "", str(text or ""))


def required_main_blocks_from_document(document: dict[str, Any]) -> tuple[tuple[str, str], ...]:
    """Return the source-derived fixed heading contract declared by the spec."""

    values = document.get("fixed_main_headings")
    if not isinstance(values, list) or not (
        MIN_FIXED_MAIN_HEADINGS <= len(values) <= MAX_FIXED_MAIN_HEADINGS
    ):
        return REQUIRED_MAIN_BLOCKS
    return tuple(
        (kind, str(text or "").strip())
        for kind, text in zip(fixed_main_heading_kinds(len(values)), values)
    )


def fixed_category_sections_from_document(
    document: dict[str, Any],
) -> tuple[tuple[str, str], ...]:
    """Return (heading, category) pairs for every fixed category section."""

    blocks = required_main_blocks_from_document(document)
    sections = [
        (text, category)
        for (_kind, text), category in zip(blocks[1:4], ("存续基金", "新设基金", "参股公司"))
    ]
    sections.extend((text, "SPV项目") for _kind, text in blocks[4:-1])
    return tuple(sections)


def infer_fixed_main_blocks(
    heading_entries: list[tuple[str, str]],
) -> tuple[tuple[str, str], ...]:
    """Infer the fixed contract from DOCX headings when no spec is supplied.

    Without a specification there is no source snapshot to compare names against,
    so only the *shape* can be certified: the three fixed categories, then one or
    more consecutive SPV slots.  Slot names observed in the document are accepted
    as-is; reconciling them with the source template still requires --spec.
    """

    section_one: list[str] = []
    for kind, text in heading_entries:
        normalized = normalize_heading(text)
        if kind == "h1" and normalized == normalize_heading(FIXED_LAST_H1):
            break
        if kind == "h2":
            section_one.append(normalized)
    slots: list[str] = []
    for offset, ordinal in enumerate(SPV_SLOT_ORDINALS):
        index = len(FIXED_LEADING_H2) + offset
        if index >= len(section_one) or not spv_slot_pattern(ordinal).fullmatch(section_one[index]):
            break
        slots.append(section_one[index])
    if not slots:
        return REQUIRED_MAIN_BLOCKS
    headings = [FIXED_FIRST_H1, *FIXED_LEADING_H2, *slots, FIXED_LAST_H1]
    return tuple(zip(fixed_main_heading_kinds(len(headings)), headings))


def validate_heading_contract_list(
    values: Any,
    field_name: str,
    findings: Findings,
) -> list[str] | None:
    """Validate one source/effective fixed-heading contract list.

    The list is the first-level heading, the three fixed categories, one or more
    consecutive SPV slots, and the closing first-level heading.
    """

    if not isinstance(values, list) or not (
        MIN_FIXED_MAIN_HEADINGS <= len(values) <= MAX_FIXED_MAIN_HEADINGS
    ):
        findings.error(
            f"document.{field_name} must contain between {MIN_FIXED_MAIN_HEADINGS} and "
            f"{MAX_FIXED_MAIN_HEADINGS} headings (one first-level heading, 存续基金／新设基金／参股公司, "
            "one or more SPV slots, and the closing first-level heading)"
        )
        return None
    headings = [str(item or "").strip() for item in values]
    if any(not item for item in headings):
        findings.error(f"document.{field_name} cannot contain an empty heading")
        return None
    normalized = [normalize_heading(item) for item in headings]
    if len(normalized) != len(set(normalized)):
        findings.error(f"document.{field_name} must be unique after normalization")

    invariant = {0: FIXED_FIRST_H1}
    invariant.update({index + 1: text for index, text in enumerate(FIXED_LEADING_H2)})
    invariant[len(normalized) - 1] = FIXED_LAST_H1
    for index, expected in invariant.items():
        if normalized[index] != normalize_heading(expected):
            findings.error(
                f"document.{field_name}[{index}] must preserve the template heading: {expected}"
            )
    spv_headings = normalized[len(FIXED_LEADING_H2) + 1 : -1]
    for offset, heading in enumerate(spv_headings):
        ordinal = SPV_SLOT_ORDINALS[offset]
        if not spv_slot_pattern(ordinal).fullmatch(heading):
            findings.error(
                f"document.{field_name}[{offset + len(FIXED_LEADING_H2) + 1}] must preserve the "
                f"source SPV slot as （{ordinal}）SPV项目 or （{ordinal}）<项目名称>SPV项目"
            )
    return headings


CHINESE_DATE_PATTERN = re.compile(r"^(\d{4})年(\d{1,2})月(\d{1,2})日$")


def parse_chinese_date(value: str) -> date | None:
    """Parse the 公文 date form YYYY年M月D日; return None when it is not that form."""

    match = CHINESE_DATE_PATTERN.fullmatch(re.sub(r"[\s\u3000]+", "", str(value or "")))
    if match is None:
        return None
    try:
        return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    except ValueError:
        return None


def validate_document_dates(document: dict[str, Any], findings: Findings) -> None:
    """Keep 报告年度、数据截止日期、成文日期 and 印发日期 mutually consistent.

    A cutoff later than the issue date would report data that does not yet exist
    when the document is signed, so it is an error rather than a warning.
    """

    labels = {
        "cutoff_date": "数据截止日期",
        "issue_date": "成文日期",
        "print_date": "印发日期",
    }
    parsed: dict[str, date] = {}
    for key, label in labels.items():
        raw = str(document.get(key) or "").strip()
        if not raw:
            continue
        value = parse_chinese_date(raw)
        if value is None:
            findings.error(f"document.{key}（{label}）must use the 公文 form YYYY年M月D日: {raw}")
            continue
        parsed[key] = value

    cutoff = parsed.get("cutoff_date")
    issue = parsed.get("issue_date")
    print_date = parsed.get("print_date")
    if cutoff is not None and issue is not None and cutoff > issue:
        findings.error(
            "document.cutoff_date（数据截止日期）cannot be later than document.issue_date（成文日期）: "
            f"{document.get('cutoff_date')} > {document.get('issue_date')}"
        )
    if print_date is not None and issue is not None and print_date < issue:
        findings.error(
            "document.print_date（印发日期）cannot be earlier than document.issue_date（成文日期）: "
            f"{document.get('print_date')} < {document.get('issue_date')}"
        )

    report_year = str(document.get("report_year") or "").strip()
    if re.fullmatch(r"20\d{2}", report_year):
        for key in ("legal_basis", "recipient"):
            text = str(document.get(key) or "")
            years = {match.group(1) for match in re.finditer(r"(20\d{2})年度", text)}
            mismatched = sorted(years - {report_year})
            if mismatched:
                findings.error(
                    f"document.{key} names 年度 {', '.join(mismatched)} but document.report_year is {report_year}"
                )
        if cutoff is not None and str(cutoff.year) != report_year:
            findings.warning(
                f"document.cutoff_date year {cutoff.year} differs from document.report_year {report_year}; "
                "confirm the reporting period and the data cutoff are intended to differ"
            )


OFFICIAL_STYLE_RULES: tuple[tuple[re.Pattern[str], str], ...] = (
    (
        re.compile(r"截止(?=\d{4}年|\d{1,2}月|目前|当前|报告期末|期末)"),
        "公文表示时间点应使用“截至”，“截止”用于事项终止（如报名截止）",
    ),
    (
        re.compile(r"[0-9A-Za-z%％】》）][ \t　]+[，。；：、％%）》】]"),
        "数字或西文与其后的全角标点之间存在多余空格",
    ),
    (
        re.compile(r"\d[ \t　]+[%％]"),
        "数值与百分号之间存在多余空格",
    ),
    (
        re.compile(r"[一-鿿][()]|[(][一-鿿]"),
        "中文语境内混用半角圆括号，公文应使用全角（）",
    ),
    (
        re.compile(r"[一-鿿][,;:!?]"),
        "中文语境内混用半角标点，公文应使用全角标点",
    ),
)


def iter_spec_texts(spec: dict[str, Any]) -> Iterable[tuple[str, str]]:
    """Yield (location, text) for every author-written string in the spec."""

    document = spec.get("document") or {}
    for key in ("recipient", "legal_basis"):
        value = str(document.get(key) or "").strip()
        if value:
            yield f"document.{key}", value
    scopes: list[tuple[str, list[dict[str, Any]]]] = [
        ("main body", spec.get("main_blocks") or [])
    ]
    scopes.extend(
        (f"attachment {attachment.get('id')}", attachment.get("blocks") or [])
        for attachment in spec.get("attachments") or []
    )
    for scope, blocks in scopes:
        for index, block in enumerate(blocks, start=1):
            text = str(block.get("text") or "").strip()
            if text:
                yield f"{scope} block {index}", text
            for row_index, row in enumerate(block.get("rows") or [], start=1):
                for cell in row:
                    cell_text = str(cell or "").strip()
                    if cell_text:
                        yield f"{scope} block {index} row {row_index}", cell_text


def validate_official_style(spec: dict[str, Any], findings: Findings) -> None:
    """Report 公文 wording and punctuation slips as warnings.

    These are heuristics: 截止 is correct for a deadline, and a half-width
    bracket can be part of a legal name.  Every hit still needs a human decision.
    """

    reported: set[tuple[str, str]] = set()
    for location, text in iter_spec_texts(spec):
        for pattern, message in OFFICIAL_STYLE_RULES:
            match = pattern.search(text)
            if match is None:
                continue
            key = (location, message)
            if key in reported:
                continue
            reported.add(key)
            findings.warning(f"{location}: {message}（“{match.group(0).strip()}”）")


def validate_imprint_metadata(document: dict[str, Any], findings: Findings) -> None:
    """版记（印发机关和印发日期）is optional, but must be complete when used."""

    printer = str(document.get("printer") or "").strip()
    print_date = str(document.get("print_date") or "").strip()
    if bool(printer) != bool(print_date):
        findings.error(
            "document.printer（印发机关）and document.print_date（印发日期）must be provided together "
            "or both omitted"
        )


def validate_fixed_heading_contract(document: dict[str, Any], findings: Findings) -> None:
    """Lock effective headings to the source snapshot unless change was authorized."""

    source_headings = validate_heading_contract_list(
        document.get("source_fixed_main_headings"),
        "source_fixed_main_headings",
        findings,
    )
    effective_headings = validate_heading_contract_list(
        document.get("fixed_main_headings"),
        "fixed_main_headings",
        findings,
    )
    if source_headings is None or effective_headings is None:
        return
    changed = source_headings != effective_headings
    authorized = document.get("heading_change_authorized")
    if not isinstance(authorized, bool):
        findings.error("document.heading_change_authorized must be true or false")
        return
    note = str(document.get("heading_change_note") or "").strip()
    if changed and not authorized:
        findings.error(
            "Effective fixed headings differ from the source heading snapshot without "
            "document.heading_change_authorized=true"
        )
    if changed and (len(note) < 8 or re.fullmatch(r"(?i)(?:none|no\s+change|unchanged|无|无变更)", note)):
        findings.error("An authorized fixed-heading change requires a specific heading_change_note")
    if not changed and authorized:
        findings.error(
            "document.heading_change_authorized must be false when effective headings equal the source snapshot"
        )


def spec_fingerprint(spec: dict[str, Any]) -> str:
    canonical = json.dumps(spec, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def block_headings(blocks: Iterable[dict[str, Any]]) -> list[str]:
    return [str(block.get("text") or "").strip() for block in blocks if str(block.get("type") or "").lower() in {"h1", "h2", "h3", "h4"}]


def block_heading_entries(blocks: Iterable[dict[str, Any]]) -> list[tuple[str, str]]:
    return [
        (str(block.get("type") or "").lower(), str(block.get("text") or "").strip())
        for block in blocks
        if str(block.get("type") or "").lower() in {"h1", "h2", "h3", "h4"}
    ]


def find_placeholders(text: str) -> list[str]:
    found: list[str] = []
    for pattern in PLACEHOLDER_PATTERNS:
        found.extend(match.group(0) for match in pattern.finditer(text))
    return found


def validate_heading_scope(headings: list[str], scope: str, findings: Findings) -> None:
    normalized = [normalize_heading(heading) for heading in headings if heading]
    duplicates = sorted(key for key, count in Counter(normalized).items() if count > 1)
    for duplicate in duplicates:
        findings.error(f"Duplicate heading in {scope}: {duplicate}")
    for heading in headings:
        if not heading:
            findings.error(f"Empty heading in {scope}")


def validate_required_heading_sequence(
    heading_entries: list[tuple[str, str]],
    scope: str,
    findings: Findings,
    *,
    required_blocks: Iterable[tuple[str, str]] = REQUIRED_MAIN_BLOCKS,
) -> None:
    """Require the fixed framework's heading text, levels, counts, and order."""

    required_blocks = tuple(required_blocks)
    required_headings = tuple(text for _kind, text in required_blocks)
    normalized = [(kind, text.strip()) for kind, text in heading_entries]
    positions: list[int] = []
    sequence_is_complete = True
    for required_kind, required_text in required_blocks:
        key = (required_kind, required_text.strip())
        count = normalized.count(key)
        if count != 1:
            same_text_levels = sorted(
                kind for kind, text in normalized if text == required_text.strip()
            )
            level_note = f"; found at levels {same_text_levels}" if same_text_levels else ""
            findings.error(
                f"Required main heading must occur exactly once in {scope}: "
                f"{required_kind} {required_text} (found {count}{level_note})"
            )
            sequence_is_complete = False
        else:
            positions.append(normalized.index(key))
    if sequence_is_complete and positions != sorted(positions):
        findings.error(
            f"Fixed main headings are out of order in {scope}; expected: "
            + " -> ".join(required_headings)
        )

    h1_texts = [text for kind, text in normalized if kind == "h1"]
    expected_h1 = [required_blocks[0][1].strip(), required_blocks[-1][1].strip()]
    if h1_texts != expected_h1:
        findings.error(f"Unexpected first-level heading structure in {scope}")

    second_h1_key = ("h1", required_blocks[-1][1].strip())
    if second_h1_key in normalized:
        second_h1_index = normalized.index(second_h1_key)
        section_one_h2 = [text for kind, text in normalized[:second_h1_index] if kind == "h2"]
        expected_section_one_h2 = [text.strip() for kind, text in required_blocks if kind == "h2"]
        if section_one_h2 != expected_section_one_h2:
            findings.error(f"Unexpected second-level heading structure under section one in {scope}")


def chinese_ordinal(text: str) -> int | None:
    digits = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
    if text in digits:
        return digits[text]
    if text == "十":
        return 10
    if "十" in text:
        left, right = text.split("十", 1)
        tens = digits.get(left, 1) if left else 1
        ones = digits.get(right, 0) if right else 0
        return tens * 10 + ones
    return None


def heading_ordinal(kind: str, text: str) -> int | None:
    if kind == "h2":
        match = re.match(r"^（([一二三四五六七八九十]+)）", text)
        return chinese_ordinal(match.group(1)) if match else None
    if kind == "h3":
        match = re.match(r"^(\d+)[.．]", text)
        return int(match.group(1)) if match else None
    if kind == "h4":
        match = re.match(r"^（(\d+)）", text)
        return int(match.group(1)) if match else None
    return None


def validate_consecutive_heading_numbers(
    heading_entries: list[tuple[str, str]], scope: str, findings: Findings
) -> None:
    """Check H2 within H1, H3 within H2, and H4 within H3."""

    parent_for = {"h2": "h1", "h3": "h2", "h4": "h3"}
    for child_kind, parent_kind in parent_for.items():
        parent_positions = [
            index for index, (kind, _text) in enumerate(heading_entries) if kind == parent_kind
        ]
        for parent_offset, start in enumerate(parent_positions):
            end = parent_positions[parent_offset + 1] if parent_offset + 1 < len(parent_positions) else len(heading_entries)
            numbers = [
                heading_ordinal(kind, text)
                for kind, text in heading_entries[start + 1 : end]
                if kind == child_kind
            ]
            if not numbers:
                continue
            expected = list(range(1, len(numbers) + 1))
            if numbers != expected:
                parent_text = heading_entries[start][1]
                findings.error(
                    f"Nonconsecutive {child_kind.upper()} numbering under '{parent_text}' in {scope}: "
                    f"expected {expected}, got {numbers}"
                )


def iter_string_values(value: Any) -> Iterable[str]:
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for child in value.values():
            yield from iter_string_values(child)
    elif isinstance(value, list):
        for child in value:
            yield from iter_string_values(child)


def block_text(block: dict[str, Any]) -> str:
    block_type = str(block.get("type") or "").lower()
    if block_type == "table":
        values: list[str] = []
        values.extend(str(item) for item in block.get("header") or [])
        for row in block.get("rows") or []:
            values.extend(str(item) for item in row)
        return "\n".join(values)
    return str(block.get("text") or "")


def is_zero_project_statement(text: str, category: str) -> bool:
    """Return whether text explicitly states that the period has no category projects."""

    compact = re.sub(r"[\s\u3000，,：:]+", "", text)
    project_label = re.escape(category) if category.endswith("项目") else rf"{re.escape(category)}(?:类)?项目"
    return bool(
        re.search(
            rf"(?:本年度|报告期内|本报告期内|本期)(?:公司)?"
            rf"(?:无|没有|未有|不存在|未发生|未开展|未实施|未设立|未新设)(?:任何)?{project_label}",
            compact,
        )
    )


def is_only_zero_project_statement(text: str, category: str) -> bool:
    sentences = [item for item in re.split(r"[。！？；;\n]+", text) if item.strip()]
    return bool(sentences) and all(is_zero_project_statement(item, category) for item in sentences)


def validate_fixed_category_sections(
    main_blocks: list[dict[str, Any]],
    registry: list[dict[str, Any]],
    findings: Findings,
    *,
    fixed_sections: Iterable[tuple[str, str]] = FIXED_CATEGORY_SECTIONS,
) -> None:
    """Keep every fixed category section meaningful and consistent with its registry count."""

    fixed_sections = tuple(fixed_sections)

    def fixed_section_for_heading(text: str) -> tuple[str, str] | None:
        normalized = normalize_heading(text)
        for heading, category in fixed_sections:
            expected = normalize_heading(heading)
            ordinal = expected.split("）", 1)[0] + "）"
            if normalized == expected or (
                normalized.startswith(ordinal) and normalized.endswith(normalize_heading(category))
            ):
                return heading, category
        return None

    sections: dict[str, list[dict[str, Any]]] = {}
    active_category: str | None = None
    for block in main_blocks:
        block_type = str(block.get("type") or "").lower()
        if block_type in {"h1", "h2"}:
            fixed_section = fixed_section_for_heading(str(block.get("text") or ""))
            active_category = fixed_section[1] if block_type == "h2" and fixed_section else None
            if active_category is not None:
                sections.setdefault(active_category, [])
            continue
        if active_category is not None:
            sections[active_category].append(block)

    registry_counts = Counter(str(project.get("category") or "").strip() for project in registry)
    spv_slots = sum(1 for _heading, category in fixed_sections if category == "SPV项目")
    if spv_slots > 1 and registry_counts["SPV项目"] < spv_slots:
        findings.error(
            f"The fixed framework declares {spv_slots} SPV slots, but project_registry lists only "
            f"{registry_counts['SPV项目']} SPV项目 project(s); every named SPV slot needs its own registry row"
        )
    for heading, category in fixed_sections:
        if category not in sections:
            # The fixed-heading validator reports the missing or malformed heading.
            continue
        substantive_blocks = [
            block
            for block in sections[category]
            if str(block.get("type") or "").lower() in {"p", "h3"}
            and str(block.get("text") or "").strip()
        ]
        if not substantive_blocks:
            findings.error(f"Fixed category section {heading} must contain a non-empty p or h3 block")
            continue

        substantive_texts = [str(block.get("text") or "").strip() for block in substantive_blocks]
        zero_statements = [
            text for text in substantive_texts if is_zero_project_statement(text, category)
        ]
        project_count = registry_counts[category]
        project_label = category if category.endswith("项目") else f"{category}项目"
        if project_count == 0:
            meaningful_blocks = [
                block
                for block in sections[category]
                if str(block.get("type") or "").lower()
                in {"p", "h3", "h4", "caption", "tnote", "table"}
                and block_text(block).strip()
            ]
            if not zero_statements:
                findings.error(
                    f"Registry has no {category} projects; section {heading} must explicitly state 本年度无{project_label}"
                )
            if any(
                str(block.get("type") or "").lower() != "p"
                or not is_only_zero_project_statement(block_text(block), category)
                for block in meaningful_blocks
            ):
                findings.error(
                    f"Registry has no {category} projects; section {heading} may contain only "
                    f"explicit zero-project paragraph(s)"
                )
        elif zero_statements:
            findings.error(
                f"Registry has {project_count} {category} project(s), but section {heading} contains a zero-project declaration"
            )


def numeric_tokens(text: str) -> set[str]:
    tokens: set[str] = set()
    for match in NUMBER_TOKEN_PATTERN.finditer(text):
        raw = match.group(0).replace(",", "")
        try:
            value = Decimal(raw)
        except InvalidOperation:
            continue
        canonical = format(value.normalize(), "f")
        if canonical == "-0":
            canonical = "0"
        tokens.add(canonical)
    return tokens


def numeric_token_sequence(text: str, *, semantic_order: bool = False) -> list[str]:
    if semantic_order:
        text = re.sub(r"\d{4}[-年/]\d{1,2}[-月/]\d{1,2}日?", "", text)
    result: list[str] = []
    seen: set[str] = set()
    for match in NUMBER_TOKEN_PATTERN.finditer(text):
        raw = match.group(0).replace(",", "")
        try:
            canonical = format(Decimal(raw).normalize(), "f")
        except InvalidOperation:
            continue
        if canonical == "-0":
            canonical = "0"
        if semantic_order:
            try:
                as_integer = int(Decimal(canonical))
            except (InvalidOperation, ValueError):
                as_integer = -1
            if Decimal(canonical) == as_integer and 1900 <= as_integer <= 2099:
                continue
        if canonical not in seen:
            seen.add(canonical)
            result.append(canonical)
    return result


def validate_numeric_order_against_facts(
    payload: str,
    fact_ids: Iterable[str],
    *,
    scope: str,
    locator: str,
    fact_map: dict[str, dict[str, Any]],
    findings: Findings,
) -> None:
    block_sequence = numeric_token_sequence(payload, semantic_order=True)
    for fact_id in fact_ids:
        fact = fact_map.get(fact_id)
        if fact is None:
            continue
        fact_sequence = numeric_token_sequence(
            f"{fact.get('value') or ''}\n{fact.get('period') or ''}",
            semantic_order=True,
        )
        common = set(block_sequence) & set(fact_sequence)
        if len(common) < 2:
            continue
        actual = [token for token in block_sequence if token in common]
        expected = [token for token in fact_sequence if token in common]
        if actual != expected:
            findings.error(
                f"Numeric order in {scope} {locator} is inconsistent with {fact_id}: "
                f"expected {expected}, got {actual}"
            )


def matching_fact_assertions(
    payload: str,
    fact_ids: Iterable[str],
    fact_map: dict[str, dict[str, Any]],
) -> dict[str, list[str]]:
    """Return exact, whitespace-normalized ledger assertions present in payload."""

    normalized_payload = normalize_assertion_text(payload)
    matches: dict[str, list[str]] = {}
    for fact_id in fact_ids:
        fact = fact_map.get(fact_id)
        if fact is None:
            continue
        assertions = fact.get("assertions")
        if not isinstance(assertions, list):
            continue
        matches[fact_id] = [
            str(assertion)
            for assertion in assertions
            if normalize_assertion_text(str(assertion))
            and normalize_assertion_text(str(assertion)) in normalized_payload
        ]
    return matches


def report_sentences(payload: str) -> list[str]:
    """Split prose into evidence-bearing sentences while retaining comma clauses."""

    return [item.strip() for item in re.split(r"[。！？；;\r\n]+", payload) if item.strip()]


def status_claim_clauses(payload: str) -> list[str]:
    """Return comma-level clauses that make a status, direction, or outcome claim."""

    return [
        item.strip()
        for item in re.split(r"[，,。！？；;\r\n]+", payload)
        if item.strip() and STATUS_CLAIM_PATTERN.search(item)
    ]


def clause_has_assertion(
    clause: str,
    fact_ids: Iterable[str],
    fact_map: dict[str, dict[str, Any]],
) -> bool:
    """Accept an atomic clause only when it and a ledger assertion overlap exactly."""

    normalized_clause = normalize_assertion_text(clause)
    for fact_id in fact_ids:
        fact = fact_map.get(fact_id)
        if fact is None or not isinstance(fact.get("assertions"), list):
            continue
        for assertion in fact["assertions"]:
            normalized_assertion = normalize_assertion_text(str(assertion))
            if normalized_assertion and (
                normalized_assertion in normalized_clause
                or normalized_clause in normalized_assertion
            ):
                return True
    return False


def is_structural_sentence(sentence: str) -> bool:
    """Return True only for a pure cross-reference or unit label."""

    compact = normalize_assertion_text(sentence)
    if re.fullmatch(r"[^，,；;。]{0,40}见附件[一二三四五六七八九十\d]+", compact):
        return True
    if re.fullmatch(r"单位[:：].+", compact):
        return True
    return False


def is_context_clause(clause: str) -> bool:
    """Return True for a standalone time/transition fragment with no claim."""

    compact = normalize_assertion_text(clause)
    return bool(
        re.fullmatch(
            r"(?:报告期内|本报告期内|本期|截至报告期末|截至期末|下一步|其中|同时|此外)",
            compact,
        )
    )


def validate_block_fact_references(
    blocks: Iterable[dict[str, Any]],
    *,
    scope: str,
    fact_map: dict[str, dict[str, Any]],
    findings: Findings,
    template_mode: bool,
) -> set[str]:
    """Validate block-to-ledger links and return the referenced fact IDs."""

    referenced: set[str] = set()
    allowed_destinations = {"both", "main body" if scope == "main body" else "attachment"}
    for index, block in enumerate(blocks, start=1):
        block_type = str(block.get("type") or "").lower()
        refs_value = block.get("fact_ids")
        refs: list[str] = []
        if refs_value is not None:
            if not isinstance(refs_value, list):
                findings.error(f"{scope} block {index} fact_ids must be a list")
            else:
                refs = [str(item or "").strip() for item in refs_value]
                if "" in refs:
                    findings.error(f"{scope} block {index} contains an empty fact_id reference")
                if len(refs) != len(set(refs)):
                    findings.error(f"{scope} block {index} contains duplicate fact_id references")

        exempt_reason = str(block.get("evidence_exempt_reason") or "").strip().lower()
        if exempt_reason:
            if refs:
                findings.error(f"{scope} block {index} cannot have both fact_ids and evidence_exempt_reason")
            if block_type != "tnote" or exempt_reason != "unit-label":
                findings.error(
                    f"Unsupported evidence_exempt_reason in {scope} block {index}: {exempt_reason}"
                )
        elif block_type in {"p", "tnote", "table"} and not refs:
            findings.error(f"Substantive {block_type} in {scope} block {index} is missing fact_ids")

        payload = block_text(block)
        has_material_number = bool(MATERIAL_NUMBER_PATTERN.search(payload))
        if block_type == "table" and BARE_NUMBER_PATTERN.search(payload):
            has_material_number = True
        if block_type in {"p", "tnote", "table"} and has_material_number and not refs:
            findings.error(f"Material numeric content in {scope} block {index} is missing fact_ids")

        for fact_id in refs:
            if not fact_id:
                continue
            referenced.add(fact_id)
            fact = fact_map.get(fact_id)
            if fact is None:
                findings.error(f"{scope} block {index} references unknown fact_id {fact_id}")
                continue
            destination = str(fact.get("destination") or "").strip().lower()
            if destination not in allowed_destinations:
                findings.error(
                    f"{scope} block {index} references {fact_id}, whose destination is {destination or '<empty>'}"
                )
            status = str(fact.get("status") or "").strip().lower()
            if not template_mode and status in {"conflicting", "stale", "missing"}:
                findings.error(f"Production block references unresolved {status} fact {fact_id}")

        assertion_matches = matching_fact_assertions(payload, refs, fact_map)
        for fact_id in refs:
            if fact_id in fact_map and not assertion_matches.get(fact_id):
                findings.error(
                    f"{scope} block {index} contains no exact ledger assertion for referenced {fact_id}"
                )

        if block_type in {"p", "tnote"} and refs:
            for sentence_number, sentence in enumerate(report_sentences(payload), start=1):
                if is_structural_sentence(sentence):
                    continue
                if template_mode and TEMPLATE_MARKER_PATTERN.search(sentence):
                    continue
                sentence_matches = matching_fact_assertions(sentence, refs, fact_map)
                if not any(sentence_matches.values()):
                    findings.error(
                        f"Unsupported sentence in {scope} block {index}, sentence {sentence_number}: "
                        f"{sentence[:100]}"
                    )
            for clause_number, clause in enumerate(status_claim_clauses(payload), start=1):
                if template_mode and TEMPLATE_MARKER_PATTERN.search(clause):
                    continue
                if not clause_has_assertion(clause, refs, fact_map):
                    findings.error(
                        f"Unsupported status clause in {scope} block {index}, clause {clause_number}: "
                        f"{clause[:100]}"
                    )
            if not (template_mode and TEMPLATE_MARKER_PATTERN.search(payload)):
                factual_clauses = [
                    item.strip()
                    for item in re.split(r"[，,。！？；;\r\n]+", payload)
                    if item.strip()
                ]
                for clause_number, clause in enumerate(factual_clauses, start=1):
                    if is_structural_sentence(clause) or is_context_clause(clause):
                        continue
                    if not clause_has_assertion(clause, refs, fact_map):
                        findings.error(
                            f"Unsupported factual clause in {scope} block {index}, "
                            f"clause {clause_number}: {clause[:100]}"
                        )

        # Attachment references are structural locators, not report facts.
        numeric_payload = re.sub(r"附件\s*[一二三四五六七八九十\d]+", "附件", payload)
        block_numbers = numeric_tokens(numeric_payload)
        if block_numbers and refs:
            matched_assertion_numbers: set[str] = set()
            for assertions in assertion_matches.values():
                for assertion in assertions:
                    matched_assertion_numbers.update(numeric_tokens(assertion))
            uncovered = sorted(
                block_numbers - matched_assertion_numbers,
                key=lambda item: (Decimal(item), item),
            )
            if uncovered:
                findings.error(
                    f"Numeric tokens in {scope} block {index} are not covered by matched assertions: "
                    + ", ".join(uncovered[:20])
                )
        if block_type in {"p", "tnote"} and refs:
            validate_numeric_order_against_facts(
                numeric_payload,
                refs,
                scope=scope,
                locator=f"block {index}",
                fact_map=fact_map,
                findings=findings,
            )
        if block_type == "table" and BARE_NUMBER_PATTERN.search(payload):
            rows = block.get("rows") or []
            row_fact_ids = block.get("row_fact_ids")
            if not isinstance(row_fact_ids, list) or len(row_fact_ids) != len(rows):
                findings.error(
                    f"Numeric table in {scope} block {index} must provide row_fact_ids for every row"
                )
            else:
                for row_index, (row, row_refs_value) in enumerate(zip(rows, row_fact_ids), start=1):
                    row_payload = "\n".join(str(item) for item in row)
                    row_numbers = numeric_tokens(row_payload)
                    if not row_numbers:
                        continue
                    if not isinstance(row_refs_value, list) or not row_refs_value:
                        findings.error(
                            f"Numeric table row {row_index} in {scope} block {index} has no row_fact_ids"
                        )
                        continue
                    row_refs = [str(item or "").strip() for item in row_refs_value]
                    if "" in row_refs or len(row_refs) != len(set(row_refs)):
                        findings.error(
                            f"Invalid row_fact_ids in {scope} block {index}, row {row_index}"
                        )
                    if any(fact_id not in refs for fact_id in row_refs):
                        findings.error(
                            f"row_fact_ids in {scope} block {index}, row {row_index} must be included in fact_ids"
                        )
                    row_assertion_matches = matching_fact_assertions(
                        row_payload, row_refs, fact_map
                    )
                    row_evidence_numbers: set[str] = set()
                    for fact_id in row_refs:
                        if fact_id not in fact_map:
                            findings.error(
                                f"{scope} block {index}, row {row_index} references unknown fact_id {fact_id}"
                            )
                            continue
                        if not row_assertion_matches.get(fact_id):
                            findings.error(
                                f"{scope} block {index}, row {row_index} contains no exact ledger "
                                f"assertion for referenced {fact_id}"
                            )
                        for assertion in row_assertion_matches.get(fact_id, []):
                            row_evidence_numbers.update(numeric_tokens(assertion))
                    uncovered_row = sorted(
                        row_numbers - row_evidence_numbers,
                        key=lambda item: (Decimal(item), item),
                    )
                    if uncovered_row:
                        findings.error(
                            f"Numeric tokens in {scope} block {index}, row {row_index} are not covered: "
                            + ", ".join(uncovered_row)
                        )
                    validate_numeric_order_against_facts(
                        row_payload,
                        row_refs,
                        scope=scope,
                        locator=f"block {index}, row {row_index}",
                        fact_map=fact_map,
                        findings=findings,
                    )
    return referenced


def xml_text(payload: bytes) -> str:
    try:
        root = ET.fromstring(payload)
    except ET.ParseError:
        return payload.decode("utf-8", errors="ignore")
    return "".join(root.itertext())


def validate_spec_container_types(spec: dict[str, Any], findings: Findings) -> bool:
    valid = True
    if "layout" in spec and not isinstance(spec.get("layout"), dict):
        findings.error("layout must be an object")
        valid = False
    if not isinstance(spec.get("document"), dict):
        findings.error("document must be an object")
        valid = False
    for key in ("main_blocks", "attachments", "project_registry", "sources", "fact_ledger"):
        value = spec.get(key)
        if not isinstance(value, list):
            findings.error(f"{key} must be a list")
            valid = False
            continue
        for index, item in enumerate(value, start=1):
            if not isinstance(item, dict):
                findings.error(f"{key}[{index}] must be an object")
                valid = False
    if not valid:
        return False

    block_groups: list[tuple[str, list[Any]]] = [("main_blocks", spec["main_blocks"])]
    for index, attachment in enumerate(spec["attachments"], start=1):
        blocks = attachment.get("blocks")
        if not isinstance(blocks, list):
            findings.error(f"attachments[{index}].blocks must be a list")
            valid = False
            continue
        block_groups.append((f"attachments[{index}].blocks", blocks))
    for group_name, blocks in block_groups:
        for index, block in enumerate(blocks, start=1):
            if not isinstance(block, dict):
                findings.error(f"{group_name}[{index}] must be an object")
                valid = False
                continue
            if "fact_ids" in block and not isinstance(block.get("fact_ids"), list):
                findings.error(f"{group_name}[{index}].fact_ids must be a list")
                valid = False
            block_type = str(block.get("type") or "").lower()
            if "bold" in block and not isinstance(block.get("bold"), bool):
                findings.error(f"{group_name}[{index}].bold must be true or false")
                valid = False
            if "align" in block:
                align = block.get("align")
                valid_alignments = {"left", "right", "center", "justify"}
                if isinstance(align, str):
                    if align not in valid_alignments:
                        findings.error(f"{group_name}[{index}].align is unsupported: {align}")
                        valid = False
                elif block_type != "table" or not isinstance(align, list) or any(
                    not isinstance(item, str) or item not in valid_alignments for item in align
                ):
                    findings.error(
                        f"{group_name}[{index}].align must be a supported string"
                        + (" or per-column list" if block_type == "table" else "")
                    )
                    valid = False
            if block_type in {"h1", "h2", "h3", "h4"} and MATERIAL_NUMBER_PATTERN.search(
                str(block.get("text") or "")
            ):
                findings.error(f"{group_name}[{index}] heading cannot contain a material numeric claim")
                valid = False
            if block_type == "table":
                header = block.get("header")
                if not isinstance(header, list) or not header:
                    findings.error(f"{group_name}[{index}].header must be a list")
                    valid = False
                rows = block.get("rows")
                if not isinstance(rows, list):
                    findings.error(f"{group_name}[{index}].rows must be a list")
                    valid = False
                elif any(not isinstance(row, list) for row in rows):
                    findings.error(f"{group_name}[{index}] table rows must all be lists")
                    valid = False
                elif isinstance(header, list) and any(len(row) != len(header) for row in rows):
                    findings.error(
                        f"{group_name}[{index}] table rows must match the header column count"
                    )
                    valid = False
                widths = block.get("widths")
                if widths is not None:
                    if not isinstance(widths, list):
                        findings.error(f"{group_name}[{index}].widths must be a list")
                        valid = False
                    else:
                        if isinstance(header, list) and len(widths) != len(header):
                            findings.error(
                                f"{group_name}[{index}].widths must match the header column count"
                            )
                            valid = False
                        for width in widths:
                            if (
                                isinstance(width, bool)
                                or not isinstance(width, (int, float))
                                or not math.isfinite(float(width))
                                or float(width) <= 0
                            ):
                                findings.error(
                                    f"{group_name}[{index}].widths must contain positive finite numbers"
                                )
                                valid = False
                                break
                if "row_fact_ids" in block and not isinstance(block.get("row_fact_ids"), list):
                    findings.error(f"{group_name}[{index}].row_fact_ids must be a list")
                    valid = False
    return valid


def validate_spec(spec: dict[str, Any], findings: Findings, *, template_mode: bool) -> None:
    if not validate_spec_container_types(spec, findings):
        return
    if str(spec.get("schema_version") or "") != "1.0":
        findings.error("schema_version must be 1.0")
    if template_mode and spec.get("template_only") is not True:
        findings.error("--template-mode requires template_only=true in the specification")

    layout = spec.get("layout") or {}
    for key in (
        "page_width_cm",
        "page_height_cm",
        "top_margin_cm",
        "bottom_margin_cm",
        "left_margin_cm",
        "right_margin_cm",
        "header_distance_cm",
        "footer_distance_cm",
    ):
        try:
            value = float(layout.get(key))
        except (TypeError, ValueError):
            findings.error(f"layout.{key} must be a positive number")
            continue
        if value <= 0:
            findings.error(f"layout.{key} must be a positive number")

    document = spec.get("document") or {}
    for key in (
        "company",
        "report_year",
        "cutoff_date",
        "document_number",
        "recipient",
        "legal_basis",
        "issuer",
        "issue_date",
        "heading_contract_source",
        "heading_change_note",
    ):
        if not str(document.get(key) or "").strip():
            findings.error(f"Missing document.{key}")
    validate_fixed_heading_contract(document, findings)
    report_year = str(document.get("report_year") or "")
    if not re.fullmatch(r"20\d{2}", report_year):
        findings.error("document.report_year must be a four-digit year")
    validate_document_dates(document, findings)
    validate_imprint_metadata(document, findings)
    validate_official_style(spec, findings)

    main_blocks = spec.get("main_blocks") or []
    main_headings = block_headings(main_blocks)
    validate_heading_scope(main_headings, "main body", findings)
    main_heading_entries = block_heading_entries(main_blocks)
    required_main_blocks = required_main_blocks_from_document(document)
    validate_required_heading_sequence(
        main_heading_entries,
        "specification main body",
        findings,
        required_blocks=required_main_blocks,
    )
    validate_consecutive_heading_numbers(main_heading_entries, "specification main body", findings)

    for index, block in enumerate(main_blocks, start=1):
        block_type = str(block.get("type") or "").lower()
        if block_type not in BLOCK_TYPES and not block_type.startswith("_"):
            findings.error(f"Unsupported main block type at index {index}: {block_type}")

    attachments = spec.get("attachments") or []
    attachment_ids = [str(item.get("id") or "") for item in attachments]
    if "" in attachment_ids:
        findings.error("Every attachment must have a non-empty id")
    if len(attachment_ids) != len(set(attachment_ids)):
        findings.error("Attachment ids must be unique")
    try:
        attachment_numbers = [int(item.get("number") or 0) for item in attachments]
    except (TypeError, ValueError):
        attachment_numbers = []
        findings.error("Every attachment number must be an integer")
    expected_numbers = list(range(1, len(attachments) + 1))
    if attachment_numbers != expected_numbers:
        findings.error(f"Attachment numbers must be consecutive and ordered: expected {expected_numbers}, got {attachment_numbers}")
    titles = [normalize_heading(str(item.get("title") or "")) for item in attachments]
    if "" in titles:
        findings.error("Every attachment must have a title")
    if len(titles) != len(set(titles)):
        findings.error("Attachment titles must be unique")
    for attachment in attachments:
        attachment_scope = f"attachment {attachment.get('id')}"
        attachment_blocks = attachment.get("blocks") or []
        validate_heading_scope(block_headings(attachment_blocks), attachment_scope, findings)
        for index, block in enumerate(attachment_blocks, start=1):
            block_type = str(block.get("type") or "").lower()
            if block_type not in BLOCK_TYPES and not block_type.startswith("_"):
                findings.error(f"Unsupported {attachment_scope} block type at index {index}: {block_type}")

    registry = spec.get("project_registry") or []
    project_ids = [str(item.get("project_id") or "") for item in registry]
    project_names = [normalize_heading(str(item.get("official_name") or "")) for item in registry]
    if "" in project_ids or "" in project_names:
        findings.error("Every registry row must have project_id and official_name")
    if len(project_ids) != len(set(project_ids)):
        findings.error("project_registry.project_id values must be unique")
    if len(project_names) != len(set(project_names)):
        findings.error("project_registry.official_name values must be unique after normalization")
    valid_attachment_ids = set(attachment_ids)
    for project in registry:
        project_id = str(project.get("project_id") or "").strip() or "<unknown>"
        for key in (
            "official_name",
            "category",
            "status",
            "investment_entity",
            "commitment",
            "paid_in",
            "interest",
            "reporting_period",
            "attachment_id",
            "change_flag",
        ):
            if not str(project.get(key) or "").strip():
                findings.error(f"Project {project_id} is missing {key}")
        category = str(project.get("category") or "").strip()
        if category and category not in VALID_PROJECT_CATEGORIES:
            findings.error(
                f"Project {project_id} has unsupported category {category}; expected one of "
                + ", ".join(sorted(VALID_PROJECT_CATEGORIES))
            )
        attachment_id = str(project.get("attachment_id") or "")
        if attachment_id and attachment_id not in valid_attachment_ids:
            findings.error(f"Project {project.get('project_id')} references missing attachment {attachment_id}")

    validate_fixed_category_sections(
        main_blocks,
        registry,
        findings,
        fixed_sections=fixed_category_sections_from_document(document),
    )

    sources = spec.get("sources") or []
    valid_project_ids = set(project_ids) | {"portfolio"}
    source_id_list = [str(source.get("source_id") or "").strip() for source in sources]
    if "" in source_id_list:
        findings.error("Every source must have a non-empty source_id")
    duplicate_source_ids = sorted(key for key, count in Counter(source_id_list).items() if key and count > 1)
    for source_id in duplicate_source_ids:
        findings.error(f"Duplicate source_id: {source_id}")
    source_ids = set(source_id_list)
    source_project_map: dict[str, set[str]] = {}
    for source in sources:
        source_id = str(source.get("source_id") or "").strip() or "<unknown>"
        for key in ("filename", "period", "document_date", "approval_status", "authority_level", "status"):
            if not str(source.get(key) or "").strip():
                findings.error(f"Source {source_id} is missing {key}")
        source_project_ids = source.get("project_ids")
        if not isinstance(source_project_ids, list) or not source_project_ids:
            findings.error(f"Source {source_id} project_ids must be a non-empty list")
        else:
            normalized_source_projects = [str(item or "").strip() for item in source_project_ids]
            source_project_map[source_id] = set(normalized_source_projects)
            if "" in normalized_source_projects:
                findings.error(f"Source {source_id} contains an empty project_id")
            if len(normalized_source_projects) != len(set(normalized_source_projects)):
                findings.error(f"Source {source_id} contains duplicate project_ids")
            for project_id in normalized_source_projects:
                if project_id and project_id not in valid_project_ids:
                    findings.error(f"Source {source_id} references unknown project_id {project_id}")
        if not template_mode and str(source.get("status") or "").strip().lower() == "synthetic":
            findings.error(f"Production specification uses synthetic source {source_id}")

    fact_map: dict[str, dict[str, Any]] = {}
    for fact in spec.get("fact_ledger") or []:
        fact_id = str(fact.get("fact_id") or "").strip()
        if not fact_id:
            findings.error("Fact ledger row is missing fact_id")
        elif fact_id in fact_map:
            findings.error(f"Duplicate fact_id: {fact_id}")
        else:
            fact_map[fact_id] = fact
        for key in ("project_id", "metric", "value", "unit", "period", "scope", "source_id", "locator", "status", "destination"):
            if str(fact.get(key) or "").strip() == "":
                findings.error(f"Fact {fact_id or '<unknown>'} is missing {key}")
        assertions = fact.get("assertions")
        if not isinstance(assertions, list) or not assertions:
            findings.error(f"Fact {fact_id or '<unknown>'} assertions must be a non-empty list")
        else:
            normalized_assertions = [
                re.sub(r"[\s\u3000]+", "", str(assertion or "")) for assertion in assertions
            ]
            if "" in normalized_assertions:
                findings.error(f"Fact {fact_id or '<unknown>'} contains an empty assertion")
            if len(normalized_assertions) != len(set(normalized_assertions)):
                findings.error(f"Fact {fact_id or '<unknown>'} contains duplicate assertions")
            evidence_text = "\n".join(
                str(fact.get(key) or "")
                for key in ("metric", "value", "unit", "period", "scope", "locator", "formula")
            )
            evidence_numbers = numeric_tokens(evidence_text)
            for assertion in assertions:
                uncovered_assertion_numbers = numeric_tokens(str(assertion)) - evidence_numbers
                if uncovered_assertion_numbers:
                    findings.error(
                        f"Fact {fact_id or '<unknown>'} assertion contains numbers absent from its evidence: "
                        + ", ".join(sorted(uncovered_assertion_numbers))
                    )
        project_id = str(fact.get("project_id") or "").strip()
        if project_id and project_id not in valid_project_ids:
            findings.error(f"Fact {fact_id} references unknown project_id {project_id}")
        source_id = str(fact.get("source_id") or "")
        if source_id and source_id not in source_ids:
            findings.error(f"Fact {fact_id} references unknown source_id {source_id}")
        elif source_id and project_id and project_id not in source_project_map.get(source_id, set()):
            findings.error(
                f"Fact {fact_id} project_id {project_id} is outside source {source_id} project_ids"
            )
        status = str(fact.get("status") or "").strip().lower()
        if status and status not in VALID_FACT_STATUSES:
            findings.error(f"Fact {fact_id} has unsupported status {status}")
        if status == "calculated" and not str(fact.get("formula") or "").strip():
            findings.error(f"Calculated fact {fact_id} is missing formula")
        destination = str(fact.get("destination") or "").strip().lower()
        if destination and destination not in VALID_FACT_DESTINATIONS:
            findings.error(f"Fact {fact_id} has unsupported destination {destination}")
        if not template_mode and status in {"conflicting", "stale", "missing"} and destination != "excluded":
            findings.error(f"Production specification contains unresolved {status} fact {fact_id}")

    main_referenced_facts = validate_block_fact_references(
        main_blocks,
        scope="main body",
        fact_map=fact_map,
        findings=findings,
        template_mode=template_mode,
    )
    attachment_referenced_facts: set[str] = set()
    for attachment in attachments:
        attachment_referenced_facts.update(
            validate_block_fact_references(
                attachment.get("blocks") or [],
                scope=f"attachment {attachment.get('id')}",
                fact_map=fact_map,
                findings=findings,
                template_mode=template_mode,
            )
        )
    referenced_facts = main_referenced_facts | attachment_referenced_facts
    for fact_id, fact in fact_map.items():
        destination = str(fact.get("destination") or "").strip().lower()
        if destination == "main body" and fact_id not in main_referenced_facts:
            findings.error(f"Fact {fact_id} is destined for the main body but is not referenced there")
        elif destination == "attachment" and fact_id not in attachment_referenced_facts:
            findings.error(f"Fact {fact_id} is destined for an attachment but is not referenced there")
        elif destination == "both":
            if fact_id not in main_referenced_facts:
                findings.error(f"Fact {fact_id} is destined for both scopes but is not referenced in the main body")
            if fact_id not in attachment_referenced_facts:
                findings.error(f"Fact {fact_id} is destined for both scopes but is not referenced in an attachment")
        if not template_mode and destination == "pending user decision":
            findings.error(f"Production specification contains pending user decision fact {fact_id}")

    spec_text = json.dumps(spec, ensure_ascii=False)
    if not template_mode:
        if spec.get("template_only") is True:
            findings.error("template_only specifications cannot be validated as production reports")
        placeholders = sorted(set(find_placeholders(spec_text)))
        if placeholders:
            findings.error(f"Unresolved placeholders in specification: {', '.join(placeholders[:10])}")
        template_markers = sorted(
            {
                match.group(0)
                for value in iter_string_values(spec)
                for match in TEMPLATE_MARKER_PATTERN.finditer(value)
            }
        )
        if template_markers:
            findings.error(f"Template-only markers in production specification: {', '.join(template_markers[:10])}")
    findings.note(
        f"Specification: {len(registry)} projects, {len(attachments)} attachments, "
        f"{len(spec.get('fact_ledger') or [])} fact rows, {len(referenced_facts)} referenced fact rows"
    )


def iter_docx_text(doc: Document) -> Iterable[str]:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            yield paragraph.text.strip()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    yield text


def normalized_item(kind: str, values: Iterable[str]) -> str:
    normalized_values = [re.sub(r"[\s\u3000]+", "", str(value)) for value in values]
    return f"{kind}:" + "\u241f".join(normalized_values)


def iter_ordered_docx_items(doc: Document) -> Iterable[tuple[str, str]]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = paragraph.text.strip()
            if text:
                yield "p", text
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            cells = [cell.text for row in table.rows for cell in row.cells]
            yield "table", "\u241f".join(cells)


def expected_block_item(block: dict[str, Any]) -> str | None:
    block_type = str(block.get("type") or "").lower()
    if block_type == "table":
        cells = [str(item) for item in block.get("header") or []]
        cells.extend(str(item) for row in block.get("rows") or [] for item in row)
        return normalized_item("table", cells)
    if block_type in {"h1", "h2", "h3", "h4", "p", "caption", "tnote"}:
        text = str(block.get("text") or "").strip()
        return normalized_item("p", [text]) if text else None
    return None


def docx_scoped_items(
    doc: Document, expected_spec: dict[str, Any]
) -> tuple[dict[str, list[str]], list[str]]:
    number_to_scope = {
        int(attachment.get("number") or 0): f"attachment {attachment.get('id')}"
        for attachment in expected_spec.get("attachments") or []
    }
    scopes: dict[str, list[str]] = {"main body": []}
    current_scope = "main body"
    unexpected_labels: list[str] = []
    items = list(iter_ordered_docx_items(doc))
    # 版记（印发机关和印发日期）sits on the last page, after the final attachment.
    # It belongs to no attachment scope, so peel it off before scope assignment.
    document = expected_spec.get("document") or {}
    printer = re.sub(r"[\s　]+", "", str(document.get("printer") or ""))
    print_date = re.sub(r"[\s　]+", "", str(document.get("print_date") or ""))
    if printer and print_date and items and items[-1][0] == "table":
        tail = re.sub(r"[\s　]+", "", items[-1][1])
        if tail == f"{printer}␟{print_date}印发":
            items = items[:-1]
    for kind, value in items:
        if kind == "p":
            compact = re.sub(r"[\s\u3000]+", "", value)
            match = re.fullmatch(r"附件(\d+)", compact)
            if match:
                number = int(match.group(1))
                next_scope = number_to_scope.get(number)
                if next_scope is None:
                    unexpected_labels.append(value)
                    current_scope = f"unexpected attachment {number}"
                else:
                    current_scope = next_scope
                scopes.setdefault(current_scope, [])
        if kind == "table":
            scopes.setdefault(current_scope, []).append(
                "table:" + re.sub(r"[\s\u3000]+", "", value)
            )
        else:
            scopes.setdefault(current_scope, []).append(normalized_item("p", [value]))
    return scopes, unexpected_labels


def validate_scoped_docx_content(
    doc: Document,
    expected_spec: dict[str, Any],
    findings: Findings,
) -> None:
    scopes, unexpected_labels = docx_scoped_items(doc, expected_spec)
    for label in unexpected_labels:
        findings.error(f"DOCX contains an unexpected standalone attachment label: {label}")

    expected_scopes: dict[str, list[str]] = {
        "main body": [
            item
            for block in expected_spec.get("main_blocks") or []
            if (item := expected_block_item(block)) is not None
        ]
    }
    for attachment in expected_spec.get("attachments") or []:
        scope = f"attachment {attachment.get('id')}"
        expected_items = [normalized_item("p", [str(attachment.get("title") or "")])]
        expected_items.extend(
            item
            for block in attachment.get("blocks") or []
            if (item := expected_block_item(block)) is not None
        )
        expected_scopes[scope] = expected_items

    for scope, expected_items in expected_scopes.items():
        actual_items = scopes.get(scope)
        if actual_items is None:
            findings.error(f"DOCX is missing scope: {scope}")
            continue
        expected_counts = Counter(expected_items)
        actual_counts = Counter(actual_items)
        for item, expected_count in expected_counts.items():
            actual_count = actual_counts.get(item, 0)
            if actual_count != expected_count:
                readable = item.split(":", 1)[-1].replace("\u241f", " | ")
                findings.error(
                    f"DOCX {scope} content count mismatch (expected {expected_count}, found {actual_count}): "
                    f"{readable[:120]}"
                )
        cursor = 0
        for item in expected_items:
            try:
                cursor = actual_items.index(item, cursor) + 1
            except ValueError:
                readable = item.split(":", 1)[-1].replace("\u241f", " | ")
                findings.error(f"DOCX {scope} content is missing or out of order: {readable[:120]}")
                break

        # Count/order checks alone permit a contradictory extra paragraph to be
        # inserted beside the correct paragraph.  Lock the substantive span to
        # the exact spec sequence.  Main-body envelope text (red head, recipient,
        # opening, attachment list, signature) is intentionally outside the span.
        if expected_items:
            try:
                span_start = actual_items.index(expected_items[0])
                span_end = len(actual_items) - 1 - actual_items[::-1].index(expected_items[-1])
            except ValueError:
                continue
            substantive_span = actual_items[span_start : span_end + 1]
            if substantive_span != expected_items:
                findings.error(
                    f"DOCX {scope} contains unexpected, stale, or contradictory content "
                    "inside the specification-controlled span"
                )

        if scope.startswith("attachment "):
            attachment_id = scope.split(" ", 1)[1]
            attachment = next(
                (
                    item
                    for item in expected_spec.get("attachments") or []
                    if str(item.get("id") or "") == attachment_id
                ),
                None,
            )
            if attachment is None:
                continue
            label_item = normalized_item("p", [f"附件{attachment.get('number')}"])
            controlled_items = list(actual_items)
            if controlled_items and controlled_items[0] == label_item:
                controlled_items = controlled_items[1:]
            if controlled_items != expected_items:
                findings.error(
                    f"DOCX {scope} contains content outside its exact attachment specification"
                )

    main_actual = scopes.get("main body") or []
    main_expected = expected_scopes.get("main body") or []
    if main_actual and main_expected:
        document = expected_spec.get("document") or {}
        company = str(document.get("company") or "").strip()
        report_year = str(document.get("report_year") or "").strip()
        document_number = str(document.get("document_number") or "").strip()
        signer = str(document.get("signer") or "").strip()
        recipient = str(document.get("recipient") or "").strip()
        legal_basis = str(document.get("legal_basis") or "").strip()
        expected_prefix = [
            normalized_item("p", [f"{company}文件"]),
            normalized_item("table", [document_number, signer]),
            normalized_item(
                "p",
                [f"{company}关于{report_year}年度股权投资项目投后情况报告"],
            ),
        ]
        expected_prefix.extend(
            normalized_item("p", [value])
            for value in (recipient, legal_basis)
            if value
        )
        expected_prefix.append(main_expected[0])
        if main_actual[: len(expected_prefix)] != expected_prefix:
            findings.error(
                "DOCX main-body opening envelope must exactly preserve the red head, document row, "
                "two-line title, recipient, opening, and first fixed heading order"
            )

        suffix_items: list[str] = []
        attachments = expected_spec.get("attachments") or []
        if attachments:
            first = attachments[0]
            suffix_items.append(
                normalized_item("p", [f"附件：1.{first.get('title', '')}"])
            )
            for number, attachment in enumerate(attachments[1:], start=2):
                suffix_items.append(
                    normalized_item("p", [f"{number}.{attachment.get('title', '')}"])
                )
        issuer = str(document.get("issuer") or document.get("company") or "").strip()
        issue_date = str(document.get("issue_date") or "").strip()
        if issuer:
            suffix_items.append(normalized_item("p", [issuer]))
        if issue_date:
            suffix_items.append(normalized_item("p", [issue_date]))
        contact_parts: list[str] = []
        contact_name = str(document.get("contact_name") or "").strip()
        contact_phone = str(document.get("contact_phone") or "").strip()
        if contact_name:
            contact_parts.append(f"联系人：{contact_name}")
        if contact_phone:
            contact_parts.append(f"联系电话：{contact_phone}")
        if contact_parts:
            suffix_items.append(normalized_item("p", [f"（{'  '.join(contact_parts)}）"]))

        try:
            final_main_index = len(main_actual) - 1 - main_actual[::-1].index(main_expected[-1])
        except ValueError:
            final_main_index = -1
        if final_main_index >= 0 and main_actual[final_main_index + 1 :] != suffix_items:
            findings.error(
                "DOCX main-body closing boundary contains unexpected content between the "
                "final specification block, attachment list, and signature"
            )


def paragraph_heading_kind(paragraph: Any) -> str:
    for properties in (
        paragraph._p.pPr,
        getattr(getattr(paragraph, "style", None), "element", None).pPr
        if getattr(getattr(paragraph, "style", None), "element", None) is not None
        else None,
    ):
        outline = properties.find(qn("w:outlineLvl")) if properties is not None else None
        if outline is not None:
            try:
                level = int(outline.get(qn("w:val")))
            except (TypeError, ValueError):
                level = -1
            if 0 <= level <= 3:
                return f"h{level + 1}"

    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
    style_match = re.match(r"^(?:Heading|标题)\s*([1-4])$", style_name, re.IGNORECASE)
    if style_match:
        return f"h{style_match.group(1)}"

    run = next((item for item in paragraph.runs if item.text.strip()), None)
    if run is None:
        return "unknown"
    properties = run._element.rPr
    fonts = properties.find(qn("w:rFonts")) if properties is not None else None
    if fonts is None:
        style_rpr = getattr(getattr(paragraph, "style", None), "element", None)
        style_rpr = style_rpr.rPr if style_rpr is not None else None
        fonts = style_rpr.find(qn("w:rFonts")) if style_rpr is not None else None
    east_asia = fonts.get(qn("w:eastAsia")) if fonts is not None else ""
    bold_element = properties.find(qn("w:b")) if properties is not None else None
    bold_raw = bold_element.get(qn("w:val")) if bold_element is not None else None
    is_bold = run.bold is True or (
        run.bold is None
        and bold_element is not None
        and str(bold_raw or "1").strip().lower() not in {"0", "false", "off", "no"}
    )
    if east_asia == "黑体" and not is_bold:
        return "h1"
    if east_asia in {"楷体", "楷体_GB2312"} and is_bold:
        return "h2"
    if east_asia in {"仿宋", "仿宋_GB2312"} and is_bold:
        return "h3"
    if east_asia in {"仿宋", "仿宋_GB2312"} and not is_bold:
        return "h4"
    return "unknown"


def _format_property_candidates(doc: Document, paragraph: Paragraph, run: Run) -> Iterable[Any]:
    run_properties = run._element.rPr
    if run_properties is not None:
        yield run_properties
    run_style = getattr(run, "style", None)
    if run_style is not None and getattr(run_style, "element", None) is not None:
        if run_style.element.rPr is not None:
            yield run_style.element.rPr
    paragraph_properties = paragraph._p.pPr
    if paragraph_properties is not None:
        paragraph_mark_properties = paragraph_properties.find(qn("w:rPr"))
        if paragraph_mark_properties is not None:
            yield paragraph_mark_properties
    paragraph_style = getattr(paragraph, "style", None)
    if paragraph_style is not None and getattr(paragraph_style, "element", None) is not None:
        if paragraph_style.element.rPr is not None:
            yield paragraph_style.element.rPr
    normal = doc.styles["Normal"]
    if normal.element.rPr is not None:
        yield normal.element.rPr


def _effective_run_format(
    doc: Document, paragraph: Paragraph, run: Run
) -> tuple[str, str, float | None, bool, str | None]:
    east_asia_font = ""
    western_font = ""
    size: float | None = None
    bold: bool | None = None
    color: str | None = None
    for properties in _format_property_candidates(doc, paragraph, run):
        if not east_asia_font or not western_font:
            fonts = properties.find(qn("w:rFonts"))
            if fonts is not None:
                if not east_asia_font:
                    east_asia_font = str(
                        fonts.get(qn("w:eastAsia"))
                        or ""
                    )
                if not western_font:
                    western_font = str(
                        fonts.get(qn("w:ascii"))
                        or fonts.get(qn("w:hAnsi"))
                        or ""
                    )
        if size is None:
            size_element = properties.find(qn("w:sz"))
            if size_element is not None:
                try:
                    size = int(size_element.get(qn("w:val"))) / 2
                except (TypeError, ValueError):
                    pass
        if bold is None:
            bold_element = properties.find(qn("w:b"))
            if bold_element is not None:
                bold = str(bold_element.get(qn("w:val")) or "1").lower() not in {
                    "0",
                    "false",
                    "off",
                    "no",
                }
        if color is None:
            color_element = properties.find(qn("w:color"))
            if color_element is not None:
                color = str(color_element.get(qn("w:val")) or "").upper() or None
    return east_asia_font, western_font, size, bool(bold), color


def _effective_line_spacing_pt(doc: Document, paragraph: Paragraph) -> tuple[float | None, str]:
    candidates: list[Any] = []
    if paragraph._p.pPr is not None:
        candidates.append(paragraph._p.pPr)
    paragraph_style = getattr(paragraph, "style", None)
    if paragraph_style is not None and getattr(paragraph_style, "element", None) is not None:
        if paragraph_style.element.pPr is not None:
            candidates.append(paragraph_style.element.pPr)
    normal = doc.styles["Normal"]
    if normal.element.pPr is not None:
        candidates.append(normal.element.pPr)
    for properties in candidates:
        spacing = properties.find(qn("w:spacing"))
        if spacing is None or spacing.get(qn("w:line")) is None:
            continue
        try:
            return int(spacing.get(qn("w:line"))) / 20, str(
                spacing.get(qn("w:lineRule")) or "auto"
            ).lower()
        except (TypeError, ValueError):
            continue
    return None, ""


def _effective_first_line_indent(
    doc: Document, paragraph: Paragraph
) -> tuple[float | None, str]:
    """Return the effective indent value and its OOXML unit contract."""

    candidates: list[Any] = []
    if paragraph._p.pPr is not None:
        candidates.append(paragraph._p.pPr)
    paragraph_style = getattr(paragraph, "style", None)
    if paragraph_style is not None and getattr(paragraph_style, "element", None) is not None:
        if paragraph_style.element.pPr is not None:
            candidates.append(paragraph_style.element.pPr)
    normal = doc.styles["Normal"]
    if normal.element.pPr is not None:
        candidates.append(normal.element.pPr)

    for properties in candidates:
        indent = properties.find(qn("w:ind"))
        if indent is None:
            continue
        chars = indent.get(qn("w:firstLineChars"))
        points = indent.get(qn("w:firstLine"))
        hanging = indent.get(qn("w:hanging")) or indent.get(qn("w:hangingChars"))
        if chars is not None:
            try:
                value = int(chars) / 100
            except (TypeError, ValueError):
                return None, "invalid-character"
            return value, "mixed" if points is not None or hanging is not None else "character"
        if points is not None:
            try:
                return int(points) / 20, "point"
            except (TypeError, ValueError):
                return None, "invalid-point"
        if hanging is not None:
            return None, "hanging"
    return None, "missing"


def _validate_paragraph_typography(
    doc: Document,
    paragraph: Paragraph,
    *,
    label: str,
    kind: str,
    findings: Findings,
    font_override: str | None = None,
    size_override: float | None = None,
    bold_override: bool | None = None,
    line_override: float | None = None,
    western_override: str = DEFAULT_WESTERN_FONT,
    first_line_chars_override: float | None = None,
) -> None:
    expected_font, expected_size, expected_bold, expected_line = TEMPLATE_TYPOGRAPHY[kind]
    if font_override is not None:
        expected_font = font_override
    if size_override is not None:
        expected_size = size_override
    if bold_override is not None:
        expected_bold = bold_override
    if line_override is not None:
        expected_line = line_override
    runs = [run for run in paragraph.runs if run.text.strip()]
    if not runs:
        findings.error(f"DOCX typography check found no visible run for {label}")
        return
    for run in runs:
        east_asia_font, western_font, size, bold, _color = _effective_run_format(
            doc, paragraph, run
        )
        if east_asia_font != expected_font:
            findings.error(
                f"DOCX {label} uses East Asian font {east_asia_font or '<missing>'}; expected {expected_font}"
            )
        if western_font != western_override:
            findings.error(
                f"DOCX {label} uses Western font {western_font or '<missing>'}; "
                f"expected {western_override}"
            )
        if size is None or abs(size - expected_size) > 0.1:
            actual_size = "<missing>" if size is None else f"{size:g} pt"
            findings.error(f"DOCX {label} uses size {actual_size}; expected {expected_size:g} pt")
        if bold != expected_bold:
            findings.error(
                f"DOCX {label} bold={str(bold).lower()}; expected {str(expected_bold).lower()}"
            )
    if expected_line is not None:
        line_pt, line_rule = _effective_line_spacing_pt(doc, paragraph)
        if line_pt is None or abs(line_pt - expected_line) > 0.1 or line_rule != "exact":
            actual_line = "<missing>" if line_pt is None else f"{line_pt:g} pt/{line_rule or '<missing>'}"
            findings.error(
                f"DOCX {label} line spacing is {actual_line}; expected exactly {expected_line:g} pt"
            )
    if first_line_chars_override is not None:
        indent_value, indent_unit = _effective_first_line_indent(doc, paragraph)
        if (
            indent_unit != "character"
            or indent_value is None
            or abs(indent_value - first_line_chars_override) > 0.01
        ):
            actual_indent = "<missing>" if indent_value is None else f"{indent_value:g}"
            findings.error(
                f"DOCX {label} first-line indent is {actual_indent}/{indent_unit}; "
                f"expected {first_line_chars_override:g} characters via w:firstLineChars"
            )


def validate_docx_typography(
    doc: Document,
    expected_spec: dict[str, Any],
    findings: Findings,
) -> None:
    """Validate the key source-template typography that protects page integrity."""

    metadata = expected_spec.get("document") or {}
    company = str(metadata.get("company") or "").strip()
    year = str(metadata.get("report_year") or "").strip()
    paragraphs = list(doc.paragraphs)

    def find_paragraph(text: str, start: int = 0) -> tuple[int, Paragraph] | None:
        target = normalize_assertion_text(text)
        for index in range(start, len(paragraphs)):
            if normalize_assertion_text(paragraphs[index].text) == target:
                return index, paragraphs[index]
        return None

    redhead_text = f"{company}文件"
    redhead_match = find_paragraph(redhead_text)
    if redhead_match is None:
        findings.error("DOCX typography check cannot locate the red-head issuer line")
    else:
        _index, redhead = redhead_match
        _validate_paragraph_typography(
            doc, redhead, label="red-head issuer", kind="redhead", findings=findings
        )
        run = next((item for item in redhead.runs if item.text.strip()), None)
        if run is not None:
            properties = run._element.rPr
            width = properties.find(qn("w:w")) if properties is not None else None
            fit_text = properties.find(qn("w:fitText")) if properties is not None else None
            width_value = width.get(qn("w:val")) if width is not None else None
            fit_value = fit_text.get(qn("w:val")) if fit_text is not None else None
            _east_asia, _western, _size, _bold, color = _effective_run_format(
                doc, redhead, run
            )
            if color != "FF0000":
                findings.error(f"DOCX red-head issuer color is {color or '<missing>'}; expected FF0000")
            if width_value != "37" or fit_value != "8195":
                findings.error(
                    "DOCX red-head issuer must preserve w:w=37 and w:fitText=8195"
                )

    title_text = f"{company}关于{year}年度股权投资项目投后情况报告"
    title_match = find_paragraph(title_text)
    if title_match is None:
        findings.error("DOCX typography check cannot locate the two-line main title")
    else:
        _validate_paragraph_typography(
            doc, title_match[1], label="main title", kind="title", findings=findings
        )

    envelope_paragraphs: list[tuple[str, str, dict[str, Any]]] = [
        ("recipient", str(metadata.get("recipient") or "").strip(), {}),
        (
            "opening basis",
            str(metadata.get("legal_basis") or "").strip(),
            {"first_line_chars_override": 2.0},
        ),
        ("issuer signature", str(metadata.get("issuer") or company).strip(), {}),
        ("issue date", str(metadata.get("issue_date") or "").strip(), {}),
    ]
    attachments = expected_spec.get("attachments") or []
    if attachments:
        envelope_paragraphs.append(
            (
                "attachment list item 1",
                f"附件：1.{attachments[0].get('title', '')}",
                {"first_line_chars_override": 2.0},
            )
        )
        envelope_paragraphs.extend(
            (f"attachment list item {number}", f"{number}.{attachment.get('title', '')}", {})
            for number, attachment in enumerate(attachments[1:], start=2)
        )
    for attachment in attachments:
        envelope_paragraphs.append(
            (f"attachment {attachment.get('id')} label", f"附件{attachment.get('number')}", {})
        )
    contact_parts: list[str] = []
    if str(metadata.get("contact_name") or "").strip():
        contact_parts.append(f"联系人：{str(metadata.get('contact_name')).strip()}")
    if str(metadata.get("contact_phone") or "").strip():
        contact_parts.append(f"联系电话：{str(metadata.get('contact_phone')).strip()}")
    if contact_parts:
        envelope_paragraphs.append(
            (
                "contact line",
                f"（{'  '.join(contact_parts)}）",
                {"font_override": "楷体_GB2312"},
            )
        )
    for label, text, overrides in envelope_paragraphs:
        if not text:
            continue
        match = find_paragraph(text)
        if match is not None:
            _validate_paragraph_typography(
                doc,
                match[1],
                label=label,
                kind="p",
                findings=findings,
                **overrides,
            )

    document_row = next(
        (
            table
            for table in doc.tables
            if len(table.rows) == 1
            and len(table.columns) == 2
            and normalize_assertion_text(table.cell(0, 0).text)
            == normalize_assertion_text(str(metadata.get("document_number") or ""))
        ),
        None,
    )
    if document_row is None:
        findings.error("DOCX typography check cannot locate the document-number/signer row")
    else:
        _validate_paragraph_typography(
            doc,
            document_row.cell(0, 0).paragraphs[0],
            label="document number",
            kind="p",
            findings=findings,
        )
        signer_text = str(metadata.get("signer") or "").strip()
        if signer_text:
            _validate_paragraph_typography(
                doc,
                document_row.cell(0, 1).paragraphs[0],
                label="signer",
                kind="p",
                findings=findings,
                font_override="楷体_GB2312",
            )

    printer = str(metadata.get("printer") or "").strip()
    print_date = str(metadata.get("print_date") or "").strip()
    if printer and print_date:
        imprint_row = next(
            (
                table
                for table in doc.tables
                if len(table.rows) == 1
                and len(table.columns) == 2
                and normalize_assertion_text(table.cell(0, 0).text)
                == normalize_assertion_text(printer)
                and normalize_assertion_text(table.cell(0, 1).text)
                == normalize_assertion_text(f"{print_date}印发")
            ),
            None,
        )
        if imprint_row is None:
            findings.error(
                "DOCX typography check cannot locate the 版记 印发机关／印发日期 row"
            )
        else:
            _validate_paragraph_typography(
                doc,
                imprint_row.cell(0, 0).paragraphs[0],
                label="imprint issuer",
                kind="imprint",
                findings=findings,
            )
            _validate_paragraph_typography(
                doc,
                imprint_row.cell(0, 1).paragraphs[0],
                label="imprint date",
                kind="imprint",
                findings=findings,
            )

    cursor = 0
    expected_blocks = list(expected_spec.get("main_blocks") or [])
    for attachment in attachments:
        title = str(attachment.get("title") or "").strip()
        if title:
            title_match = find_paragraph(title, cursor)
            if title_match is not None:
                cursor, paragraph = title_match
                _validate_paragraph_typography(
                    doc,
                    paragraph,
                    label=f"attachment {attachment.get('id')} title",
                    kind="title",
                    findings=findings,
                )
        expected_blocks.extend(attachment.get("blocks") or [])

    cursor = 0
    for block_number, block in enumerate(expected_blocks, start=1):
        kind = str(block.get("type") or "").lower()
        if kind not in {"h1", "h2", "h3", "h4", "p", "caption", "tnote"}:
            continue
        text = str(block.get("text") or "").strip()
        match = find_paragraph(text, cursor)
        if match is None:
            continue
        cursor, paragraph = match
        cursor += 1
        expected_kind = kind
        if kind == "p" and bool(block.get("bold")):
            # The standard body supports explicit bold emphasis; keep the source
            # font/size/spacing while honoring the declared weight.
            _validate_paragraph_typography(
                doc,
                paragraph,
                label=f"spec block {block_number}",
                kind="p",
                findings=findings,
                bold_override=True,
                first_line_chars_override=2.0,
            )
            continue
        _validate_paragraph_typography(
            doc,
            paragraph,
            label=f"spec block {block_number}",
            kind=expected_kind,
            findings=findings,
            first_line_chars_override=(2.0 if kind in {"h1", "h2", "h3", "h4", "p"} else None),
        )

    table_blocks = [
        block
        for attachment in attachments
        for block in attachment.get("blocks") or []
        if str(block.get("type") or "").lower() == "table"
    ]
    unused_tables = list(doc.tables)
    for table_number, block in enumerate(table_blocks, start=1):
        expected_cells = [str(item) for item in block.get("header") or []]
        expected_cells.extend(str(item) for row in block.get("rows") or [] for item in row)
        expected_key = normalize_assertion_text("\u241f".join(expected_cells))
        match_index = next(
            (
                index
                for index, table in enumerate(unused_tables)
                if normalize_assertion_text(
                    "\u241f".join(cell.text for row in table.rows for cell in row.cells)
                )
                == expected_key
            ),
            None,
        )
        if match_index is None:
            continue
        table = unused_tables.pop(match_index)
        for row_number, row in enumerate(table.rows):
            row_properties = row._tr.find(qn("w:trPr"))
            if row_number == 0 and (
                row_properties is None
                or row_properties.find(qn("w:tblHeader")) is None
            ):
                findings.error(f"DOCX table {table_number} header row is not marked to repeat")
            if row_properties is None or row_properties.find(qn("w:cantSplit")) is None:
                findings.error(f"DOCX table {table_number} row {row_number + 1} may split across pages")
            for cell_number, cell in enumerate(row.cells, start=1):
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    _validate_paragraph_typography(
                        doc,
                        paragraph,
                        label=f"table {table_number} row {row_number + 1} cell {cell_number}",
                        kind="tnote",
                        findings=findings,
                        bold_override=row_number == 0,
                    )

    footer_contracts = (
        ("odd", "right", WD_ALIGN_PARAGRAPH.RIGHT, lambda section: section.footer),
        ("even", "left", WD_ALIGN_PARAGRAPH.LEFT, lambda section: section.even_page_footer),
    )
    page_field_count = 0
    for section_number, section in enumerate(doc.sections, start=1):
        for footer_name, edge_name, expected_alignment, resolve_footer in footer_contracts:
            footer = resolve_footer(section)
            page_paragraphs: list[Paragraph] = []
            for paragraph in footer.paragraphs:
                paragraph_has_page = False
                for run_element in paragraph._p.findall(qn("w:r")):
                    instruction = "".join(
                        str(item.text or "")
                        for item in run_element.findall(qn("w:instrText"))
                    )
                    visible = "".join(
                        str(item.text or "") for item in run_element.findall(qn("w:t"))
                    )
                    if re.search(r"\bPAGE\b", instruction):
                        paragraph_has_page = True
                        page_field_count += 1
                    if not visible and not instruction:
                        continue
                    run = Run(run_element, paragraph)
                    east_asia, western, size, bold, _color = _effective_run_format(
                        doc, paragraph, run
                    )
                    expected_font, expected_size, expected_bold, _line = TEMPLATE_TYPOGRAPHY[
                        "footer"
                    ]
                    if (
                        east_asia != expected_font
                        or western != expected_font
                        or size is None
                        or abs(size - expected_size) > 0.1
                        or bold != expected_bold
                    ):
                        findings.error(
                            f"DOCX section {section_number} {footer_name}-page footer run typography "
                            f"is {east_asia or '<missing>'}/{western or '<missing>'}/"
                            f"{('<missing>' if size is None else f'{size:g} pt')}/"
                            f"bold={str(bold).lower()}; expected {expected_font}/{expected_font}/"
                            f"{expected_size:g} pt/bold={str(expected_bold).lower()}"
                        )
                if paragraph_has_page:
                    page_paragraphs.append(paragraph)
            if len(page_paragraphs) != 1:
                findings.error(
                    f"DOCX section {section_number} {footer_name}-page footer must contain "
                    f"exactly one PAGE field; found {len(page_paragraphs)}"
                )
                continue
            page_paragraph = page_paragraphs[0]
            if page_paragraph.alignment != expected_alignment:
                findings.error(
                    f"DOCX section {section_number} {footer_name}-page footer is not aligned "
                    f"to the outside {edge_name} edge"
                )
            visible_text = "".join(
                str(item.text or "") for item in page_paragraph._p.findall(".//" + qn("w:t"))
            )
            if visible_text != "- 1 -":
                findings.error(
                    f"DOCX section {section_number} {footer_name}-page footer format is "
                    f"{visible_text!r}; expected '- 1 -'"
                )
    if page_field_count == 0:
        findings.error("DOCX typography check found no PAGE field run in odd/even footers")


def validate_docx(
    path: Path,
    findings: Findings,
    *,
    template_mode: bool,
    public_safe: bool,
    deny_terms: Iterable[str] = (),
    expected_spec: dict[str, Any] | None = None,
) -> None:
    if not path.is_file():
        findings.error(f"DOCX not found: {path}")
        return
    try:
        doc = Document(path)
    except Exception as exc:
        findings.error(f"DOCX could not be opened: {exc}")
        return
    if len(doc.sections) != 1:
        findings.warning(f"Expected one section; found {len(doc.sections)}")
    section = doc.sections[0]
    layout = expected_spec.get("layout") if expected_spec is not None else {}
    layout = layout if isinstance(layout, dict) else {}
    expected_cm = {
        "page_width": float(layout.get("page_width_cm", 21.0)),
        "page_height": float(layout.get("page_height_cm", 29.7)),
        "top_margin": float(layout.get("top_margin_cm", 3.7)),
        "bottom_margin": float(layout.get("bottom_margin_cm", 3.5)),
        "left_margin": float(layout.get("left_margin_cm", 2.8)),
        "right_margin": float(layout.get("right_margin_cm", 2.6)),
        "header_distance": float(layout.get("header_distance_cm", 1.5)),
        "footer_distance": float(layout.get("footer_distance_cm", 1.75)),
    }
    for attribute, expected in expected_cm.items():
        actual = getattr(section, attribute).cm
        if abs(actual - expected) > 0.06:
            findings.error(f"DOCX {attribute} is {actual:.2f} cm; expected {expected:.2f} cm")

    text_items = list(iter_docx_text(doc))
    full_text = "\n".join(text_items)
    identifier = str(doc.core_properties.identifier or "").strip()
    identifier_match = re.fullmatch(r"soe-post-investment-report:sha256:([0-9a-f]{64})", identifier)
    if identifier_match is None:
        findings.error("DOCX is missing a valid report-spec SHA-256 identifier")
    elif expected_spec is not None:
        expected_fingerprint = spec_fingerprint(expected_spec)
        if identifier_match.group(1) != expected_fingerprint:
            findings.error("DOCX report-spec fingerprint does not match the supplied specification")

    if expected_spec is not None:
        validate_scoped_docx_content(doc, expected_spec, findings)
        validate_docx_typography(doc, expected_spec, findings)
        normalized_doc_text = re.sub(r"[\s\u3000]+", "", full_text)
        document_metadata = expected_spec.get("document") or {}
        for key in (
            "company",
            "report_year",
            "cutoff_date",
            "document_number",
            "signer",
            "recipient",
            "legal_basis",
            "issuer",
            "issue_date",
            "contact_name",
            "contact_phone",
            "printer",
            "print_date",
        ):
            value = str(document_metadata.get(key) or "").strip()
            if value and re.sub(r"[\s\u3000]+", "", value) not in normalized_doc_text:
                findings.error(f"DOCX does not contain expected document.{key}: {value}")

        expected_blocks: list[tuple[str, dict[str, Any]]] = [
            ("main body", block) for block in expected_spec.get("main_blocks") or []
        ]
        for attachment in expected_spec.get("attachments") or []:
            attachment_id = str(attachment.get("id") or "<unknown>")
            number = str(attachment.get("number") or "")
            title = str(attachment.get("title") or "").strip()
            if number and re.sub(r"[\s\u3000]+", "", f"附件{number}") not in normalized_doc_text:
                findings.error(f"DOCX is missing attachment page label 附件{number}")
            if title and re.sub(r"[\s\u3000]+", "", title) not in normalized_doc_text:
                findings.error(f"DOCX is missing attachment title {title}")
            expected_blocks.extend(
                (f"attachment {attachment_id}", block) for block in attachment.get("blocks") or []
            )

        for scope, block in expected_blocks:
            block_type = str(block.get("type") or "").lower()
            values: list[str]
            if block_type == "table":
                values = [str(item) for item in block.get("header") or []]
                values.extend(str(item) for row in block.get("rows") or [] for item in row)
            elif block_type in {"h1", "h2", "h3", "h4", "p", "caption", "tnote"}:
                values = [str(block.get("text") or "")]
            else:
                values = []
            for value in values:
                normalized_value = re.sub(r"[\s\u3000]+", "", value)
                if normalized_value and normalized_value not in normalized_doc_text:
                    findings.error(f"DOCX is missing expected {scope} content: {value[:80]}")

        for project in expected_spec.get("project_registry") or []:
            official_name = str(project.get("official_name") or "").strip()
            if official_name and re.sub(r"[\s\u3000]+", "", official_name) not in normalized_doc_text:
                findings.error(f"DOCX is missing registry project name {official_name}")

    main_headings: list[str] = []
    main_heading_entries: list[tuple[str, str]] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        # The generated attachment page begins with the standalone label "附件1".
        # The earlier attachment list begins with "附件：1." and must remain inside
        # the main-body scope, so do not use a broad startswith("附件") boundary.
        if re.match(r"^附件\s*[：:]\s*1(?:[.．、]|\s|$)", text) or re.fullmatch(r"附件\s*(?:1|一)", text):
            break
        if HEADING_PATTERN.match(text):
            main_headings.append(text)
            main_heading_entries.append((paragraph_heading_kind(paragraph), text))
    validate_heading_scope(main_headings, "DOCX main body", findings)
    docx_required_blocks = (
        required_main_blocks_from_document(expected_spec.get("document") or {})
        if expected_spec is not None
        else infer_fixed_main_blocks(main_heading_entries)
    )
    validate_required_heading_sequence(
        main_heading_entries,
        "DOCX main body",
        findings,
        required_blocks=docx_required_blocks,
    )
    validate_consecutive_heading_numbers(main_heading_entries, "DOCX main body", findings)
    if not template_mode:
        placeholders = sorted(set(find_placeholders(full_text)))
        if placeholders:
            findings.error(f"Unresolved placeholders in DOCX: {', '.join(placeholders[:10])}")
        template_markers = sorted(set(match.group(0) for match in TEMPLATE_MARKER_PATTERN.finditer(full_text)))
        if template_markers:
            findings.error(f"Template-only markers in production DOCX: {', '.join(template_markers[:10])}")

    with ZipFile(path) as archive:
        names = set(archive.namelist())
        lower_names = {name.lower(): name for name in names}
        if any(name.endswith("vbaproject.bin") for name in lower_names):
            findings.error("DOCX contains a macro project")
        embedded = sorted(name for name in names if name.startswith("word/embeddings/"))
        if embedded:
            findings.error(f"DOCX contains embedded source objects: {', '.join(embedded)}")
        active_x = sorted(name for name in names if name.lower().startswith("word/activex/"))
        if active_x:
            findings.error(f"DOCX contains ActiveX content: {', '.join(active_x)}")

        xml_names = sorted(name for name in names if name.lower().endswith(".xml"))
        oversized_xml = [
            name for name in xml_names if archive.getinfo(name).file_size > 64 * 1024 * 1024
        ]
        if oversized_xml:
            findings.error(f"DOCX contains oversized XML parts: {', '.join(oversized_xml[:10])}")
        xml_payloads = {
            name: archive.read(name) for name in xml_names if name not in oversized_xml
        }
        package_text = "\n".join(xml_text(payload) for payload in xml_payloads.values())
        raw_word_xml = "\n".join(
            payload.decode("utf-8", errors="ignore")
            for name, payload in xml_payloads.items()
            if name.startswith("word/")
        )
        if re.search(r"<w:altChunk\b", raw_word_xml, re.IGNORECASE):
            findings.error("DOCX contains altChunk content")
        if re.search(r"\bDDE(?:AUTO)?\b", package_text, re.IGNORECASE):
            findings.error("DOCX contains a DDE field")

        external_relationships: list[str] = []
        for relationship_name in sorted(name for name in names if name.endswith(".rels")):
            try:
                root = ET.fromstring(archive.read(relationship_name))
            except ET.ParseError:
                findings.error(f"DOCX relationship part could not be parsed: {relationship_name}")
                continue
            for relationship in root:
                if str(relationship.attrib.get("TargetMode") or "").lower() == "external":
                    target = str(relationship.attrib.get("Target") or "<empty>")
                    external_relationships.append(f"{relationship_name} -> {target}")
        if external_relationships:
            findings.error(
                "DOCX contains external relationships: " + "; ".join(external_relationships[:10])
            )

        header_parts = sorted(
            name
            for name in names
            if name.startswith("word/header") and name.endswith(".xml")
        )
        for header_name in header_parts:
            try:
                header_root = ET.fromstring(archive.read(header_name))
            except ET.ParseError:
                findings.error(f"DOCX header part could not be parsed: {header_name}")
                continue
            header_text = "".join(
                str(node.text or "") for node in header_root.iter(qn("w:t"))
            ).strip()
            has_visual_content = any(
                header_root.find(".//" + qn(tag)) is not None
                for tag in ("w:drawing", "w:pict", "w:object", "w:fldChar")
            )
            if header_text or has_visual_content:
                findings.error(
                    f"DOCX header must contain no content: {header_name}"
                )

        footer_xml = "".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )
        if "PAGE" not in footer_xml:
            findings.error("DOCX footer is missing a PAGE field")
        settings_xml = archive.read("word/settings.xml").decode("utf-8", errors="ignore") if "word/settings.xml" in names else ""
        if "evenAndOddHeaders" not in settings_xml:
            findings.error("DOCX does not enable odd/even outside footers")
        if "attachedTemplate" in settings_xml:
            findings.error("DOCX settings contain an attached template reference")

        if public_safe:
            embedded_fonts = sorted(
                name
                for name in names
                if name.lower().startswith("word/fonts/") or name.lower().endswith(".odttf")
            )
            font_table_xml = (
                archive.read("word/fontTable.xml").decode("utf-8", errors="ignore")
                if "word/fontTable.xml" in names
                else ""
            )
            if embedded_fonts or re.search(r"<w:embed(?:Regular|Bold|Italic|BoldItalic)\b", font_table_xml):
                findings.error("Public artifact contains embedded font data")
            comments = sorted(
                name for name in names if name.startswith("word/comments") and name.endswith(".xml")
            )
            if comments:
                findings.error(f"Public artifact contains comments: {', '.join(comments)}")
            for label, pattern in PUBLIC_PATTERNS.items():
                hits = sorted(set(pattern.findall(package_text)))
                if hits:
                    if label in {"mobile phone number", "landline phone number", "email address"}:
                        findings.error(f"Possible {label}s in public artifact: {', '.join(hits[:10])}")
                    else:
                        findings.error(f"Possible {label} in public artifact")
            for term in sorted({str(item).strip() for item in deny_terms if str(item).strip()}):
                if term in package_text:
                    findings.error(f"Public artifact contains deny-listed term: {term}")
    findings.note(f"DOCX: {len(doc.paragraphs)} body paragraphs, {len(doc.tables)} tables")


def validate_pdf(
    path: Path,
    findings: Findings,
    *,
    template_mode: bool,
    expected_spec: dict[str, Any] | None = None,
) -> None:
    if not path.is_file():
        findings.error(f"PDF not found: {path}")
        return
    try:
        from pypdf import PdfReader
    except ImportError:
        findings.error("pypdf is required for PDF page validation: python -m pip install pypdf")
        return
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        findings.error(f"PDF could not be opened: {exc}")
        return
    total_pages = len(reader.pages)
    page_texts = [page.extract_text() or "" for page in reader.pages]
    content_page_texts = [
        re.sub(r"(?m)^\s*-\s*\d+\s*-\s*$", "", page_text)
        for page_text in page_texts
    ]
    # Remove standalone footer page numbers before concatenation so a paragraph
    # split across pages can still be reconciled as one continuous spec block.
    normalized_pdf_text = re.sub(r"[\s\u3000]+", "", "\n".join(content_page_texts))
    if expected_spec is None:
        message = "PDF content certification requires --spec so the rendered content can be reconciled"
        if template_mode:
            findings.warning(message)
        else:
            findings.error(message)
    else:
        document_metadata = expected_spec.get("document") or {}
        for key in (
            "company",
            "report_year",
            "cutoff_date",
            "document_number",
            "signer",
            "recipient",
            "legal_basis",
            "issuer",
            "issue_date",
            "contact_name",
            "contact_phone",
            "printer",
            "print_date",
        ):
            value = str(document_metadata.get(key) or "").strip()
            if value and re.sub(r"[\s\u3000]+", "", value) not in normalized_pdf_text:
                findings.error(f"PDF does not contain expected document.{key}: {value}")
        for scope, blocks in [
            ("main body", expected_spec.get("main_blocks") or []),
            *[
                (f"attachment {attachment.get('id')}", attachment.get("blocks") or [])
                for attachment in expected_spec.get("attachments") or []
            ],
        ]:
            for block in blocks:
                block_type = str(block.get("type") or "").lower()
                if block_type == "table":
                    values = [str(item) for item in block.get("header") or []]
                    values.extend(str(item) for row in block.get("rows") or [] for item in row)
                elif block_type in {"h1", "h2", "h3", "h4", "p", "caption", "tnote"}:
                    values = [str(block.get("text") or "")]
                else:
                    values = []
                for value in values:
                    normalized_value = re.sub(r"[\s\u3000]+", "", value)
                    if normalized_value and normalized_value not in normalized_pdf_text:
                        findings.error(f"PDF is missing expected {scope} content: {value[:80]}")
        for attachment in expected_spec.get("attachments") or []:
            number = str(attachment.get("number") or "")
            title = str(attachment.get("title") or "").strip()
            if number and re.sub(r"[\s\u3000]+", "", f"附件{number}") not in normalized_pdf_text:
                findings.error(f"PDF is missing attachment page label 附件{number}")
            if title and re.sub(r"[\s\u3000]+", "", title) not in normalized_pdf_text:
                findings.error(f"PDF is missing attachment title {title}")

    first_attachment_page: int | None = None
    for page_number, page_text in enumerate(page_texts, start=1):
        lines = [re.sub(r"[\s\u3000]+", "", line) for line in page_text.splitlines()]
        # Match the standalone attachment-page label only. Main-body prose may
        # legitimately contain references such as "详见附件1" and the attachment
        # list contains "附件：1."; neither marks the page boundary.
        if any(re.fullmatch(r"附件(?:一|1)", line) for line in lines):
            first_attachment_page = page_number
            break
    if first_attachment_page is None:
        message = "Could not locate the standalone 附件1 page label in extracted PDF text; the 10-page main-body limit is unproven"
        if template_mode:
            findings.warning(message)
        else:
            findings.error(message)
    else:
        main_pages = first_attachment_page - 1
        if main_pages > 10:
            findings.error(f"Main body is {main_pages} pages; hard maximum is 10")
        elif main_pages not in {5, 6}:
            findings.warning(f"Main body is {main_pages} pages; normal target is 5–6")
        findings.note(f"Rendered pages: main body {main_pages}, total {total_pages}")
    if total_pages == 0:
        findings.error("PDF contains no pages")


def validate_render_manifest(
    path: Path,
    *,
    docx: Path,
    pdf: Path,
    findings: Findings,
) -> None:
    if not path.is_file():
        findings.error(f"Render-binding manifest not found: {path}")
        return
    try:
        payload = json.loads(path.read_text(encoding="utf-8-sig"))
    except Exception as exc:
        findings.error(f"Render-binding manifest could not be read: {exc}")
        return
    if not isinstance(payload, dict) or str(payload.get("schema_version") or "") != "1.0":
        findings.error("Render-binding manifest schema_version must be 1.0")
        return
    if str(payload.get("input_name") or "") != docx.name:
        findings.error("Render-binding manifest input_name does not match the DOCX")
    if str(payload.get("pdf_name") or "") != pdf.name:
        findings.error("Render-binding manifest pdf_name does not match the PDF")
    try:
        actual_docx_hash = file_sha256(docx)
        actual_pdf_hash = file_sha256(pdf)
    except OSError as exc:
        findings.error(f"Render-binding hash could not be computed: {exc}")
        return
    if str(payload.get("input_sha256") or "") != actual_docx_hash:
        findings.error("Rendered PDF manifest does not match the validated DOCX hash")
    if str(payload.get("pdf_sha256") or "") != actual_pdf_hash:
        findings.error("Rendered PDF manifest does not match the supplied PDF hash")
    if str(payload.get("renderer") or "") not in {"word", "libreoffice"}:
        findings.error("Render-binding manifest has an unsupported renderer")

    try:
        page_count = int(payload.get("pages"))
    except (TypeError, ValueError):
        page_count = 0
        findings.error("Render-binding manifest pages must be a positive integer")
    if page_count <= 0:
        findings.error("Render-binding manifest pages must be a positive integer")
    try:
        from pypdf import PdfReader

        actual_page_count = len(PdfReader(str(pdf)).pages)
    except Exception as exc:
        actual_page_count = 0
        findings.error(f"Actual PDF page count could not be verified for render binding: {exc}")
    if actual_page_count and page_count != actual_page_count:
        findings.error(
            f"Render-binding manifest pages ({page_count}) do not match the PDF ({actual_page_count})"
        )

    png_files = payload.get("png_files")
    if not isinstance(png_files, list) or not png_files:
        findings.error("Render-binding manifest must contain page PNG hashes for visual inspection")
    else:
        if page_count > 0 and len(png_files) != page_count:
            findings.error(
                f"Render-binding manifest has {len(png_files)} PNGs for {page_count} PDF pages"
            )
        if actual_page_count and len(png_files) != actual_page_count:
            findings.error(
                f"Render-binding manifest has {len(png_files)} PNGs for {actual_page_count} actual PDF pages"
            )
        manifest_root = path.parent.resolve()
        resolved_pngs: list[Path] = []
        page_numbers: list[int] = []
        for index, item in enumerate(png_files, start=1):
            if not isinstance(item, dict):
                findings.error(f"Render-binding PNG entry {index} must be an object")
                continue
            relative = Path(str(item.get("path") or ""))
            if not str(relative) or relative.is_absolute():
                findings.error(f"Render-binding PNG entry {index} must use a relative path")
                continue
            resolved = (manifest_root / relative).resolve()
            try:
                resolved.relative_to(manifest_root)
            except ValueError:
                findings.error(f"Render-binding PNG entry {index} escapes the manifest directory")
                continue
            resolved_pngs.append(resolved)
            page_match = re.search(r"-page-(\d+)\.png$", resolved.name, re.IGNORECASE)
            if page_match is None:
                findings.error(f"Render-binding PNG filename has no page number: {resolved.name}")
            else:
                page_numbers.append(int(page_match.group(1)))
            if not resolved.is_file():
                findings.error(f"Render-binding PNG file not found: {resolved}")
                continue
            try:
                with resolved.open("rb") as handle:
                    png_header = handle.read(24)
            except OSError as exc:
                findings.error(f"Render-binding PNG could not be read: {resolved.name}: {exc}")
                continue
            if (
                len(png_header) < 24
                or png_header[:8] != b"\x89PNG\r\n\x1a\n"
                or png_header[12:16] != b"IHDR"
                or int.from_bytes(png_header[16:20], "big") <= 0
                or int.from_bytes(png_header[20:24], "big") <= 0
            ):
                findings.error(f"Render-binding page image is not a valid PNG header: {resolved.name}")
            if str(item.get("sha256") or "") != file_sha256(resolved):
                findings.error(f"Render-binding PNG hash mismatch: {resolved.name}")
        if len(resolved_pngs) != len(set(resolved_pngs)):
            findings.error("Render-binding manifest contains duplicate PNG paths")
        expected_page_numbers = list(range(1, actual_page_count + 1)) if actual_page_count else []
        if expected_page_numbers and page_numbers != expected_page_numbers:
            findings.error(
                f"Render-binding PNG page sequence is invalid: expected {expected_page_numbers}, got {page_numbers}"
            )
    findings.note(f"Render binding: {docx.name} -> {pdf.name} ({page_count or 'unknown'} pages)")


def print_findings(findings: Findings) -> None:
    for label, values in (("ERROR", findings.errors), ("WARNING", findings.warnings), ("NOTE", findings.notes)):
        for value in values:
            print(f"[{label}] {value}")
    print(f"Summary: {len(findings.errors)} error(s), {len(findings.warnings)} warning(s), {len(findings.notes)} note(s)")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--spec", type=Path, help="Report specification JSON")
    parser.add_argument("--docx", type=Path, help="Generated DOCX")
    parser.add_argument("--pdf", type=Path, help="Rendered PDF")
    parser.add_argument(
        "--render-manifest",
        type=Path,
        help="DOCX-to-PDF binding manifest (default: <pdf>.render.json)",
    )
    parser.add_argument(
        "--template-mode",
        action="store_true",
        help="Allow visible placeholders only when the specification declares template_only=true",
    )
    parser.add_argument(
        "--public-safe",
        action="store_true",
        help="Run heuristic public-artifact checks across body, headers, footers, properties, and relationships",
    )
    parser.add_argument(
        "--deny-term",
        action="append",
        default=[],
        help="Reject an exact sensitive term in the DOCX package; repeat as needed",
    )
    parser.add_argument(
        "--denylist",
        type=Path,
        action="append",
        default=[],
        help="UTF-8 text file of exact sensitive terms, one per line; repeat as needed",
    )
    return parser.parse_args()


def main() -> int:
    configure_utf8_stdio()
    args = parse_args()
    if not any((args.spec, args.docx, args.pdf)):
        raise SystemExit("Provide at least one of --spec, --docx, or --pdf")
    findings = Findings()
    spec: dict[str, Any] | None = None
    template_mode_active = False
    if args.template_mode and args.spec is None:
        findings.error("--template-mode requires --spec with template_only=true")
    deny_terms = list(args.deny_term)
    for denylist_path in args.denylist:
        try:
            deny_terms.extend(
                line.strip()
                for line in denylist_path.read_text(encoding="utf-8-sig").splitlines()
                if line.strip() and not line.lstrip().startswith("#")
            )
        except Exception as exc:
            findings.error(f"Denylist could not be read ({denylist_path}): {exc}")
    if args.spec:
        try:
            loaded_spec = json.loads(args.spec.read_text(encoding="utf-8-sig"))
        except Exception as exc:
            findings.error(f"Specification could not be read: {exc}")
        else:
            if not isinstance(loaded_spec, dict):
                findings.error("Specification root must be a JSON object")
            else:
                spec = loaded_spec
        if spec is not None:
            validate_spec(spec, findings, template_mode=args.template_mode)
            template_mode_active = args.template_mode and spec.get("template_only") is True
    if args.docx:
        validate_docx(
            args.docx,
            findings,
            template_mode=template_mode_active,
            public_safe=args.public_safe or bool(deny_terms),
            deny_terms=deny_terms,
            expected_spec=spec,
        )
    if args.pdf:
        validate_pdf(args.pdf, findings, template_mode=template_mode_active, expected_spec=spec)
        if args.docx is None:
            findings.error("PDF certification requires --docx and a render-binding manifest")
        else:
            manifest_path = args.render_manifest or args.pdf.with_suffix(args.pdf.suffix + ".render.json")
            validate_render_manifest(
                manifest_path,
                docx=args.docx,
                pdf=args.pdf,
                findings=findings,
            )
    elif args.render_manifest is not None:
        findings.error("--render-manifest requires --pdf and --docx")
    print_findings(findings)
    return 1 if findings.errors else 0


if __name__ == "__main__":
    sys.exit(main())
