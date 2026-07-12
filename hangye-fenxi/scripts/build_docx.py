#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_docx.py —— 立项报告章节 Word 生成器（三技能统一公文排版 v2）

本脚本是 hangye-fenxi / zhuying-yewu-fenxi / gongsi-qingkuang 三个技能共用的
排版渲染器，三份拷贝内容完全一致；修改任何一份时必须同步另外两份。
统一公文版式基准（依据集团《关于规范行文格式的通知》所附模板 + 用户表格规则）：

- A4 页面（上 3.7cm、下 3.5cm、左 2.8cm、右 2.6cm）
- 封面主标题（方正小标宋简体二号 22pt、居中、固定行距 30 磅）；默认不生成封面
- 正文（仿宋_GB2312 三号 16pt、两端对齐、首行缩进 2 字符、固定行距 28 磅）
- 四级编号标题（一、/（一）/ 1. /（1））均首行缩进 2 字符、与段落平齐：
  一级黑体三号不加粗；二级楷体_GB2312 三号加粗；三级仿宋_GB2312 三号加粗；
  四级仿宋_GB2312 三号不加粗（集团规范：四级标题及正文不加粗，三技能统一）
- 核心结论/段首论点句：整段加粗（type=p, bold=true）
- 表格（统一表格规范）：
    * 全表统一仿宋_GB2312、五号 10.5pt
    * 仅首行（表头）加粗，浅蓝底 #D9E2F3；其余单元格不加粗、无底纹
    * 所有单元格内容水平居中 + 垂直居中
    * 细灰边框 #BFBFBF
    * 表头行跨页重复（tblHeader）
    * 宽度按窗口自动调整（pct 100% + tblLayout=autofit），随页宽缩放不溢出；
      table 块的 widths 作为列宽比例（转百分比），不写死磅值
- 表注（type=tnote）："单位：万元""注：……"等，仿宋_GB2312 五号、不缩进，
  align 可选 left/right/center（默认 left；"单位"行惯例放表格上方右对齐）
- 页脚页码（奇偶页外侧，四号宋体，格式 -1-）
- 封面 / 目录：默认关闭（章节通常并入整份立项报告）；独立成文时在 content
  里给 cover 或把 toc 设为 true

数据缺口纪律：资料包、会议纪要、权威网络检索均无法获得的数据，一律不进入
本脚本渲染的报告正文与表格（也不要用 note 块写"待核查"占位）；全部缺口在
对话回复中以"资料缺口与待核查清单"形式提示用户。note 块仅在用户明确要求
在文中标注缺口时使用。

设计原则：脚本只管"排版"，不管"写作"。调用方负责把写好的章节内容按下面的
BLOCK 结构组织好传进来，脚本对内容零编造。

生成前先运行 `python scripts/ensure_fonts.py` 确保公文字体已安装。

依赖：python-docx  →  pip install python-docx --break-system-packages

用法：
    python build_docx.py content.json out.docx
或在 python 中：
    from build_docx import build
    build(content_dict, "out.docx")

content 结构（dict）：
{
  "cover": {                              # 可选，默认不生成封面
     "title_lines": ["关于投资XX公司的", "立项报告"],
     "org": "XX投资管理有限公司",
     "date": "2026年7月"
  },
  "toc": false,                           # 是否插入目录域，默认 false
  "summary": "要点概述正文（可选，多段用 \\n 分隔）",
  "summary_title": "要点概述",            # 可选，默认"要点概述"
  "blocks": [ ...见下... ]                # 章节正文，按顺序渲染
}

blocks 里每个元素是一个 dict，type 决定渲染方式：
  {"type":"h1","text":"二、所属行业分析"}          # 一级标题（自带"二、"前缀）
  {"type":"h2","text":"（一）行业发展现状"}          # 二级
  {"type":"h3","text":"1.市场规模与增长趋势"}        # 三级
  {"type":"h4","text":"（1）政策驱动"}              # 四级（不加粗）
  {"type":"p","text":"正文段落……"}                 # 普通段落
  {"type":"p","text":"……","bold":true}            # 加粗段落（核心结论/段首论点句）
  {"type":"bullet","items":["要点1","要点2"]}       # 项目符号列表
  {"type":"tnote","text":"单位：万元","align":"right"}   # 表注，仿宋五号
  {"type":"table","header":["列1","列2"],           # 表格；header 可省略（无表头）
      "rows":[["a","b"],["c","d"]]}                 # autofit 到窗口，列等分
  {"type":"table","header":[...],"rows":[...],
      "widths":[3,6]}                               # 可选：列宽比例（随窗口缩放）
  {"type":"note","text":"【待进一步核实】"}          # 灰色提示；默认不使用（见缺口纪律）
  {"type":"pagebreak"}                             # 分页

标题层级与编号：脚本不自动编号，"二、""（一）"等前缀由你写在 text 里。
"""

import sys, json
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 版式常量（三技能统一公文格式） ----------
FONT_TITLE = "方正小标宋简体"
FONT_BODY = "仿宋_GB2312"
FONT_H1 = "黑体"            # 使用系统原本黑体，不随技能打包替换
FONT_H2 = "楷体_GB2312"
FONT_FOOTER = "宋体"
FONT_EN = "Times New Roman"
BODY_SZ = 16               # 三号
H1_SZ = 16                 # 一级标题黑体三号，不加粗
H2_SZ = 16                 # 二级标题楷体_GB2312 三号，加粗
H3_SZ = 16                 # 三级标题仿宋_GB2312 三号，加粗
H4_SZ = 16                 # 四级标题仿宋_GB2312 三号，不加粗（三技能统一）
COVER_TITLE_SZ = 22        # 二号方正小标宋简体
COVER_ORG_SZ = 16          # 落款/机构三号仿宋_GB2312
FOOTER_SZ = 14             # 四号宋体
BODY_LINE_PT = 28
TITLE_LINE_PT = 30
TABLE_BORDER = "BFBFBF"    # 表格边框灰
HEADER_FILL = "D9E2F3"     # 表头浅蓝底（三技能统一）
CONTENT_WIDTH = 8844       # A4：21cm - 2.8cm - 2.6cm ≈ 15.6cm（DXA≈8844）
TABLE_SZ = 10.5            # 表格统一五号，所有单元格；表注同


def _set_run_font(run, size=BODY_SZ, bold=False, color=None, font_name=FONT_BODY):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_EN
    # 中文字体需要单独设 eastAsia
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), FONT_EN)
    rfonts.set(qn('w:hAnsi'), FONT_EN)
    rfonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_para(doc, text, size=BODY_SZ, bold=False, align=None, color=None,
              space_before=0, space_after=6, outline=None, line=None,
              font_name=FONT_BODY):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line:
        if isinstance(line, (int, float)) and line > 5:
            pf.line_spacing = Pt(line)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        else:
            pf.line_spacing = line
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color, font_name=font_name)
    if outline is not None:
        _set_outline_level(p, outline)
    return p


def _set_outline_level(paragraph, level):
    """给段落打 outlineLevel，使其能进 TOC（0=H1,1=H2,...）。"""
    pPr = paragraph._p.get_or_add_pPr()
    ol = pPr.find(qn('w:outlineLvl'))
    if ol is None:
        ol = OxmlElement('w:outlineLvl')
        pPr.append(ol)
    ol.set(qn('w:val'), str(level))


def _indent_first_line(paragraph, chars=2):
    """正文段落及各级标题首行缩进 2 字符（标题与段落平齐）。"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))


def _style_cell(cell, fill=None, color=TABLE_BORDER, sz=4,
                margins=(60, 60, 100, 100), valign='center'):
    """一次性构建 tcPr 的子元素，严格按 OOXML schema 顺序：
       tcBorders → shd → tcMar → vAlign 。顺序错了 Word 校验会报错。
       valign 默认 center：单元格内容垂直居中（配合段落水平居中，实现表格内容整体居中）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for tag in ('w:tcBorders', 'w:shd', 'w:tcMar', 'w:vAlign'):
        for e in tcPr.findall(qn(tag)):
            tcPr.remove(e)
    # 1) 边框
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)
    # 2) 底色
    if fill:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    # 3) 内边距
    top, bottom, left, right = margins
    m = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        m.append(el)
    tcPr.append(m)
    # 4) 垂直居中（水平居中在 _fill_cell 的段落上设置）
    if valign:
        v = OxmlElement('w:vAlign')
        v.set(qn('w:val'), valign)
        tcPr.append(v)


def _fill_cell(cell, text, bold=False, header=False, size=TABLE_SZ):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中；垂直居中见 _style_cell
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run(str(text) if text is not None else "")
    _set_run_font(run, size=size, bold=bold or header, font_name=FONT_BODY)
    _style_cell(cell, fill=HEADER_FILL if header else None)


def _mark_header_row(row):
    """表头行设 tblHeader：表格跨页时每页重复表头。"""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        th = OxmlElement('w:tblHeader')
        th.set(qn('w:val'), 'true')
        trPr.append(th)


def _add_table(doc, header, rows, widths=None):
    ncols = len(header) if header else (len(rows[0]) if rows else 1)
    if not widths or len(widths) != ncols:
        widths = [CONTENT_WIDTH // ncols] * ncols   # 默认等分，作为列宽比例
    nrows = (1 if header else 0) + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True                 # 根据窗口自动调整表格
    _set_tbl_autofit(table)              # 表宽=窗口 100%（pct）+ tblLayout=autofit
    r = 0
    if header:
        for j, h in enumerate(header):
            _fill_cell(table.rows[0].cells[j], h, header=True, size=TABLE_SZ)
        _mark_header_row(table.rows[0])
        r = 1
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            _fill_cell(table.rows[r + i].cells[j], val, size=TABLE_SZ)
    # 列宽改用百分比：保持 widths 的设计比例，随窗口宽度整体缩放
    total = sum(widths) or 1
    col_pct = [max(1, round(w / total * 5000)) for w in widths]  # 5000 = 100.00%
    for row in table.rows:
        for j, pct in enumerate(col_pct):
            _set_cell_width_pct(row.cells[j], pct)
    return table


def _set_tbl_autofit(table):
    """表格宽度设为窗口（页面内容区）宽度的 100%，并启用 autofit 布局，
       实现 Word『根据窗口自动调整表格』。tblW 由 add_table 生成，改其属性即可，
       位置天然合法（tblStyle 之后、jc 之前）。"""
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')          # 5000 = 100.00%
    tblW.set(qn('w:type'), 'pct')
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'autofit')


def _set_cell_width_pct(cell, pct):
    """单元格首选宽度用百分比（tcW type=pct），配合表格 autofit 随窗口缩放。
       tcW 在 CT_TcPr 中须位于最前，故先清已有再 insert(0)。"""
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcW')):
        tcPr.remove(e)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(pct))
    tcW.set(qn('w:type'), 'pct')
    tcPr.insert(0, tcW)


def _add_toc(doc):
    """插入 Word 原生 TOC 域；打开文档后需手动/自动更新（右键→更新域）。"""
    _add_para(doc, "目录", size=H1_SZ, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_after=8, line=BODY_LINE_PT, font_name=FONT_H1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（打开文档后右键此处『更新域』生成目录与页码）"
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    for el in (fldChar1, instr, fldChar2, t, fldChar3):
        run._r.append(el)


def _add_page_number_footer(section):
    _add_footer_page_field(section.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    try:
        _add_footer_page_field(section.even_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)
    except Exception:
        pass


def _add_footer_page_field(p, align):
    p.alignment = align
    run = p.add_run()
    run.add_text("-")
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    for el in (f1, instr, f2):
        run._r.append(el)
    run.add_text("-")
    _set_run_font(run, size=FOOTER_SZ, font_name=FONT_FOOTER)


def _enable_odd_even_footers(doc):
    settings = doc.settings.element
    even = settings.find(qn('w:evenAndOddHeaders'))
    if even is None:
        settings.append(OxmlElement('w:evenAndOddHeaders'))


def _set_a4_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def _fix_zoom(doc):
    """python-docx 默认 settings.xml 里的 w:zoom 可能缺 percent 属性，补上。"""
    try:
        settings = doc.settings.element
        zoom = settings.find(qn('w:zoom'))
        if zoom is not None and zoom.get(qn('w:percent')) is None:
            zoom.set(qn('w:percent'), '100')
    except Exception:
        pass


def build(content, out_path):
    doc = Document()
    _fix_zoom(doc)
    _enable_odd_even_footers(doc)
    sec = doc.sections[0]
    _set_a4_page(sec)
    _add_page_number_footer(sec)

    # 设置 Normal 默认字体
    normal = doc.styles['Normal']
    normal.font.name = FONT_EN
    normal.font.size = Pt(BODY_SZ)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    # ---------- 封面（可选，默认不生成） ----------
    cover = content.get("cover")
    if cover:
        for _ in range(6):
            doc.add_paragraph()
        for line in cover.get("title_lines", ["立项报告"]):
            _add_para(doc, line, size=COVER_TITLE_SZ, bold=False,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
                      line=TITLE_LINE_PT, font_name=FONT_TITLE)
        for _ in range(6):
            doc.add_paragraph()
        if cover.get("org"):
            _add_para(doc, cover["org"], size=COVER_ORG_SZ, bold=False,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
                      line=BODY_LINE_PT, font_name=FONT_BODY)
        if cover.get("date"):
            _add_para(doc, cover["date"], size=COVER_ORG_SZ, bold=False,
                      align=WD_ALIGN_PARAGRAPH.CENTER, line=BODY_LINE_PT,
                      font_name=FONT_BODY)
        doc.add_page_break()

    # ---------- 目录（可选，默认关闭） ----------
    if content.get("toc", False):
        _add_toc(doc)
        doc.add_page_break()

    # ---------- 要点概述（可选） ----------
    if content.get("summary"):
        _add_para(doc, content.get("summary_title", "要点概述"),
                  size=H1_SZ, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_after=8, outline=0, line=BODY_LINE_PT,
                  font_name=FONT_H1)
        for para in str(content["summary"]).split("\n"):
            if para.strip():
                p = _add_para(doc, para.strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              line=BODY_LINE_PT, font_name=FONT_BODY)
                _indent_first_line(p)
        doc.add_page_break()

    # ---------- 章节正文 ----------
    ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
             "center": WD_ALIGN_PARAGRAPH.CENTER}
    for blk in content.get("blocks", []):
        t = blk.get("type")
        if t == "h1":
            p = _add_para(doc, blk["text"], size=H1_SZ, bold=False,
                          space_before=8, space_after=6, outline=0,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_H1)
            _indent_first_line(p)
        elif t == "h2":
            p = _add_para(doc, blk["text"], size=H2_SZ, bold=True,
                          space_before=6, space_after=4, outline=1,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_H2)
            _indent_first_line(p)
        elif t == "h3":
            p = _add_para(doc, blk["text"], size=H3_SZ, bold=True,
                          space_before=4, space_after=4, outline=2,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_BODY)
            _indent_first_line(p)
        elif t == "h4":
            # 四级标题不加粗（集团规范：四级标题及正文不加粗，三技能统一）
            p = _add_para(doc, blk["text"], size=H4_SZ, bold=False,
                          space_before=2, space_after=2, outline=3,
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_BODY)
            _indent_first_line(p)
        elif t == "p":
            p = _add_para(doc, blk["text"], bold=blk.get("bold", False),
                          align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=BODY_LINE_PT,
                          font_name=FONT_BODY)
            _indent_first_line(p)
        elif t == "bullet":
            for item in blk.get("items", []):
                p = doc.add_paragraph(style=None)
                p.paragraph_format.left_indent = Twips(480)
                p.paragraph_format.line_spacing = Pt(BODY_LINE_PT)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                run = p.add_run("• " + str(item))
                _set_run_font(run, size=BODY_SZ, font_name=FONT_BODY)
        elif t == "tnote":
            # 表注："单位：万元""注：……"，仿宋五号、不缩进；单位行惯例右对齐
            _add_para(doc, blk["text"], size=TABLE_SZ, bold=False,
                      align=ALIGN.get(blk.get("align", "left")),
                      space_before=0, space_after=2, line=18,
                      font_name=FONT_BODY)
        elif t == "table":
            _add_table(doc, blk.get("header"), blk.get("rows", []), blk.get("widths"))
            doc.add_paragraph()  # 表后空行
        elif t == "note":
            # 灰色提示：默认不使用——缺口一律在对话中提示，不写入报告（见文件头"数据缺口纪律"）
            _add_para(doc, blk["text"], color="808080", line=BODY_LINE_PT,
                      font_name=FONT_BODY)
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            # 未知类型，按普通段落处理，避免内容丢失
            if blk.get("text"):
                _add_para(doc, blk["text"], align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line=BODY_LINE_PT, font_name=FONT_BODY)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python build_docx.py content.json out.docx")
        sys.exit(1)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        data = json.load(f)
    path = build(data, sys.argv[2])
    print("已生成:", path)
